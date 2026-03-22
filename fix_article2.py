"""
fix_article2.py — финальная полировка (peer review 4й раунд)

1. companion paper — добавить сноску с DOI/статусом
2. cnRNA предсказание #7 — укоротить в таблице, перенести детали в §3.3
3. CAFD терминология — уточнить "CAFD = подкласс IDI"
4. IDI-гипотеза — добавить независимость ядра от конкретного механизма
"""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "CDATA_Theory_Full_Article.docx"
doc = Document(PATH)


def replace_in_para(para, old, new):
    if old not in para.text:
        return False
    full = para.text.replace(old, new, 1)
    for r in para.runs:
        r._r.getparent().remove(r._r)
    run = para.add_run(full)
    return True


def replace_in_cell(cell, old, new):
    for para in cell.paragraphs:
        if old in para.text:
            return replace_in_para(para, old, new)
    return False


def insert_paragraph_after(ref_para, text):
    new_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    ref_para._p.addnext(new_p)
    return new_p


# ═══════════════════════════════════════════════════════════════════
# FIX 1: Companion paper — добавить ссылку [33] и обновить in-text
# ═══════════════════════════════════════════════════════════════════
# Находим все упоминания "companion paper" и добавляем [33]
fixed_companion = 0
for para in doc.paragraphs:
    t = para.text
    # Добавляем [33] только там, где "companion paper" без ссылки
    if 'companion paper' in t and '[33]' not in t:
        new_t = t.replace(
            'companion paper',
            'companion paper [33]',
            1  # только первое вхождение в каждом параграфе
        )
        if new_t != t:
            replace_in_para(para, t, new_t)
            fixed_companion += 1

# Добавляем ссылку [33] в список литературы (после [32])
ref32_para = None
for para in doc.paragraphs:
    if para.text.strip().startswith('[32]'):
        ref32_para = para
        break

if ref32_para:
    REF33 = (
        "[33]  Tkemaladze J. Cell-DT: A Computational Digital Twin for "
        "Simulating Centriolar Damage Accumulation and Stem Cell Aging. "
        "Preprints.org 2026; submitted simultaneously. "
        "doi: 10.20944/preprints202603.xxxx.v1 [companion paper, under review]"
    )
    insert_paragraph_after(ref32_para, REF33)

print(f"FIX 1: companion paper → [33] added ({fixed_companion} in-text + ref entry) ✓")


# ═══════════════════════════════════════════════════════════════════
# FIX 2a: cnRNA предсказание #7 — укоротить в таблице
# ═══════════════════════════════════════════════════════════════════
pred_table = doc.tables[4]
row7 = pred_table.rows[7]

OLD_CNRNA = (
    "Centrosomal RNA (cnRNA) species with reverse-transcriptase domains, "
    "localized to the proximal centriolar scaffold (≤200 nm from the cartwheel), "
    "are detectable by single-centriole RNA-seq in adult mammalian stem cells "
    "(HSCs or NSCs). Note: cnRNA-class molecules were first identified in "
    "mollusks (Alliegro et al.); this prediction extends the finding to adult "
    "mammalian somatic stem cells."
)
NEW_CNRNA_SHORT = (
    "Centrosomal RNA (cnRNA) species with reverse-transcriptase domains are "
    "detectable by single-centriole RNA-seq at the proximal scaffold "
    "(≤200 nm from cartwheel) of adult mammalian HSCs or NSCs "
    "(see §3.3 for molecular context and localization rationale)"
)

r2a = replace_in_cell(row7.cells[0], OLD_CNRNA, NEW_CNRNA_SHORT)
print(f"FIX 2a: cnRNA prediction shortened in table → {r2a} ✓")


# ═══════════════════════════════════════════════════════════════════
# FIX 2b: Перенести cnRNA-детали в §3.3
# Добавить после фразы о молекулярных кандидатах IDI
# ═══════════════════════════════════════════════════════════════════
CAFD_PARA_FRAGMENT = (
    "centriole-associated fate determinant proteins (CAFDs) such as pSTAT3 and YAP/TAZ."
)
CNRNA_DETAIL = (
    "Note on cnRNA candidates: centrosomal RNA species were first experimentally "
    "identified in the surf clam Spisula solidissima (Alliegro et al., 2006), "
    "where a centrosomal RNA fraction (including cnRNA11 with a reverse-transcriptase "
    "domain) was enriched at the spindle apparatus. CDATA extends this observation: "
    "the IDI hypothesis predicts that analogous cnRNA-class molecules should be "
    "detectable at the proximal scaffold of mammalian stem cell centrioles (≤200 nm "
    "from the cartwheel), identifiable by single-centriole RNA-seq and distinguishable "
    "from PCM-associated mRNAs by STED/FISH localization."
)

inserted_cnrna_detail = False
for para in doc.paragraphs:
    if CAFD_PARA_FRAGMENT in para.text:
        insert_paragraph_after(para, CNRNA_DETAIL)
        inserted_cnrna_detail = True
        break

print(f"FIX 2b: cnRNA detail para inserted into §3.3 → {inserted_cnrna_detail} ✓")


# ═══════════════════════════════════════════════════════════════════
# FIX 3: CAFD терминология — уточнить как подкласс IDI
# ═══════════════════════════════════════════════════════════════════
OLD_CAFD = (
    "centriole-associated fate determinant proteins (CAFDs) such as pSTAT3 and YAP/TAZ."
)
NEW_CAFD = (
    "centriole-associated fate determinant proteins (CAFDs — a candidate "
    "protein subclass of IDI molecules) such as pSTAT3 and YAP/TAZ."
)
fixed_cafd = False
for para in doc.paragraphs:
    if OLD_CAFD in para.text:
        replace_in_para(para, OLD_CAFD, NEW_CAFD)
        fixed_cafd = True
        break

print(f"FIX 3: CAFD → 'protein subclass of IDI' clarified → {fixed_cafd} ✓")


# ═══════════════════════════════════════════════════════════════════
# FIX 4: IDI-гипотеза — добавить независимость ядра от механизма
# ═══════════════════════════════════════════════════════════════════
OLD_SPECULATIVE = (
    "The exact mechanism of irreversible epigenetic action remains unknown; "
    "direct DNA integration via cnRNA reverse-transcriptase activity is "
    "mechanistically compelling but speculative; alternative mechanisms — "
    "piRNA-directed CpG methylation, direct chromatin binding of protein CAFDs "
    "— cannot currently be excluded."
)
NEW_SPECULATIVE = (
    "The exact mechanism of irreversible epigenetic action remains unknown; "
    "direct DNA integration via cnRNA reverse-transcriptase activity is "
    "mechanistically compelling but speculative; alternative mechanisms — "
    "piRNA-directed CpG methylation, direct chromatin binding of protein CAFDs "
    "— cannot currently be excluded. Importantly, the core IDI hypothesis — "
    "that two molecularly distinct sets of fate-determining molecules are "
    "differentially depleted (M-set by oxidative damage, D-set by divisional "
    "release) — does not depend on the specific downstream epigenetic mechanism; "
    "the existence, asymmetry, and depletion dynamics of the two sets constitute "
    "the falsifiable core, addressed by predictions (ii), (iii), and (iv) in "
    "Section 8."
)

fixed_speculative = False
for para in doc.paragraphs:
    if OLD_SPECULATIVE in para.text:
        replace_in_para(para, OLD_SPECULATIVE, NEW_SPECULATIVE)
        fixed_speculative = True
        break

print(f"FIX 4: IDI core independence note added → {fixed_speculative} ✓")


# ═══════════════════════════════════════════════════════════════════
# FIX 5: В Hallmarks Track G — "CAFD manipulation" → уточнить
# ═══════════════════════════════════════════════════════════════════
# Track G track column might have "CAFD manipulation"
hallmarks_table = doc.tables[2]
fixed_cafd_table = False
for row in hallmarks_table.rows:
    if 'Hormonal' in row.cells[0].text:
        for cell in row.cells:
            if 'CAFD' in cell.text:
                replace_in_cell(
                    cell,
                    'CAFD manipulation',
                    'CAFD/IDI-targeted manipulation'
                )
                fixed_cafd_table = True
print(f"FIX 5: CAFD in hallmarks table → 'CAFD/IDI-targeted' (if present): {fixed_cafd_table} ✓")


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
doc.save(PATH)
print("\nAll fixes applied. Saved:", PATH)
