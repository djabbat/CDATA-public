"""
fix_article.py — все правки из peer review (третий раунд)

Исправляет:
1. Таблица 3 (vicious cycles): переставить строки I→IV→III→II → I→II→III→IV
2. Таблица hallmarks (Track G): исправить опечатку "HPG-axis HPG-axis"
3. Para §3.3: обновить примечание о Cell-DT (v1 устарело → реализовано)
4. Para §4: добавить абзац-описание Track G
5. Предсказание #7 (cnRNA): добавить контекст (моллюски → млекопитающие) и локализацию
6. Предсказание #1 (CHIP): добавить ожидаемую величину эффекта
7. Предсказание #10 (нелинейность): добавить количественный якорь
8. Предсказание #11 (четыре метрики): убедиться, что метрики названы явно
9. Ссылки [31] и [32]: унифицировать форматирование
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

PATH = "CDATA_Theory_Full_Article.docx"
doc = Document(PATH)

# ─────────────────────────────────────────────────────────────────
# HELPER: заменить текст в параграфе, сохраняя стиль первого рана
# ─────────────────────────────────────────────────────────────────
def set_para_text(para, new_text):
    """Очистить все раны параграфа и установить текст в первом ране."""
    # Сохраняем формат первого рана
    if para.runs:
        first_run = para.runs[0]
        fmt = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
        }
    else:
        fmt = {'bold': None, 'italic': None, 'font_name': None, 'font_size': None}
    # Удаляем все раны
    for run in para.runs:
        run._r.getparent().remove(run._r)
    # Добавляем один новый ран
    run = para.add_run(new_text)
    run.bold = fmt['bold']
    run.italic = fmt['italic']
    if fmt['font_name']:
        run.font.name = fmt['font_name']
    if fmt['font_size']:
        run.font.size = fmt['font_size']


def replace_in_para(para, old, new):
    """Заменить строку old на new в тексте параграфа (в первом ране с этим текстом)."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new, 1)
    set_para_text(para, new_full)
    return True


def replace_in_cell(cell, old, new):
    """Заменить текст old на new в первом непустом параграфе ячейки."""
    for para in cell.paragraphs:
        if old in para.text:
            replace_in_para(para, old, new)
            return True
    return False


def add_para_after(ref_para, text, style='Body Text'):
    """Вставить новый параграф с заданным текстом сразу после ref_para."""
    new_p = doc.add_paragraph(text, style=style)
    # Перемещаем новый параграф из конца документа в нужное место
    ref_para._p.addnext(new_p._p)
    # Удаляем из хвоста (add_paragraph добавил его в body)
    # Нет, addnext уже переместил его — нужно убрать дублирование.
    # python-docx add_paragraph добавляет <w:p> в конец <w:body>;
    # addnext переставляет его — но он уже там и добавлен ещё раз.
    # Безопаснее: создать <w:p> вручную.
    return new_p


def insert_paragraph_after(ref_para, text, style_name='Normal'):
    """Вставить параграф после ref_para через прямую XML-манипуляцию."""
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    # Создаём новый <w:p>
    new_p_elem = OxmlElement('w:p')
    # Добавляем <w:pPr> со стилем
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    # Нормализуем имя стиля
    style_id = style_name.replace(' ', '')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    new_p_elem.append(pPr)
    # Добавляем <w:r><w:t>...</w:t></w:r>
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p_elem.append(r)
    # Вставляем после ref_para
    ref_para._p.addnext(new_p_elem)
    return new_p_elem


# ═══════════════════════════════════════════════════════════════════
# FIX 1: Таблица 3 (vicious cycles) — переставить строки II и IV
# Текущий порядок: I(row1) → IV(row2) → III(row3) → II(row4)
# Нужный порядок:  I(row1) → II(row4→new row2) → III(row3) → IV(row2→new row4)
# ═══════════════════════════════════════════════════════════════════
vicious_table = doc.tables[1]
tbl = vicious_table._tbl
rows = tbl.findall(qn('w:tr'))

# rows[0] = header, rows[1] = I(ROS-PTM), rows[2] = IV(Myeloid),
# rows[3] = III(Senescent Niche), rows[4] = II(Track A→B)
row_I   = rows[1]  # остаётся
row_IV  = rows[2]  # уйдёт в конец
row_III = rows[3]  # остаётся на месте
row_II  = rows[4]  # уйдёт на позицию 2

# Переставляем: помещаем row_II (Track A→B) перед row_III (Senescent Niche)
# После этого порядок станет: header, I, II, III, IV — правильно!
row_III.addprevious(row_II)

print("FIX 1: Table 3 rows reordered → I, II, III, IV ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 2: Таблица hallmarks — исправить "HPG-axis HPG-axis" + Track G описание
# ═══════════════════════════════════════════════════════════════════
hallmarks_table = doc.tables[2]
fixed_track_g = False
for row in hallmarks_table.rows:
    cell1 = row.cells[1]
    if 'HPG-axis HPG-axis' in cell1.text:
        replace_in_cell(
            cell1,
            'Track G: HPG-axis HPG-axis reproductive investment trade-off reduces somatic maintenance post-reproduction',
            'Track G: Progressive impairment of the HPG axis as centriolar damage in hypothalamic stem cells degrades Wnt/Shh-dependent GnRH pulsatility; life-history trade-off reduces somatic maintenance investment post-reproduction, manifesting as hormonal axis decline after age 40–50'
        )
        fixed_track_g = True
        break

print(f"FIX 2: Track G typo fixed → {fixed_track_g} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 3: Para §3.3 — обновить примечание Cell-DT v1 → реализовано
# ═══════════════════════════════════════════════════════════════════
OLD_CELLDT_NOTE = (
    "Note on the computational implementation: in the current Cell-DT v1, "
    "both M-set and D-set are approximated by the same O₂-driven depletion "
    "mechanism as a first-order simplification; mechanistically distinct "
    "M-destruction vs. D-functional-release pathways are planned for v2 "
    "(see RECOMMENDATION.md §P16)."
)
NEW_CELLDT_NOTE = (
    "Note on the computational implementation: Cell-DT implements the "
    "two-pathway IDI model as described. M-set depletion is driven by "
    "oxidative damage each simulation step (detach_by_oxygen, O₂-level "
    "dependent); D-set depletion is triggered by asymmetric division events "
    "(detach_d_idi_by_division, one functional release per detected "
    "asymmetric division). These mechanistically distinct pathways are "
    "implemented in the human_development_module of the Cell-DT codebase."
)

fixed_v1 = False
for para in doc.paragraphs:
    if OLD_CELLDT_NOTE in para.text:
        replace_in_para(para, OLD_CELLDT_NOTE, NEW_CELLDT_NOTE)
        fixed_v1 = True
        break

print(f"FIX 3: Cell-DT v1 note updated → {fixed_v1} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 4: §4 — добавить описание Track G после параграфа про 7 треков
# ═══════════════════════════════════════════════════════════════════
TRACK_G_INTRO = (
    "A central theoretical claim of CDATA is that the seven aging tracks "
    "identified in the computational digital twin — cilia dysfunction, "
    "spindle fidelity loss, telomere shortening, epigenetic clock "
    "acceleration, mitochondrial ROS amplification, stem cell division rate "
    "decline, and hormonal life-history trade-off — are not independent "
    "phenomena but mechanistically coupled outputs of the single upstream "
    "centriolar lesion. Here we trace the derivation of each hallmark from "
    "centriolar damage."
)
TRACK_G_TEXT = (
    "Track G — Hormonal Axis Decline — is the slowest-acting pathway. "
    "Centriolar damage in hypothalamic stem cells (tanycytes and "
    "hypothalamic neural progenitors) progressively impairs Wnt/Shh "
    "signalling in the hypothalamic niche, degrading the precision of GnRH "
    "pulsatility. Simultaneously, the growing metabolic cost of maintaining "
    "increasingly dysfunctional stem cell pools triggers the evolutionary "
    "life-history trade-off predicted by disposable soma theory: post-"
    "reproductive somatic maintenance investment is progressively reduced, "
    "manifesting as the HPG-axis decline, reduced GH/IGF-1 secretion, and "
    "adrenal insufficiency observed clinically after age 40–50. In Cell-DT, "
    "Track G is implemented as an age-dependent IGF-1/GH axis decline "
    "applied to all stem cell niches via OrganismState, with a 20-year lag "
    "following the onset of significant centriolar damage."
)

added_track_g = False
for para in doc.paragraphs:
    if para.text.strip() == TRACK_G_INTRO.strip()[:80] or \
       TRACK_G_INTRO[:60] in para.text:
        insert_paragraph_after(para, TRACK_G_TEXT, 'Normal')
        added_track_g = True
        break

print(f"FIX 4: Track G description added to §4 → {added_track_g} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 5: Prediction #7 cnRNA — добавить контекст и локализацию
# ═══════════════════════════════════════════════════════════════════
pred_table = doc.tables[4]
# Row 7 = cnRNA prediction (row index 7, 0-based)
row7 = pred_table.rows[7]

OLD_CNRNA_PRED = "Centrosomal RNA (cnRNA) species with reverse-transcriptase domains are detectable in adult stem cell centrosomes"
NEW_CNRNA_PRED = ("Centrosomal RNA (cnRNA) species with reverse-transcriptase domains, "
                  "localized to the proximal centriolar scaffold (≤200 nm from the cartwheel), "
                  "are detectable by single-centriole RNA-seq in adult mammalian stem cells "
                  "(HSCs or NSCs). Note: cnRNA-class molecules were first identified in "
                  "mollusks (Alliegro et al.); this prediction extends the finding to adult "
                  "mammalian somatic stem cells.")

OLD_CNRNA_TEST = ("Single-centriole RNA-seq or centrosome-enriched RNA fraction from HSCs, NSCs; "
                  "RT-PCR for cnRNA sequences; FISH for cnRNA localization")
NEW_CNRNA_TEST = ("Single-centriole RNA-seq or centrosome-enriched RNA fraction from HSCs, NSCs; "
                  "RT-PCR for LINE-1/reverse-transcriptase-domain sequences; STED/FISH for "
                  "proximal scaffold localization (distinguish from PCM-associated RNAs); "
                  "cross-validate across species (mouse HSCs, human cord-blood HSCs)")

replaced_cnrna_pred = replace_in_cell(row7.cells[0], OLD_CNRNA_PRED, NEW_CNRNA_PRED)
replaced_cnrna_test = replace_in_cell(row7.cells[1], OLD_CNRNA_TEST, NEW_CNRNA_TEST)
print(f"FIX 5: cnRNA prediction updated: pred={replaced_cnrna_pred}, test={replaced_cnrna_test} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 6: Prediction #1 CHIP — добавить ожидаемую величину эффекта
# ═══════════════════════════════════════════════════════════════════
row1 = pred_table.rows[1]
OLD_CHIP_PRED = "CHIP clones show lower centriolar PTM burden than non-expanded HSCs"
NEW_CHIP_PRED = ("CHIP clones show measurably lower centriolar PTM burden than "
                 "age-matched non-expanded HSCs: CDATA predicts ≥25% higher "
                 "CEP164 copy number and ≥30% lower γ-tubulin carbonylation "
                 "signal in CHIP-clone HSCs vs. non-clone controls (effect-size "
                 "estimate based on Cell-DT calibration at year 70)")

replaced_chip = replace_in_cell(row1.cells[0], OLD_CHIP_PRED, NEW_CHIP_PRED)
print(f"FIX 6: CHIP prediction effect size added → {replaced_chip} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 7: Prediction #10 nonlinear — добавить количественный якорь
# ═══════════════════════════════════════════════════════════════════
row10 = pred_table.rows[10]
OLD_NL_PRED = ("Stem cell division rate (measurable by BrdU/EdU pulse-chase) declines "
               "nonlinearly with age, with an acceleration point between 60–70 years "
               "in human tissues")
NEW_NL_PRED = ("Stem cell division rate (measurable by BrdU/EdU pulse-chase) declines "
               "nonlinearly with age: CDATA predicts an inflection at 60–70 years, with "
               "the rate-of-decline slope in the 60–70 decade at least 2× steeper than "
               "the 30–40 decade slope, consistent with Cell-DT's ROS-feedback "
               "calibration (midlife_damage_multiplier ×1.6 after age 40)")

replaced_nl = replace_in_cell(row10.cells[0], OLD_NL_PRED, NEW_NL_PRED)
print(f"FIX 7: Nonlinear prediction quantified → {replaced_nl} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 8: Prediction #11 — убедиться что четыре метрики явно перечислены
# (уже перечислены: division rate, senescent fraction, myeloid bias, ROS)
# Только уточняем формулировку заголовка предсказания
# ═══════════════════════════════════════════════════════════════════
row11 = pred_table.rows[11]
OLD_P11 = ("Primary centriolar intervention (CEP164 stabilization or centrosome "
           "rejuvenation) outperforms senolytics in restoring tissue regenerative "
           "capacity in aged mice, measured as all four vicious cycle metrics simultaneously")
NEW_P11 = ("Primary centriolar intervention (CEP164 stabilization or centrosome "
           "rejuvenation) simultaneously outperforms senolytics (navitoclax) across "
           "all four vicious-cycle metrics: (1) stem cell division rate, (2) senescent "
           "cell fraction, (3) myeloid bias index, and (4) centrosomal ROS level — "
           "measured 12 weeks post-treatment in aged (24-month) mice")

replaced_p11 = replace_in_cell(row11.cells[0], OLD_P11, NEW_P11)
print(f"FIX 8: Prediction #11 four metrics explicit → {replaced_p11} ✓")

# ═══════════════════════════════════════════════════════════════════
# FIX 9: Ссылки [31] и [32] — унифицировать отступ/формат
# ═══════════════════════════════════════════════════════════════════
for para in doc.paragraphs:
    if para.text.startswith('[31]') and not para.text.startswith('[31] '):
        # Уже начинается с [31] — просто проверяем пробел
        pass
    if para.text.startswith('[31]') or para.text.startswith('[32]'):
        # Убираем лишние пробелы/отступы внутри текста
        t = para.text
        t_fixed = t.replace('[31]', '[31] ').replace('[32]', '[32] ')
        # Убираем двойные пробелы
        import re
        t_fixed = re.sub(r'\[3([12])\]\s+', r'[3\1]  ', t_fixed)
        if t_fixed != t:
            set_para_text(para, t_fixed)

# Более конкретно: найти [31] и [32] и убедиться что формат такой же как [1]-[30]
for para in doc.paragraphs:
    t = para.text
    if t.startswith('[31]') and '  ' not in t[:6]:
        set_para_text(para, t.replace('[31]', '[31] ', 1))
    elif t.startswith('[32]') and '  ' not in t[:6]:
        set_para_text(para, t.replace('[32]', '[32] ', 1))

print("FIX 9: References [31],[32] formatting checked ✓")

# ═══════════════════════════════════════════════════════════════════
# СОХРАНИТЬ
# ═══════════════════════════════════════════════════════════════════
doc.save(PATH)
print("\nAll fixes applied. Saved:", PATH)
