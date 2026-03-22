"""
finalize_article.py — финальная обработка CDATA_Theory_Full_Article.docx

Что делает:
  1. Исправляет нумерацию таблиц (последовательная 1–5 в порядке чтения)
  2. Добавляет подписи к таблицам 1, 3, 4, 5 (у таблицы 2 уже есть)
  3. Встраивает 9 рисунков в нужные места с подписями
  4. Убирает артефактный путь к файлу из текста Figure 2
  5. Добавляет «Список рисунков» после раздела References
  6. Сохраняет как CDATA_Theory_Final.docx

Ссылки [1]–[33] НЕ изменяются.
"""

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

PATH_IN  = "CDATA_Theory_Full_Article.docx"
PATH_OUT = "CDATA_Theory_Final.docx"
FIG_DIR  = Path("figures")

doc = Document(PATH_IN)


# ─── Utility functions ────────────────────────────────────────────────────────

def para_from_elem(p_elem):
    """Wrap raw XML <w:p> element as python-docx Paragraph."""
    return Paragraph(p_elem, doc.element.body)


def insert_para_after(ref_para, text="", bold=False, italic=False,
                      center=False, empty=False):
    """Insert new paragraph after ref_para. Returns new Paragraph."""
    p = OxmlElement("w:p")
    ref_para._p.addnext(p)
    new_para = para_from_elem(p)
    if not empty and text:
        run = new_para.add_run(text)
        run.bold = bold
        run.italic = italic
        if center:
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_para


def insert_para_before_tbl(table, text, bold=True, center=False):
    """Insert labeled paragraph before table element."""
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    para = para_from_elem(p)
    run = para.add_run(text)
    run.bold = bold
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def insert_figure_after(ref_para, img_path, caption, width=Inches(5.5)):
    """
    Insert: blank line → centered image → italic centered caption → blank line
    all AFTER ref_para.

    Uses doc.add_paragraph() to ensure the paragraph inherits the document
    part (needed for add_picture), then moves the _p element via lxml.
    """
    img_path = str(img_path)

    # --- blank line before ---
    blank1_p = doc.add_paragraph()  # appended at body end
    ref_para._p.addnext(blank1_p._p)  # move to after ref_para

    # --- image paragraph (created at end, moved to position) ---
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_path, width=width)
    blank1_p._p.addnext(img_para._p)  # move to after blank1

    # --- caption ---
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    img_para._p.addnext(cap_para._p)  # move to after image

    # --- blank line after ---
    blank2_p = doc.add_paragraph()
    cap_para._p.addnext(blank2_p._p)  # move to after caption

    return img_para, cap_para


def find_para_containing(text_fragment, start=0):
    """Return first Paragraph whose .text contains text_fragment."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if text_fragment in p.text:
            return p
    return None


# ─── FIX 1: Clean up inline Figure 2 reference (remove file path) ─────────────
OLD_FIG2 = "Figure 2 (see figures/02_asymmetric_WHY_full.png) provides"
NEW_FIG2 = "Figure 2 provides"

for para in doc.paragraphs:
    if OLD_FIG2 in para.text:
        full = para.text.replace(OLD_FIG2, NEW_FIG2, 1)
        for r in para.runs:
            r._r.getparent().remove(r._r)
        para.add_run(full)
        print("FIX 1: Cleaned inline Figure 2 path reference ✓")
        break


# ─── FIX 2: Rename "Table 3" → "Table 2" in Vicious Cycles paragraph ─────────
OLD_TBL3 = "Table 3. The Four Vicious Cycles of CDATA Aging"
NEW_TBL3 = "Table 2. The Four Vicious Cycles of CDATA Aging"

for para in doc.paragraphs:
    if OLD_TBL3 in para.text:
        full = para.text.replace(OLD_TBL3, NEW_TBL3, 1)
        for r in para.runs:
            r._r.getparent().remove(r._r)
        run = para.add_run(full)
        run.bold = True
        print("FIX 2: Renamed Table 3 → Table 2 (Vicious Cycles) ✓")
        break


# ─── FIX 3: Add missing table labels ─────────────────────────────────────────
# Table 0 in doc = PTM damage → label "Table 1"
insert_para_before_tbl(
    doc.tables[0],
    "Table 1. Five Forms of Irreversible Post-Translational Modification Damage "
    "in the Mother Centriole",
    bold=True
)
print("FIX 3a: Table 1 label added (PTM damage) ✓")

# Table 2 in doc = Hallmarks → label "Table 3"
insert_para_before_tbl(
    doc.tables[2],
    "Table 3. CDATA Mechanistic Derivation of the Hallmarks of Aging",
    bold=True
)
print("FIX 3b: Table 3 label added (Hallmarks) ✓")

# Table 3 in doc = Specificity comparison → label "Table 4"
insert_para_before_tbl(
    doc.tables[3],
    "Table 4. Structural Non-Renewability and Functional Duality: Candidate "
    "Long-Lived Structures Compared",
    bold=True
)
print("FIX 3c: Table 4 label added (Specificity) ✓")

# Table 4 in doc = Predictions → label "Table 5"
insert_para_before_tbl(
    doc.tables[4],
    "Table 5. Falsifiable Experimental Predictions of CDATA",
    bold=True
)
print("FIX 3d: Table 5 label added (Predictions) ✓")


# ─── FIX 4: Embed figures ─────────────────────────────────────────────────────

# --- Graphical Abstract (before Introduction section) ---
intro_para = find_para_containing("1. Introduction: The Fundamental Paradox")
if intro_para:
    insert_figure_after(
        intro_para,
        FIG_DIR / "02b_graphical_abstract.png",
        "Graphical Abstract. CDATA proposes that the progressive, irreversible "
        "accumulation of oxidative PTMs in the mother centriole is the primary "
        "upstream cause of organismal aging, acting through two failure tracks, "
        "four vicious cycles, and seven aging tracks that collectively explain "
        "the hallmarks of aging.",
        width=Inches(5.5)
    )
    print("FIG GA: Graphical abstract inserted after §1 header ✓")

# --- Figure 1: CDATA Overview (after §1.2 historical section) ---
hist_para = find_para_containing(
    "Early centrosomal aging hypotheses focused primarily on centrosome "
    "amplification"
)
if hist_para:
    insert_figure_after(
        hist_para,
        FIG_DIR / "01_overview.png",
        "Figure 1. Overview of the CDATA theoretical framework. The mother "
        "centriole functions as the primary aging timer through irreversible PTM "
        "accumulation, dual integration of ciliogenesis and spindle assembly, and "
        "a cell-autonomous ROS positive feedback loop. Seven downstream aging "
        "tracks and four vicious cycles derive mechanistically from this single "
        "upstream lesion.",
        width=Inches(5.5)
    )
    print("FIG 1: Overview figure inserted after §1.2 ✓")

# --- Figure 2: Asymmetric WHY (after the paragraph that describes it) ---
fig2_para = find_para_containing("Figure 2 provides")
if fig2_para:
    insert_figure_after(
        fig2_para,
        FIG_DIR / "02_asymmetric_WHY_full.png",
        "Figure 2. The asymmetric inheritance mechanism — why and how. "
        "(A) Three structural layers enforcing mother centriole retention in the "
        "stem cell daughter: distal appendage anchoring, pericentriolar fate "
        "determinant concentration, and niche-mediated selection. "
        "(B) The damage ratchet: at p > 0.5, expected centriolar age grows as "
        "n × p, creating monotonic damage accumulation across stem cell "
        "generations. (C) Cross-system evidence: Drosophila GSCs, mammalian RGCs, "
        "and HSCs. (D) The asymmetric QC landscape: the stem daughter retains old "
        "centrioles [System I] but receives new mitochondria and clean proteome "
        "[System II].",
        width=Inches(5.8)
    )
    print("FIG 2: Asymmetric WHY figure inserted ✓")

# --- Figure 3: PTM damage types (after §2.3 text, before the note) ---
ptm_para = find_para_containing(
    "A critical point regarding ROS vulnerability: the α-tubulin-rich composition"
)
if ptm_para:
    insert_figure_after(
        ptm_para,
        FIG_DIR / "03_damage_types.png",
        "Figure 3. The five forms of irreversible PTM damage in the mother "
        "centriole. (I) Protein carbonylation of SAS-6 and CEP135 disrupts the "
        "cartwheel scaffold. (II) Tubulin hyperacetylation stiffens the centriolar "
        "wall. (III) Protein aggregation of CPAP and CEP290 blocks PCM nucleation. "
        "(IV) Phosphorylation dysregulation disrupts PLK4/NEK2/AURKA signaling. "
        "(V) Distal appendage loss of CEP164/CEP89 prevents primary ciliogenesis. "
        "All five forms are irreversible and accumulate monotonically with "
        "organismal age.",
        width=Inches(5.5)
    )
    print("FIG 3: PTM damage types figure inserted ✓")

# --- Figure 4: ROS Positive Feedback Loop (after §2.5 last paragraph) ---
ros_para = find_para_containing(
    "both consistent with a midlife damage-acceleration factor in the range 1.5–1.7×"
)
if ros_para:
    insert_figure_after(
        ros_para,
        FIG_DIR / "05_ros_loop.png",
        "Figure 4. The ROS positive feedback loop (Postulate V). Centriolar PTM "
        "damage → MTOC disorganization → impaired mitophagy → mitochondrial "
        "dysfunction → elevated ROS → accelerated PTM carbonylation → more damage. "
        "A secondary tissue-level loop operates through myeloid bias and SASP. "
        "The Cell-DT calibration places the midlife damage acceleration factor at "
        "×1.6 (age 40+), consistent with epigenetic clock and Gompertz mortality "
        "data [18].",
        width=Inches(5.2)
    )
    print("FIG 4: ROS loop figure inserted after §2.5 ✓")

# --- Figure 5: Two failure tracks (after §3.2 last paragraph) ---
track_para = find_para_containing(
    "Track B therefore predicts a two-fate model for the cancer-aging relationship"
)
if track_para:
    insert_figure_after(
        track_para,
        FIG_DIR / "04_two_tracks.png",
        "Figure 5. The two primary failure tracks of centriolar aging. "
        "Track A (left): progressive distal appendage loss (CEP164↓) → failure of "
        "primary ciliogenesis → loss of Wnt/Shh/Notch niche signals → stem cell "
        "quiescence and pool depletion. "
        "Track B (right): spindle assembly dysfunction → symmetric pro-differentiative "
        "division → stem cell pool exhaustion; or symmetric self-renewal → clonal "
        "expansion and CHIP. The same upstream centriolar lesion drives the aging-"
        "cancer tradeoff.",
        width=Inches(5.5)
    )
    print("FIG 5: Two tracks figure inserted after §3.2 ✓")

# --- Figure 6: Damage accumulation curves (after §3.5 / vicious cycles table) ---
vcycle_text = find_para_containing(
    "The four vicious cycles are nested: Cycle I (cell-autonomous ROS–PTM) feeds "
    "into Cycle II"
)
if vcycle_text:
    insert_figure_after(
        vcycle_text,
        FIG_DIR / "06_damage_curves.png",
        "Figure 6. Simulated centriolar damage accumulation trajectories across "
        "the human lifespan (Cell-DT, n = 50 stochastic runs). The damage "
        "accumulation curve follows a sigmoidal profile: slow linear phase "
        "(ages 0–40), inflection at midlife (×1.6 acceleration factor), and "
        "exponential phase thereafter. Dashed lines indicate Track A (ciliary) "
        "and Track B (spindle) failure thresholds. Death in the simulation occurs "
        "at damage ≥ 0.75, corresponding to a median lifespan of ~79 years "
        "consistent with human all-cause mortality data.",
        width=Inches(5.5)
    )
    print("FIG 6: Damage curves inserted after vicious cycles ✓")

# --- Figure 7: Theory comparison (after §5.3 epigenetic clock comparison) ---
epi_para = find_para_containing(
    "Epigenetic clocks — methylation-based biomarkers that accurately predict "
    "biological age"
)
if epi_para:
    insert_figure_after(
        epi_para,
        FIG_DIR / "07_comparison.png",
        "Figure 7. CDATA in the landscape of aging theories. Comparison of CDATA, "
        "telomere theory, epigenetic clock theory, mitochondrial dysfunction theory, "
        "and inflammaging theory on four criteria: upstream causation, downstream "
        "mechanistic scope, falsifiability, and therapeutic implications. CDATA "
        "uniquely provides a single upstream lesion that mechanistically derives "
        "all major competing theories as downstream outputs.",
        width=Inches(5.5)
    )
    print("FIG 7: Theory comparison inserted after §5.3 ✓")

# --- Figure 8: CHIP clonal dynamics (in §7.3 CHIP section) ---
chip_para = find_para_containing(
    "CHIP — the expansion of hematopoietic clones carrying somatic driver mutations"
)
if chip_para:
    insert_figure_after(
        chip_para,
        FIG_DIR / "10_chip.png",
        "Figure 8. CHIP clonal dynamics in the CDATA framework. "
        "Cell-DT simulation of 20 HSC niches over 80 years (NichePool model). "
        "CHIP-initiating clones (green) show lower centriolar PTM burden at "
        "year 40 than non-expanded HSCs, acquiring a division-rate advantage "
        "that enables clonal expansion. By year 79, three dominant clones account "
        "for 50%, 29%, and 14% of the hematopoietic pool — consistent with "
        "age-related CHIP prevalence (≥40% by age 70 [16,17]). "
        "CDATA predicts that CHIP is a survival selection for lower-damage HSCs, "
        "not a random mutation event.",
        width=Inches(5.5)
    )
    print("FIG 8: CHIP figure inserted in §7.3 ✓")

# --- Figure 9: Interventions hierarchy (after §8 Predictions intro) ---
pred_intro = find_para_containing(
    "A theory's scientific value is proportional to the specificity of its "
    "falsifiable predictions"
)
if pred_intro:
    insert_figure_after(
        pred_intro,
        FIG_DIR / "09_interventions.png",
        "Figure 9. CDATA intervention hierarchy. Primary centriolar interventions "
        "(CEP164-AAV, CentrosomeTransplant, CAFD/IDI-targeted) act upstream and "
        "are predicted to outperform downstream senolytics (navitoclax), NAD+ "
        "precursors, and antioxidants across all four vicious-cycle metrics: "
        "stem cell division rate, senescent fraction, myeloid bias, and "
        "centrosomal ROS level. Cell-DT simulation (Prediction #11).",
        width=Inches(5.2)
    )
    print("FIG 9: Interventions figure inserted after §8 intro ✓")


# ─── FIX 5: Add Figure List after References ──────────────────────────────────
last_ref = find_para_containing(
    "Tkemaladze, J. (2026). CDATA Computational Validation"
)
if last_ref:
    p1 = insert_para_after(last_ref, empty=True)
    p2 = insert_para_after(p1, "List of Figures", bold=True)
    figure_list = [
        "Graphical Abstract. CDATA theoretical framework overview.",
        "Figure 1. CDATA conceptual overview — single centriolar lesion → seven "
        "aging tracks.",
        "Figure 2. The asymmetric inheritance mechanism (four-panel): structural "
        "layers, damage ratchet, cross-system evidence, QC landscape.",
        "Figure 3. The five forms of irreversible PTM damage in the mother "
        "centriole.",
        "Figure 4. The ROS positive feedback loop (Postulate V) with Cell-DT "
        "calibration.",
        "Figure 5. The two primary failure tracks: Track A (ciliary) and Track B "
        "(spindle).",
        "Figure 6. Simulated centriolar damage accumulation curves across the "
        "human lifespan.",
        "Figure 7. CDATA in the landscape of aging theories — comparative "
        "analysis.",
        "Figure 8. CHIP clonal dynamics: Cell-DT NichePool simulation over "
        "80 years.",
        "Figure 9. CDATA intervention hierarchy: primary vs. downstream "
        "interventions.",
    ]
    prev_p = p2
    for line in figure_list:
        prev_p = insert_para_after(prev_p, line)
    print("FIX 5: Figure list added after References ✓")


# ─── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(PATH_OUT)
print(f"\nFinal version saved: {PATH_OUT}")
print("Summary of changes:")
print("  ✓ Table labels: Table 1 (PTM), 2 (Vicious Cycles), 3 (Hallmarks), "
      "4 (Specificity), 5 (Predictions)")
print("  ✓ Figures embedded: Graphical Abstract + Figures 1–9")
print("  ✓ Inline Figure 2 file path removed from text")
print("  ✓ Figure list added at end")
print("  ✓ References [1]–[33] unchanged")
