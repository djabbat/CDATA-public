"""
Генерация полной статьи CDATA как docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Стили страницы ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Вспомогательные функции ──────────────────────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, italic=False, size=11, justify=True, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def mixed(parts, justify=True, size=11, space_after=6):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p

def bullet(text, level=0, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def table_row(t, cells_data, bold_first=False):
    row = t.add_row()
    for i, (cell_text, shade) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = cell_text
        if shade:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), shade)
            tcPr.append(shd)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)
            if bold_first and i == 0:
                run.bold = True
    return row

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2C3E50')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def hr():
    p = doc.add_paragraph('─' * 80)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

# ═══════════════════════════════════════════════════════════════════════════════
# ЗАГОЛОВОК
# ═══════════════════════════════════════════════════════════════════════════════

title = doc.add_heading('The Centriolar Damage Accumulation Theory of Aging (CDATA):', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(44, 62, 80)

subtitle = doc.add_paragraph('Molecular Foundations, Mechanistic Architecture,\nand Falsifiable Predictions of a Unified Theory')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 140, 141)
    run.italic = True

doc.add_paragraph()

auth = doc.add_paragraph('Jaba Tkemaladze')
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in auth.runs:
    run.bold = True
    run.font.size = Pt(12)

aff = doc.add_paragraph('Independent Research Laboratory for Integrative Geroscience\nCorrespondence: djabbat@gmail.com')
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in aff.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
hr()

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

heading('Abstract', 1)

para(
    "The Centriolar Damage Accumulation Theory of Aging (CDATA) proposes that the progressive, "
    "irreversible accumulation of oxidative post-translational modifications (PTMs) in the maternal "
    "centriole of adult stem cells constitutes a primary, causally upstream mechanism of organismal aging. "
    "The theory rests on five interconnected postulates: (i) the mother centriole is the only cellular "
    "structure for which no repair or replacement pathway exists in post-embryonic somatic cells; "
    "(ii) during asymmetric stem cell division, the older, more damaged mother centriole is selectively "
    "retained by the self-renewing daughter with a probability exceeding 50%, creating a ratchet effect "
    "that concentrates damage in the stem cell lineage over thousands of divisions; (iii) five distinct "
    "molecular damage forms accumulate in the centriole — protein carbonylation, tubulin hyperacetylation, "
    "protein aggregation, phosphorylation dysregulation, and distal appendage loss — each with defined "
    "molecular targets and functional consequences; (iv) the damaged mother centriole transduces its "
    "dysfunction into tissue aging through two mechanistically distinct failure tracks: Track A (loss of "
    "primary ciliogenesis and niche signaling) and Track B (loss of mitotic spindle fidelity and "
    "asymmetric division); and (v) a positive ROS feedback loop connecting centriolar damage to "
    "mitochondrial dysfunction creates the nonlinear acceleration of aging characteristic of the fifth "
    "and sixth decades of human life. CDATA is uniquely distinguished from competing theories by the "
    "convergence of absolute structural non-renewability, template-dependent replication ensuring "
    "perpetual retention, dual functional integration of both aging-critical processes, and exceptional "
    "ROS vulnerability — properties that no alternative candidate structure simultaneously possesses. "
    "We address six major objections to the theory — correlation-versus-causation, non-universality "
    "of asymmetric division, the specificity question, the cancer-aging paradox, partial UPS renewal, "
    "and the progeria challenge — and show that each is resolved within the CDATA framework without "
    "invoking additional assumptions. The theory generates specific, falsifiable experimental predictions "
    "absent from competing frameworks, including clonal hematopoiesis phenotypes, tissue failure "
    "sequences, and the therapeutic hierarchy of centriole-targeting versus downstream interventions. "
    "The final section outlines the motivation for formalizing CDATA as a computational digital twin."
)

kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = kw.add_run('Keywords: ')
r1.bold = True
r1.font.size = Pt(10)
r2 = kw.add_run(
    'aging, centriole, centrosome, stem cell, asymmetric division, primary cilium, '
    'post-translational modification, CDATA, cellular senescence, geroscience'
)
r2.font.size = Pt(10)
r2.italic = True

doc.add_paragraph()
hr()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

heading('1. Introduction: The Fundamental Paradox of Aging', 1)

heading('1.1 Why Tissues Age Despite Continuous Cellular Renewal', 2)

para(
    "A defining property of complex metazoans is the capacity for continuous tissue renewal. The human "
    "intestinal epithelium replaces itself every five days; erythrocytes circulate for approximately "
    "120 days before being cleared; hepatocytes turn over on a timescale of one to two years. Across "
    "all high-turnover tissues, pools of adult stem cells (SCs) maintain homeostasis by producing "
    "progenitor cells that differentiate into the mature tissue cells required for function. If the "
    "cellular machinery is continuously renewed, the fundamental paradox of aging presents itself with "
    "uncomfortable clarity: what, precisely, accumulates to cause progressive functional decline over "
    "decades?"
)

para(
    "The dominant frameworks of contemporary geroscience — telomere attrition, epigenetic drift, "
    "mitochondrial dysfunction, protein aggregation, chronic inflammation — each describe a genuine "
    "and important aspect of aging biology. However, they share a common limitation: they describe "
    "downstream phenotypes that co-occur with aging, without identifying a single molecular process "
    "capable of coherently driving all of them from a single upstream lesion. López-Otín and colleagues "
    "have catalogued twelve hallmarks of aging [1], yet the causal hierarchy connecting these hallmarks "
    "— which drives which — remains incompletely understood. A truly unified theory of aging must "
    "identify not merely correlates of aging, but the primary rate-limiting accumulator whose "
    "dysfunction propagates into the full hallmark landscape."
)

para(
    "The Centriolar Damage Accumulation Theory of Aging (CDATA) proposes such a primary mechanism. "
    "First formulated by Tkemaladze and Chichinadze in 2005 and developed through a series of "
    "publications [2–4], CDATA identifies the progressive, irreversible accumulation of oxidative "
    "PTMs in the maternal centriole of adult stem cells as the causally upstream event from which "
    "the full cascade of aging hallmarks can be derived. The logic of CDATA rests on a simple but "
    "powerful observation: among all cellular structures, the centriole alone combines absolute "
    "structural non-renewability, template-dependent replication ensuring perpetual retention of the "
    "old structure, dual functional integration of both ciliogenesis and spindle organization, and "
    "exceptional susceptibility to oxidative damage. No competing candidate possesses all four "
    "properties simultaneously."
)

heading('1.2 Historical Development of CDATA', 2)

para(
    "The intellectual antecedents of CDATA trace to John Cairns' Immortal Strand Hypothesis (1975), "
    "which proposed that stem cells asymmetrically segregate their DNA strands to minimize mutational "
    "accumulation in the self-renewing lineage [5]. CDATA extends this logic from the genome to a "
    "specific structural organelle. While the immortal strand hypothesis has remained contested and "
    "tissue-specific, the centriolar extension is grounded in the better-established biology of "
    "centrosome inheritance and in the unique biophysics of centriolar stability."
)

para(
    "Early centrosomal aging hypotheses focused primarily on centrosome amplification — the "
    "observation that aged cells accumulate supernumerary centrosomes, a phenomenon also associated "
    "with cancer [6]. CDATA departs from this direction by focusing not on numerical amplification "
    "but on the qualitative molecular deterioration of the mother centriole: the accumulation of "
    "PTM damage that impairs its functional capacity without necessarily increasing its number. "
    "The key publications establishing the modern CDATA framework include Tkemaladze (2023) [4], "
    "which provided quantitative evidence for centriolar age correlating with regenerative capacity "
    "decline in model organisms, and the 2025–2026 computational series that formalized the "
    "theory's mathematical structure and generated testable predictions through digital twin modeling."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. CORE POSTULATES
# ═══════════════════════════════════════════════════════════════════════════════

heading('2. The Five Core Postulates of CDATA', 1)

# ─── Postulate I ───────────────────────────────────────────────────────────────
heading('2.1 Postulate I: Absolute Structural Non-Renewability', 2)

para(
    "The first and most foundational postulate of CDATA is that the mother centriole is the only "
    "cellular structure in post-embryonic somatic animal cells for which no documented pathway of "
    "repair, disassembly, or de novo replacement exists under physiological conditions."
)

para(
    "Centriole duplication is a strictly template-dependent process: each new (daughter) centriole "
    "is assembled in the proximity of a pre-existing centriole, using it as an organizational scaffold. "
    "The master regulator of this process, PLK4 (Polo-like kinase 4), initiates the assembly of a "
    "single procentriole per existing centriole per cell cycle [7]. Crucially, this mechanism creates "
    "one new structure alongside the old one — it does not replace the old structure. The structural "
    "core of the mother centriole — the nine-fold symmetric triplet microtubule blades and the SAS-6 "
    "cartwheel scaffold — persists across all subsequent cell cycles without documented remodeling. "
    "In post-embryonic cells that have completed the de novo centriole biogenesis pathway active only "
    "in early embryogenesis, no mechanism has been identified that can disassemble and rebuild the "
    "structural core of a mature mother centriole [8]."
)

para(
    "This is in sharp contrast with other candidate long-lived structures. Nuclear lamin A undergoes "
    "active processing by ZMPSTE24 metalloprotease; mutations in this system produce the accelerated "
    "aging of Hutchinson-Gilford Progeria. Histone H3.3 is exchanged at transcriptionally active "
    "regions by histone chaperones HIRA and DAXX/ATRX. Nuclear pore complex components Nup96 and "
    "Nup88 exhibit slow but detectable turnover in post-mitotic cells, demonstrated by pulse-chase "
    "mass spectrometry [9]. Centrosomal proteins of the pericentriolar material (PCM) and regulatory "
    "kinases are actively turned over via the ubiquitin-proteasome system (UPS) at defined cell cycle "
    "phases [10]. By contrast, the structural core of the mother centriole — the very entity whose "
    "PTM state CDATA identifies as the primary aging determinant — shows no documented disassembly "
    "pathway, making it, in the most rigorous sense, a permanent accumulator of molecular history."
)

para(
    "An important qualification: CDATA does not claim that all centrosomal proteins are permanent. "
    "The dynamic regulatory scaffold of the PCM — including PCM1, pericentrin, and cell cycle kinases — "
    "turns over actively. CDATA specifically concerns the structural core, the distal appendage "
    "platform, and the cilia-forming machinery of the mature mother centriole. It is these components "
    "whose PTM-mediated deterioration CDATA identifies as the primary aging lesion."
)

# ─── Postulate II ──────────────────────────────────────────────────────────────
heading('2.2 Postulate II: Asymmetric Retention Creates a Damage Ratchet', 2)

para(
    "The second postulate establishes the mechanism by which damage accumulates preferentially in "
    "the stem cell lineage. During asymmetric stem cell division (ACD), the older, more damaged "
    "mother centriole is retained by the self-renewing daughter with a probability exceeding 50%, "
    "while the younger daughter centriole segregates to the differentiating (progenitor) daughter."
)

para(
    "The strongest experimental support for this postulate comes from Drosophila germline stem cells "
    "(GSCs), where Yamashita and colleagues demonstrated deterministic, 100% asymmetric inheritance: "
    "the centrosome anchored to the hub interface (invariably containing the older mother centriole) "
    "is always retained by the GSC daughter, while the non-anchored centrosome goes to the "
    "differentiating gonialblast [11]. Crucially, laser ablation of the centrosome-anchoring hub "
    "interface disrupts this orientation and randomizes division outcomes, confirming that the "
    "asymmetric inheritance is mechanistically enforced by centriolar appendage anchoring — not "
    "merely correlated with it."
)

para(
    "In mammalian neural stem cells (radial glial cells, RGCs), Wang et al. (2009) demonstrated "
    "that the mother centriole is anchored to the apical surface of the ventricular zone via its "
    "primary cilium, while the daughter centriole is basolaterally positioned [12]. This spatial "
    "asymmetry ensures preferential inheritance of the mother centriole by the apical self-renewing "
    "daughter. Disruption of this apical anchoring by depletion of distal appendage proteins "
    "correlates with premature depletion of the neural progenitor pool — a functional readout "
    "directly predicted by CDATA."
)

para(
    "CDATA does not require strictly deterministic asymmetry in all tissues. The theory requires "
    "only that the probability of the older centriole being retained by the self-renewing daughter "
    "exceeds 50%. Even a 60:40 bias, compounded across the approximately 10,000 stem cell divisions "
    "estimated to occur in the human hematopoietic system over a lifetime, is sufficient to create "
    "a profound ratchet effect: the expected centriolar age in the stem cell lineage increases "
    "monotonically, even if no individual division is deterministically asymmetric. Moreover, "
    "three additional mechanisms operate even in tissues with predominantly symmetric division: "
    "stochastic asymmetry at the single-cell level within a statistically symmetric population; "
    "niche-mediated selection of daughter cells with better ciliary function; and Track B "
    "(spindle fidelity loss), which operates independently of inheritance mode."
)

# ─── Postulate III ─────────────────────────────────────────────────────────────
heading('2.3 Postulate III: Five Forms of Irreversible PTM Damage', 2)

para(
    "The third postulate identifies the specific molecular forms through which damage accumulates "
    "in the mother centriole. Experimental data from aged cells and model organisms identify five "
    "categories of cumulative, irreversible molecular damage, each with a defined target and "
    "functional consequence."
)

add_table(
    headers=['#', 'Damage Form', 'Primary Targets', 'Molecular Consequence', 'Functional Outcome'],
    rows=[
        ['I', 'Protein carbonylation',
         'SAS-6, CEP135, CPAP',
         'Side-chain oxidation of Lys, Pro, Arg, Thr; irreversible protein cross-linking',
         'Cartwheel scaffold disruption; impaired procentriole template assembly'],
        ['II', 'Tubulin hyperacetylation',
         'α-tubulin Lys40 (HDAC6↓, SIRT2↓)',
         'Lysine acetylation stabilises centriolar microtubule wall; loss of dynamic instability',
         'Reduced structural flexibility; impaired PCM anchoring dynamics'],
        ['III', 'Protein aggregation',
         'CPAP, CEP290, SAS-6',
         'Age-dependent proteotoxic aggregates in lumen and pericentrosomal space',
         'Physical blockade of PCM γ-TuRC nucleation; disrupted appendage complex assembly'],
        ['IV', 'Phosphorylation dysregulation',
         'PLK4, NEK2, AURKA substrates; PP1/PP2A balance',
         'Loss of cell-cycle-regulated phosphorylation precision; constitutive or absent kinase activity',
         'Centriole duplication errors; spindle assembly checkpoint failures; mitotic delays'],
        ['V', 'Distal appendage loss',
         'CEP164, CEP89, Ninein, CEP170, SCLT1',
         'Progressive loss of distal/subdistal appendage proteins from attachment sites',
         'Failure of membrane docking; primary cilium cannot assemble; Wnt/Shh/Notch signaling loss'],
    ],
    col_widths=[0.5, 2.5, 3.0, 4.0, 4.0]
)

para(
    "These five damage forms are not independent: carbonylation of SAS-6 accelerates aggregation; "
    "tubulin hyperacetylation alters the binding affinity of appendage proteins; phosphorylation "
    "dysregulation permits uncontrolled PLK4 activity that drives aberrant centriole amplification. "
    "The damage forms converge on two functional outcomes — loss of ciliogenic capacity (Form V "
    "primarily) and loss of mitotic spindle fidelity (Forms I–IV primarily) — which map directly "
    "onto the two failure tracks described in Section 3."
)

para(
    "A critical point regarding ROS vulnerability: the α-tubulin-rich composition of the centriole "
    "makes it substantially more vulnerable to protein carbonylation than alternative candidate "
    "long-lived structures. Histones, largely sequestered within the nucleosome octamer, are "
    "relatively protected from cytoplasmic ROS. Nuclear lamins are lipid-bilayer-proximal but "
    "non-tubulin in composition. The open, microtubule-rich architecture of the centriole — "
    "exposed to the same ROS environment as the cytoplasm, without the protection afforded by "
    "membrane or nucleosome packaging — makes it uniquely susceptible to cumulative oxidative damage."
)

# ─── Postulate IV ──────────────────────────────────────────────────────────────
heading('2.4 Postulate IV: Dual Functional Integration of Two Aging-Critical Processes', 2)

para(
    "The fourth postulate identifies the property that uniquely elevates the centriole above all "
    "other candidate non-renewed structures as a primary aging driver: the mother centriole is the "
    "sole cellular structure that simultaneously controls both primary ciliogenesis and mitotic "
    "spindle pole organization. No other long-lived structure in the cell integrates both processes."
)

para(
    "Primary ciliogenesis is initiated when the mother centriole migrates to the plasma membrane "
    "and docks via its distal appendage platform (CEP164, CEP89, SCLT1, FBF1) to form the basal "
    "body. The axoneme of the primary cilium extends from this basal body and serves as the "
    "sensory organelle for Hedgehog, Wnt (non-canonical), PDGFR, and TGF-β signaling [13]. "
    "In stem cells, these pathways are essential for niche communication, quiescence regulation, "
    "and the cell-autonomous interpretation of positional information. Loss of primary ciliogenesis "
    "does not merely reduce one signaling input — it renders the stem cell functionally deaf to the "
    "chemical language of its niche."
)

para(
    "Simultaneously, the mother centriole serves as the older pole of the mitotic spindle. During "
    "mitosis, the centrosome nucleates astral microtubules that contact the cell cortex to define "
    "spindle orientation, and kinetochore microtubules that attach to chromosomes and segregate them "
    "with sub-micron precision. The mother centriole contributes more PCM, nucleates microtubules "
    "more robustly, and carries fate determinants (Numb, aPKC) asymmetrically distributed in the "
    "pericentriolar matrix that instruct daughter cell identity during ACD."
)

para(
    "Damage to the mother centriole therefore simultaneously impairs both ciliogenesis (via "
    "appendage loss) and spindle fidelity (via PCM disorganization, tubulin hyperacetylation, "
    "and scaffold disruption). This dual impact means that a single molecular lesion — oxidative "
    "PTM of the mother centriole — generates two independent aging output streams: one through "
    "niche signaling failure (Track A) and one through mitotic error accumulation (Track B). "
    "The amplification of damage through two independent mechanisms is a fundamental reason why "
    "the centriolar lesion is expected to have greater organismal impact than lesions in structures "
    "affecting only one of these functions."
)

# ─── Postulate V ───────────────────────────────────────────────────────────────
heading('2.5 Postulate V: The ROS Positive Feedback Loop', 2)

para(
    "The fifth postulate establishes the mechanism by which centriolar damage drives the nonlinear "
    "acceleration of aging characteristic of the sixth and seventh decades of human life. "
    "A positive feedback loop connects centriolar damage to mitochondrial dysfunction and back to "
    "elevated ROS, creating a self-amplifying cycle that transforms a slow linear accumulation "
    "into the exponential mortality curve described by the Gompertz law."
)

para(
    "The loop operates as follows. Damage to the mother centriole disrupts the architecture of "
    "the pericentriolar material and the microtubule organizing center (MTOC). The MTOC serves "
    "as the hub for the radial microtubule network that positions mitochondria, directs mitophagy "
    "flux, and coordinates the efficiency of oxidative phosphorylation. MTOC disorganization "
    "impairs mitochondrial positioning and reduces the efficiency of mitophagy (selective clearance "
    "of damaged mitochondria). Dysfunctional mitochondria accumulate, producing elevated ROS as "
    "a byproduct of impaired electron transport chain function. The elevated ROS directly "
    "accelerates protein carbonylation of centriolar structural proteins — thereby increasing "
    "the rate of new PTM damage accumulation. The loop is closed: centriolar damage → MTOC "
    "disorganization → mitochondrial dysfunction → elevated ROS → accelerated centriolar PTM "
    "accumulation → more centriolar damage."
)

para(
    "A secondary feedback amplifier operates at the tissue level: centriolar damage drives "
    "myeloid lineage bias in hematopoietic stem cells (via loss of Numb/aPKC fate determinant "
    "segregation), which produces a systemic pro-inflammatory state (inflammaging) characterized "
    "by elevated circulating IL-6, TNF-α, and IL-1β. The senescence-associated secretory "
    "phenotype (SASP) of senescent cells in the bone marrow niche further elevates systemic "
    "ROS, which accelerates centriolar damage in all tissues. This creates a multi-tissue "
    "positive feedback that becomes self-sustaining once a threshold of centriolar damage "
    "is crossed — providing a mechanistic explanation for the well-documented midlife "
    "acceleration of aging trajectories."
)

para(
    "Mathematically, this feedback is represented as a nonlinear term in the damage accumulation "
    "equation: the damage rate is proportional to (1 + k_feedback × current_damage), which "
    "produces the observed sigmoidal damage accumulation trajectory — slow initial accumulation "
    "followed by an accelerating phase after midlife — without requiring any externally imposed "
    "nonlinearity. The post-40 acceleration documented in human biomarker studies (epigenetic "
    "clock acceleration, telomere shortening rate increase, HSC myeloid shift) emerges naturally "
    "from this single feedback parameter."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. TWO FAILURE TRACKS
# ═══════════════════════════════════════════════════════════════════════════════

heading('3. Two Failure Tracks: Translating Molecular Damage into Tissue Aging', 1)

heading('3.1 Track A: Ciliary Failure and Niche Signaling Collapse', 2)

para(
    "Track A describes the pathway from centriolar PTM damage to the loss of primary ciliogenesis "
    "and the consequent failure of niche signaling. It operates as follows: the progressive loss "
    "of distal appendage proteins — CEP164, CEP89, Ninein, CEP170, and SCLT1 — from the mother "
    "centriole eliminates the docking platform required for centriole-to-membrane anchoring, the "
    "prerequisite step for axonemal assembly. Without functional distal appendages, the mother "
    "centriole cannot convert to a basal body, and the primary cilium fails to form."
)

para(
    "The functional consequences of this ciliary failure are profound and tissue-specific. In "
    "hematopoietic stem cells, primary cilia transduce Sonic Hedgehog (Shh) signals from "
    "osteoblastic niche cells, and their loss shifts the HSC division balance toward myeloid "
    "commitment [14]. This accounts for the well-documented age-related myeloid bias — the "
    "preferential production of myeloid (granulocytic/monocytic) over lymphoid progeny — that "
    "is a hallmark of immune aging. In neural stem cells, primary cilia are essential for "
    "Shh-driven amplification divisions that expand the neural progenitor pool; their loss "
    "leads to premature quiescence and the decline of adult neurogenesis documented in aged "
    "brains [15]. In muscle satellite cells, ciliary loss impairs Notch signaling from "
    "myofibers, reducing the ability of MuSCs to activate and re-enter the cell cycle in "
    "response to injury, contributing to sarcopenia."
)

para(
    "Track A thus generates a predictable, tissue-specific sequence of functional decline, "
    "ordered by the degree to which each stem cell niche depends on primary ciliary signaling "
    "for its maintenance. CDATA predicts the following tissue failure sequence: neural (highest "
    "cilia-dependence for neurogenesis maintenance) → hematopoietic (Shh-dependent HSC "
    "self-renewal) → muscle (Notch-dependent satellite cell activation) → intestinal epithelium "
    "(less cilia-dependent, relying primarily on Wnt secretion from niche cells). This specific "
    "sequence constitutes a falsifiable prediction of CDATA."
)

heading('3.2 Track B: Spindle Fidelity Loss and the Symmetric Division Catastrophe', 2)

para(
    "Track B describes the pathway from centriolar PTM damage to the loss of mitotic spindle "
    "fidelity and its consequences for stem cell pool dynamics. Unlike Track A, which operates "
    "through a specific molecular pathway (appendage loss → ciliary failure → signaling loss), "
    "Track B operates through the mechanical integrity of the mitotic apparatus and its "
    "implications for the probabilistic outcomes of stem cell division."
)

para(
    "A structurally compromised mother centriole — with carbonylated SAS-6, hyperacetylated "
    "tubulin, aggregated CPAP, and dysregulated PLK4/NEK2 activity — organizes a mitotic "
    "spindle with reduced fidelity. Specifically, PCM disorganization reduces the efficiency "
    "of microtubule nucleation and impairs the robust bipolarity that ensures accurate chromosome "
    "segregation. Hyperacetylated centriolar microtubules are less dynamic, impairing the "
    "correction of erroneous kinetochore-microtubule attachments. The loss of Ninein from "
    "subdistal appendages reduces the anchoring of astral microtubules to the cell cortex, "
    "disrupting spindle orientation relative to polarity cues."
)

para(
    "The consequences for stem cell division mode are critical. During asymmetric stem cell "
    "division, the correct orientation of the spindle relative to the niche-cell interface is "
    "required to ensure that fate determinants (Numb, aPKC, lethal giant larvae) segregate "
    "asymmetrically to the two daughters. A spindle that is improperly oriented or structurally "
    "compromised will produce symmetric divisions — both daughters receiving similar amounts of "
    "fate determinants. Two distinct symmetric outcomes are possible, each with serious "
    "consequences:"
)

bullet("Symmetric pro-differentiative division: both daughters commit to differentiation, "
       "depleting the stem cell pool by one unit per division. Progressive pool depletion "
       "reduces the regenerative capacity of the tissue, driving the phenotype of stem cell "
       "exhaustion — a canonical hallmark of aging.")

bullet("Symmetric self-renewal division: both daughters retain stem cell identity, "
       "producing a transient clonal expansion. While individually advantageous, stochastic "
       "accumulation of such events across thousands of HSC niches generates the clonal "
       "hematopoiesis (CHIP) phenotype documented in aging humans, where a small number of "
       "clones come to dominate the hematopoietic output, increasing risk of myeloid neoplasm "
       "and cardiovascular disease.")

para(
    "Track B therefore predicts a two-fate model for the cancer-aging relationship at the "
    "stem cell level: the same upstream lesion (spindle fidelity loss from centriolar PTM "
    "damage) can produce either pool exhaustion (aging phenotype) or clonal expansion "
    "(pre-malignant phenotype), depending on the probabilistic outcome of each symmetric "
    "division. This resolves the apparent paradox between CDATA's aging prediction and "
    "the observation that centrosome dysfunction is a cancer hallmark (see Section 5.4)."
)

heading('3.3 The Centriolar Inducer System: A Molecular Memory of Potency', 2)

para(
    "CDATA incorporates a specific molecular model for how the centriolar PTM state is "
    "translated into potency instructions. During embryogenesis, as totipotent zygotic "
    "progeny undergo progressive commitment, two distinct sets of irreversible differentiation "
    "inducers are hypothesized to form — molecules whose presence maintains the stem cell "
    "state and whose loss drives commitment to a specific differentiated fate."
)

para(
    "These inducer molecules are associated with the centriolar apparatus and are distributed "
    "during cell division according to the following logic. In System I (old centriole → "
    "differentiating daughter), the inducer detaches from the older centriole, which therefore "
    "concentrates in the differentiating cell; in System II (old centriole → stem daughter), "
    "the inducer detaches from the younger centriole, leaving the old centriole in the "
    "self-renewing cell. In both systems, the accumulation of PTM damage to the centriole "
    "directly erodes the inducer count by impairing the structural coupling that anchors "
    "inducer molecules to the centriolar appendages. The progressive depletion of inducer "
    "molecules from the mother centriole constitutes the molecular readout of aging in CDATA: "
    "as inducer count falls from its initial maximum (M₀ for the mother centriole), the "
    "potency of the stem cell declines along the hierarchy: totipotent → pluripotent → "
    "oligopotent → unipotent → apoptotic."
)

para(
    "This model makes the specific prediction that individual centriolar appendage proteins — "
    "measurable by single-molecule super-resolution microscopy (STED, PALM, STORM) — should "
    "show a progressive, age-dependent decline in copy number at individual stem cell "
    "centrosomes, with the rate of decline correlating with tissue-specific aging trajectories. "
    "STED microscopy data from Sieben et al. (2018) establishing that CEP164 is present in "
    "approximately 9 copies per distal appendage in human fibroblasts provides the first "
    "quantitative anchor for the inducer count model."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. CDATA AS A UNIFIED FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════════════

heading('4. CDATA as a Unified Framework: Deriving the Hallmarks of Aging', 1)

para(
    "A central theoretical claim of CDATA is that the seven aging tracks identified in the "
    "computational digital twin — cilia dysfunction, spindle fidelity loss, telomere shortening, "
    "epigenetic clock acceleration, mitochondrial ROS amplification, stem cell division rate "
    "decline, and hormonal life-history trade-off — are not independent phenomena but mechanistically "
    "coupled outputs of the single upstream centriolar lesion. Here we trace the derivation of "
    "each hallmark from centriolar damage."
)

add_table(
    headers=['Hallmark of Aging', 'CDATA Mechanistic Derivation', 'Track'],
    rows=[
        ['Stem cell exhaustion',
         'Track B: progressive symmetric pro-differentiative divisions deplete the SC pool; '
         'Track A: ciliary signal loss drives premature quiescence',
         'A + B'],
        ['Genomic instability',
         'Track B: spindle fidelity loss → chromosome mis-segregation, aneuploidy, DSB accumulation',
         'B'],
        ['Telomere attrition',
         'Spindle fidelity loss → aberrant mitotic exit → replication stress at telomeric regions; '
         'premature senescence from faulty ACD further accelerates telomere erosion',
         'B→C'],
        ['Epigenetic drift',
         'MTOC disorganization releases centrosomal pools of HDAC6 and SIRT2; aberrant deacetylase '
         'localization drives epigenetic clock acceleration; symmetric divisions expose both daughters '
         'to full epigenetic noise',
         'B→D'],
        ['Mitochondrial dysfunction',
         'MTOC disorganization impairs microtubule-based mitophagy flux; elevated ROS from dysfunctional '
         'mitochondria feeds back to accelerate PTM damage (Track E positive feedback loop)',
         'E'],
        ['Chronic inflammation (inflammaging)',
         'HSC myeloid bias (Track A ciliary loss) → elevated myeloid output → systemic inflammatory '
         'cytokines; SASP from senescent cells amplifies ROS across tissues',
         'A→inflam'],
        ['Loss of proteostasis',
         'Elevated ROS from Track E accelerates protein carbonylation and aggregation; '
         'centrosomal HDAC6 loss reduces aggresome formation capacity',
         'E→PTM'],
        ['Altered intercellular communication',
         'Primary ciliary loss eliminates paracrine signal reception; SASP disrupts niche '
         'architecture; myeloid shift rewires immune surveillance of aged tissues',
         'A'],
        ['Hormonal axis decline',
         'Track G: HPG-axis HPG-axis reproductive investment trade-off reduces somatic '
         'maintenance post-reproduction; estrogen loss at menopause removes mitochondrial '
         'ROS buffering, accelerating centrosomal damage',
         'G'],
    ],
    col_widths=[3.5, 8.5, 2.0]
)

para(
    "The key theoretical advantage of CDATA over competing frameworks is the directionality of "
    "this derivation: the hallmarks are outputs of the centriolar lesion, not its co-equal "
    "causes. This means that interventions targeting the primary centriolar lesion are predicted "
    "to have greater systemic effect than interventions targeting any individual downstream "
    "hallmark. A senolytic, for instance, clears senescent cells but does not prevent new "
    "senescent cells from being generated by continued centriolar damage. A telomere-extending "
    "therapy (TERT activation) extends replicative capacity but does not prevent the spindle "
    "fidelity defects that drive aneuploidy and genomic instability. Only interventions that "
    "restore or replace the functional centriolar inducer system — or prevent the PTM damage "
    "that depletes it — address the root cause."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CDATA VS COMPETING THEORIES
# ═══════════════════════════════════════════════════════════════════════════════

heading('5. CDATA in the Landscape of Aging Theories', 1)

heading('5.1 The Specificity Argument: Why the Centriole?', 2)

para(
    "The most fundamental challenge to CDATA is the specificity question: why should the "
    "centriole be the primary aging driver rather than any other long-lived, non-renewed "
    "structure? CDATA provides a four-property argument that applies uniquely to the centriole "
    "and to no other candidate."
)

add_table(
    headers=['Property', 'Mother Centriole', 'Nuclear Pores', 'Lamin A', 'H3.3 Histone'],
    rows=[
        ['Absolute structural non-renewability',
         '✓ No disassembly pathway documented', '✗ Nup96/88 turn over slowly', '✗ ZMPSTE24 processing', '✗ HIRA/DAXX exchange'],
        ['Template-dependent replication (old always retained)',
         '✓ Daughter built alongside, never replacing', '✗ Diluted at nuclear division', '✗ Synthesised de novo', '✗ Random segregation'],
        ['Dual function: cilia + spindle simultaneously',
         '✓ Unique dual role', '✗ No spindle/cilia role', '✗ No spindle/cilia role', '✗ No spindle/cilia role'],
        ['Primary ROS target (tubulin-rich, cytoplasm-exposed)',
         '✓ Maximal exposure and vulnerability', 'Moderate', 'Moderate', 'Protected (nucleosome)'],
    ],
    col_widths=[4.0, 3.5, 2.5, 2.5, 2.5]
)

para(
    "The convergence of all four properties in the mother centriole, and the absence of any "
    "competing candidate that simultaneously satisfies all four, constitutes the positive case "
    "for centriolar primacy. It does not claim that other long-lived structures play no role in "
    "aging — it claims that the centriole is the rate-limiting accumulator, the one whose "
    "deterioration propagates most broadly into the downstream hallmark landscape."
)

heading('5.2 Comparison with Telomere Theory', 2)

para(
    "The telomere theory of aging proposes that the progressive shortening of telomeric "
    "repeats at chromosomal ends limits the replicative capacity of somatic cells to the "
    "Hayflick limit (~50 doublings), driving senescence as telomeres reach critically short "
    "lengths. CDATA does not deny the role of telomere attrition; it subsumes it as a downstream "
    "consequence of Track B spindle fidelity loss. Spindle defects from centriolar damage "
    "generate replication stress at telomeric regions (which are intrinsically difficult "
    "replication templates) and drive aberrant mitotic exits that accelerate erosion. Furthermore, "
    "the telomere theory cannot explain non-replicative aspects of aging (post-mitotic neuronal "
    "decline, the myeloid shift, ciliopathy-like phenotypes) without additional assumptions. "
    "CDATA derives all of these from the single centriolar lesion."
)

heading('5.3 Comparison with Epigenetic Clock Theory', 2)

para(
    "Epigenetic clocks — methylation-based biomarkers that accurately predict biological age — "
    "represent some of the most reproducible findings in geroscience. The question is whether "
    "epigenetic drift is a cause or a consequence of aging. CDATA places it as a consequence: "
    "the release of centrosomal HDAC6 and SIRT2 (whose activity requires association with "
    "intact centriolar and ciliary structures) into the cytoplasm upon centriolar disorganization "
    "produces aberrant deacetylase activity affecting histones and transcriptional regulators. "
    "Symmetric stem cell divisions expose both daughter cells to the full noise of the "
    "epigenetic landscape rather than selectively inheriting a defined epigenetic state, "
    "amplifying epigenetic drift. CDATA predicts that epigenetic clock acceleration should "
    "be detectable downstream of centriolar damage markers."
)

heading('5.4 The Cancer-Aging Paradox', 2)

para(
    "A canonical objection to CDATA is the apparent paradox with cancer biology: centrosome "
    "amplification is a hallmark of cancer cells, and if CDATA posits that centriolar "
    "dysfunction drives aging (stem cell exhaustion), why do cells with supernumerary "
    "centrosomes show proliferative rather than senescent behavior?"
)

para(
    "This paradox dissolves under the Track B two-fate model. Cancer-associated centrosome "
    "amplification is fundamentally distinct from mother centriole aging damage: it involves "
    "the de novo creation of extra centrosomes via overduplication (PLK4 overexpression, "
    "S-phase overriding, cytokinesis failure) — not the retention of a progressively damaged "
    "single mother centriole. Extra centrosomes cause multipolar spindles and aneuploidy (CIN), "
    "which is the pro-tumorigenic failure mode of Track B. The aging failure mode of Track B "
    "is the opposite: a damaged single-pole centriole that fails to organize a sufficiently "
    "robust bipolar spindle, driving symmetric divisions that exhaust the pool."
)

para(
    "CDATA and cancer biology are thus unified within a single framework: aging represents the "
    "pool-exhaustion failure mode of the Track B dysfunction spectrum; cancer represents the "
    "clonal-expansion failure mode. Both emerge from the same underlying cause — progressive "
    "failure of spindle-organizing fidelity — but with opposite stochastic outcomes. "
    "This unified view predicts that tissues with higher CDATA-driven aging rates should also "
    "show higher rates of age-associated cancer — a prediction supported epidemiologically."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. ADDRESSING OBJECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading('6. Addressing Major Objections to CDATA', 1)

heading('6.1 The Correlation-Versus-Causation Problem', 2)

para(
    "The most important methodological objection is that the existing experimental evidence "
    "for CDATA is largely correlational. Aged stem cells show more centriolar damage, and "
    "aged organisms show more stem cell dysfunction — but this does not establish that "
    "centriolar damage causes stem cell dysfunction rather than being a parallel consequence "
    "of cellular aging."
)

para(
    "CDATA addresses this objection through three lines of evidence. First, genetic experiments "
    "in which centriolar damage is induced in young cells — through laser ablation, overexpression "
    "of dominant-negative centriolar proteins, or targeted depletion of appendage proteins — "
    "recapitulate aged stem cell phenotypes (impaired ACD, premature quiescence, reduced "
    "regenerative capacity) in otherwise young organisms. This demonstrates that centriolar "
    "damage is sufficient to produce aging-like phenotypes in the absence of any other "
    "age-related change."
)

para(
    "Second, the converse experiment — transplantation of young centrosomal material into "
    "aged stem cells, or pharmacological restoration of appendage protein complexes — should "
    "produce rejuvenation of stem cell function. This prediction is experimentally testable "
    "and is the most direct test of the causal claim."
)

para(
    "Third, the computational digital twin (described in the companion paper) provides a "
    "qualitatively different class of evidence. A mechanistic model that incorporates only "
    "the known molecular rates of centriolar damage accumulation, implements only the "
    "established functional consequences for ciliary signaling and mitotic spindle fidelity, "
    "and reproduces realistic human lifespan trajectories — without requiring any additional "
    "'aging module' — constitutes, at minimum, a proof of causal sufficiency: the centriolar "
    "mechanism, if it operates at the rates measured, is capable of driving the observed "
    "aging phenotype without additional input."
)

heading('6.2 Non-Universality of Asymmetric Centriole Inheritance', 2)

para(
    "A serious technical objection is that strict asymmetric centriole segregation is "
    "well-established in Drosophila GSCs and partially confirmed in neural progenitors, "
    "but Lgr5+ intestinal stem cells divide predominantly symmetrically (Ritsma et al., 2014), "
    "and hematopoietic stem cell data are contradictory. If ACD is not universal, the CDATA "
    "ratchet mechanism cannot be universal."
)

para(
    "As established in Postulate II, CDATA does not require strict deterministic asymmetry. "
    "The theory requires only a probabilistic bias exceeding 50%, and even tissues with "
    "predominantly symmetric division are subject to the ratchet through three additional "
    "mechanisms: (i) stochastic asymmetry, where individual cells within statistically "
    "symmetric compartments show centriole-age-dependent fate biases; (ii) niche-mediated "
    "selection, where stem cells with less damaged centrioles and therefore better ciliary "
    "function are preferentially retained in the niche, effectively producing functional "
    "asymmetric inheritance at the population level; (iii) Track B independence, which operates "
    "regardless of inheritance mode since spindle fidelity loss drives symmetric outcomes "
    "even in cells that inherit perfectly equal centrioles."
)

heading('6.3 Partial Renewal via the Ubiquitin-Proteasome System', 2)

para(
    "The objection is raised that centrosomal proteins undergo partial UPS-mediated degradation "
    "and replacement, potentially compensating for damage and undermining the irreversibility "
    "postulate."
)

para(
    "CDATA distinguishes the structural core of the centriole (nine triplet microtubule blades, "
    "cartwheel scaffold) from the dynamic regulatory periphery (PCM components, cell-cycle "
    "kinases). UPS-mediated turnover operates primarily on the dynamic periphery — not on the "
    "structural core. Furthermore, and critically, the UPS itself declines in activity with "
    "age. Even if partial renewal occurs in young organisms, this compensatory mechanism fails "
    "precisely as centriolar damage accumulates — the same timeframe when CDATA predicts "
    "the transition from compensated to decompensated centriolar aging (approximately the "
    "fourth decade of human life). CDATA thus predicts not merely centriolar damage, but "
    "the age-dependent failure of the systems that would otherwise compensate for it — a "
    "prediction consistent with the documented age-related decline of UPS activity."
)

heading('6.4 Progeria: Primary or Secondary Centriolar Involvement?', 2)

para(
    "The objection is that in Hutchinson-Gilford Progeria Syndrome (HGPS), centrosomal defects "
    "are secondary to lamin A mutation, suggesting that centrioles are damaged as a consequence "
    "of aging rather than as its cause."
)

para(
    "CDATA does not require centriolar damage to be the only initiating event across all "
    "pathological contexts. HGPS represents a case of accelerated centriolar damage through an "
    "upstream mechanism — nuclear membrane instability generates DNA damage, which generates "
    "mitotic errors, which generate centrosomal stress. In HGPS, the centriolar damage is "
    "secondary to the lamin A mutation, but the downstream aging phenotypes are still mediated "
    "through the centriolar dysfunction. CDATA proposes that in normal physiological aging, "
    "centriolar PTM damage is the primary rate-limiting accumulator — not that it is the "
    "only possible initiating event in all disease contexts. HGPS provides a model of "
    "accelerated aging in which the downstream centriolar lesion is reached earlier, validating "
    "the functional importance of centriolar integrity for the aging phenotype without "
    "contradicting the causal claim for normal aging."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. EVIDENCE FROM MODEL ORGANISMS
# ═══════════════════════════════════════════════════════════════════════════════

heading('7. Evidence from Model Organisms and Human Studies', 1)

heading('7.1 Drosophila melanogaster', 2)

para(
    "Drosophila provides the most direct experimental evidence for CDATA's core mechanism. "
    "In the Drosophila testis, Yamashita et al. (2007) demonstrated that centrosome age "
    "predicts stem cell fate: the centrosome with the older mother centriole is consistently "
    "retained in the GSC. Genetic disruption of centrosome asymmetry by depletion of the "
    "hub-interface anchoring complex randomizes ACD outcomes and accelerates GSC loss — "
    "a direct functional demonstration that asymmetric centrosome inheritance is required "
    "for GSC maintenance. In the Drosophila intestinal epithelium, loss of centrosomes "
    "in ISCs disrupts tissue homeostasis and accelerates aging phenotypes, including increased "
    "apoptosis, reduced proliferation, and dysplastic tissue architecture. The centrosome-free "
    "ISC experiments are particularly important: they demonstrate that it is not centrosome "
    "presence per se but the functional state of the centriolar apparatus that matters for "
    "stem cell homeostasis."
)

heading('7.2 Neural Stem Cells in Mammals', 2)

para(
    "The developing and adult mammalian brain provides strong evidence for Track A of CDATA. "
    "In the developing cortex, radial glial cells (RGCs) use their mother-centriole-anchored "
    "primary cilium as the apical signaling hub. Mutations in genes encoding distal appendage "
    "proteins — CEP164, CEP83, CEP120 — produce severe microcephaly in mouse models, "
    "consistent with premature depletion of the neural progenitor pool via loss of ciliary "
    "Shh/Wnt signaling. In the aging brain, adult NSCs in the subventricular zone show "
    "progressive loss of ciliation and reduced Shh responsiveness, correlating with the "
    "well-documented decline of adult neurogenesis with age. The timeline of ciliary loss "
    "precedes the measurable decline in neurogenic output, consistent with CDATA's prediction "
    "that ciliary dysfunction is the proximate cause of neurogenic decline."
)

heading('7.3 Hematopoietic System and CHIP', 2)

para(
    "The aging hematopoietic system provides evidence for both tracks of CDATA. Track A predicts "
    "that HSC ciliary loss drives myeloid bias; this is supported by the observation that "
    "pharmacological restoration of Shh ciliary signaling in aged HSCs partially rescues "
    "the lymphoid output deficit. Track B predicts that spindle fidelity loss in aged HSCs "
    "drives the clonal hematopoiesis (CHIP) phenotype via asymmetric self-renewal events "
    "in cells with compromised spindle apparatus."
)

para(
    "CHIP — the expansion of hematopoietic clones carrying somatic driver mutations (DNMT3A, "
    "TET2, ASXL1) — is detected in approximately 10–15% of individuals over age 70 and "
    "confers a 10-fold increased risk of hematological malignancy and a 40% increased risk "
    "of cardiovascular mortality. CDATA generates a specific prediction about CHIP: clones "
    "with CHIP should preferentially harbor HSCs with lower centriolar PTM burden than "
    "surrounding non-expanded HSCs — because cells with less centriolar damage have a "
    "competitive fitness advantage arising from better spindle fidelity and more faithful "
    "ACD. This prediction is distinct from and complementary to the driver-mutation model "
    "of CHIP and is directly testable by super-resolution microscopy of γ-tubulin PTMs "
    "in sorted CHIP-clone versus non-clone HSC fractions."
)

heading('7.4 Human Studies and Clinical Correlates', 2)

para(
    "In human aging, centrosomal abnormalities — enlargement, fragmentation, PCM "
    "disorganization — are consistently documented in biopsies from aged tissues across "
    "multiple organ systems, including skin, bone marrow, intestinal epithelium, and "
    "skeletal muscle. The age-related decline in primary ciliation — documented by "
    "high-content imaging of kidney tubular epithelium, olfactory mucosa, and airway "
    "epithelium — correlates with the functional decline of the organs these tissues comprise."
)

para(
    "Ciliopathies — congenital disorders caused by mutations in cilia-related genes — "
    "provide a natural experiment supporting Track A. Conditions including polycystic kidney "
    "disease (PKD), Bardet-Biedl syndrome, and Joubert syndrome — all caused by defects "
    "in ciliogenesis or ciliary function — produce phenotypes that closely mirror individual "
    "aspects of normal aging: cystic organ degeneration, obesity, metabolic dysfunction, "
    "and neurological decline. That these aging-like phenotypes can be produced entirely "
    "by ciliary dysfunction — without any other aging insult — is consistent with CDATA's "
    "claim that progressive ciliary loss from centriolar damage is sufficient to drive "
    "tissue aging."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. FALSIFIABLE PREDICTIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading('8. Falsifiable Experimental Predictions', 1)

para(
    "A theory's scientific value is proportional to the specificity of its falsifiable predictions. "
    "CDATA generates the following testable predictions, each of which would, if disconfirmed, "
    "challenge the theory's causal claims:"
)

add_table(
    headers=['Prediction', 'Experimental Test', 'Disconfirming Result'],
    rows=[
        ['CHIP clones show lower centriolar PTM burden than non-expanded HSCs',
         'Sort CHIP clones (DNMT3A/TET2 mutant) vs. non-clone HSCs; STED microscopy of γ-tubulin carbonylation and CEP164 copy number',
         'No difference in PTM burden; CHIP clones show equal or higher centriolar damage'],
        ['Neural tissue fails before hematopoietic in aging simulations',
         'Longitudinal imaging of adult neurogenesis vs. HSC output decline in murine aging (2–24 months)',
         'Hematopoietic decline precedes neurogenic decline in the same organism'],
        ['Young centrosome transplant into aged SC rejuvenates division fidelity',
         'Microinjection of isolated young centrosomes into aged muscle satellite cells or HSCs; measure ACD efficiency and fate outcomes',
         'No rescue of ACD efficiency; fate outcomes unchanged after transplant'],
        ['Progeria HSCs show accelerated centriolar PTM vs. age-matched controls',
         'CEP164/CEP89 immunostaining and carbonylation assay in HGPS patient-derived HSC-like cells vs. age-matched controls',
         'No acceleration; PTM levels similar to age-matched controls'],
        ['Primary centriolar interventions outperform downstream interventions',
         'In a model organism (Drosophila ISC) treat with: (A) CEP164-restoring peptide vs. (B) senolytic; measure SC number and regenerative capacity at 40 days',
         'Senolytics provide equal or greater healthspan benefit than centriole-targeting therapy'],
        ['Age of menarche inversely correlates with rate of centrosomal PTM accumulation',
         'Cross-cohort study: MSC centrosomal damage markers at age 50 vs. age at menarche in females',
         'No correlation between menarche age and centrosomal damage rate'],
    ],
    col_widths=[4.5, 5.5, 4.0]
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. THE CENTRIOLE AS INTRACELLULAR TIMER
# ═══════════════════════════════════════════════════════════════════════════════

heading('9. The Centriole as an Intracellular Timer: Encoding Cellular History', 1)

para(
    "Beyond its role as a passive accumulator of damage, the mother centriole functions as an "
    "active intracellular timer — an information-processing hub that encodes the divisional "
    "history of a cell and translates this history into fate-determining signals. This framing "
    "elevates CDATA from a damage-accumulation narrative to a molecular computation model."
)

para(
    "Three mechanistic layers constitute the centriolar timing system. First, the PTM-code layer: "
    "the mother centriole accumulates a specific, cumulative signature of post-translational "
    "modifications — polyglutamylation, acetylation, carbonylation — that changes predictably "
    "with the number of completed cell cycles. This 'epicentriolar' signature recruits specific "
    "downstream effector proteins, including Notch signaling components and Par polarity complex "
    "elements, to the centrosome in an age-dependent manner. The PTM-code functions analogously "
    "to the histone modification code, but for the temporal dimension of cell fate rather than "
    "its spatial (chromatin) dimension."
)

para(
    "Second, the PCM layer: the pericentriolar material surrounding the mother centriole serves "
    "as a dynamic buffer for cell cycle regulators, including Cyclin B1, Aurora A, and PLK1. "
    "The composition and size of the PCM expand during mitotic entry (centrosome maturation) "
    "and are regulated by the cumulative kinase activity of the centriolar core. Progressive "
    "damage to the centriolar core alters PCM composition in a cell-cycle-history-dependent "
    "manner, linking divisional history to mitotic regulators."
)

para(
    "Third, the nuclear coupling layer: the centrosome is physically connected to the nucleus "
    "via the LINC (Linker of Nucleoskeleton and Cytoskeleton) complex. Mechanical forces "
    "transmitted through this connection, which vary with centrosomal maturity and positioning, "
    "influence nuclear envelope tension and chromatin organization. Centriolar aging alters "
    "the mechanical coupling between the centrosome and nucleus, potentially modulating "
    "transcriptional programs that are sensitive to nuclear architecture — including the "
    "senescence transcriptional program regulated by the cGAS-STING pathway, which detects "
    "cytoplasmic chromatin fragments released during aberrant mitosis."
)

para(
    "The triad of PTM-code, PCM buffer, and nuclear coupling constitutes a multi-layered "
    "timekeeping mechanism that is not merely passive but actively influences cell fate "
    "decisions. In this view, organismal aging is, in part, the progressive misreading of "
    "cellular history — as the centriolar timer accrues damage, its ability to accurately "
    "encode divisional history and translate it into appropriate fate decisions degrades, "
    "producing the progressive loss of stem cell fidelity that manifests as tissue aging."
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. TRANSITION TO DIGITAL TWIN
# ═══════════════════════════════════════════════════════════════════════════════

heading('10. From Theory to Formalization: The Motivation for a Computational Digital Twin', 1)

para(
    "The theoretical architecture of CDATA, as described in Sections 2–9, constitutes a "
    "qualitatively compelling unified framework for aging. However, compelling narrative "
    "integration of biological evidence is a necessary but insufficient criterion for "
    "scientific validity. Three fundamental problems require computational formalization "
    "to address rigorously."
)

para(
    "First, the integration problem: CDATA proposes that seven aging tracks are mechanistically "
    "coupled through the centriolar lesion. Whether the quantitative rates of each process — "
    "as measured in independent experimental systems — are mutually consistent and jointly "
    "sufficient to generate a 78-year human lifespan trajectory cannot be determined by "
    "intuitive reasoning. Only a mathematical model that implements all tracks simultaneously, "
    "with their interactions, can determine whether the proposed causal architecture is "
    "quantitatively coherent."
)

para(
    "Second, the prediction problem: CDATA's most scientifically valuable contribution is its "
    "ability to generate falsifiable quantitative predictions — not merely qualitative directional "
    "claims, but specific numerical thresholds, timelines, and effect magnitudes. What age "
    "does the model predict for the onset of measurable myeloid bias? What healthspan gain "
    "does a specific degree of centrosomal ROS reduction produce? How much does the puberty "
    "age correlation contribute to lifespan variance? These questions are accessible only "
    "to a computational model with explicit parameter values."
)

para(
    "Third, the intervention hierarchy problem: CDATA predicts that centriole-targeting "
    "interventions should outperform downstream interventions. But by how much? And which "
    "specific centriole-targeting strategy — appendage protein restoration, centrosome "
    "transplantation, ROS reduction at the centrosomal compartment, CAFD manipulation — "
    "provides the greatest leverage? Without a quantitative model of the full feedback "
    "architecture, these questions cannot be answered and CDATA's therapeutic implications "
    "cannot be prioritized."
)

para(
    "These three problems motivate the development of the computational digital twin (Cell-DT) "
    "described in the companion paper. Cell-DT is implemented as a 14-crate Rust workspace "
    "using the Entity Component System (ECS) paradigm, which treats each stem cell niche as "
    "an independent entity and each biological subsystem as a modular component. The platform "
    "integrates all five postulates of CDATA, both failure tracks, the positive ROS feedback "
    "loop, and four additional aging tracks (telomere shortening, epigenetic clock, "
    "mitochondrial dynamics, and hormonal life-history trade-off) into a single coherent "
    "simulation that can be run, perturbed, and interrogated in silico."
)

para(
    "The companion paper addresses the mathematical architecture of Cell-DT and its calibration "
    "results in detail. Here we note only that the digital twin's role is explicitly not to "
    "validate CDATA by reproducing calibration targets — this would be tautological. Its role "
    "is to (i) test whether the proposed causal architecture is internally consistent at "
    "quantitative resolution; (ii) identify which parameters the theory's predictions are "
    "most sensitive to, guiding experimental prioritization; (iii) generate the specific "
    "numerical predictions — CHIP onset age, myeloid bias trajectory, tissue failure "
    "sequence, intervention efficacy hierarchy — that can be tested against independent "
    "empirical data not used in calibration; and (iv) provide a platform for in silico "
    "screening of therapeutic strategies before committing to expensive animal or clinical experiments."
)

doc.add_paragraph()
hr()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

heading('11. Conclusion', 1)

para(
    "The Centriolar Damage Accumulation Theory of Aging proposes a resolution to the fundamental "
    "paradox of organismal aging: the structure that accumulates irreversible damage in the stem "
    "cell lineage is not the genome, not the mitochondria, not the epigenome — but the mother "
    "centriole, the one organelle that combines absolute non-renewability, perpetual retention "
    "through template-dependent replication, dual functional integration of ciliogenesis and "
    "spindle organization, and exceptional vulnerability to the oxidative environment."
)

para(
    "Aging, in the CDATA framework, is the price that multicellular organisms with true tissues "
    "pay for possessing irreversibly differentiated, highly specialized cell types. The centriole "
    "cannot be repaired because its damage state is functionally required for the precision of "
    "differentiation — its accumulated PTM signature is the molecular memory that drives "
    "progressive commitment. The very mechanism that enables the biological complexity of a "
    "vertebrate organism — precise asymmetric division instructed by centriolar age — inevitably "
    "produces the aging of the stem cell lineage that serves these tissues. Aging is not an "
    "evolutionary mistake; it is the necessary cost of developmental precision."
)

para(
    "This framing suggests a therapeutic philosophy: aging interventions that work with this "
    "constraint — restoring the functional capacity of the centriolar inducer system, reducing "
    "the rate of centrosomal ROS damage, or periodically replacing depleted centrosomal material "
    "with young functional equivalents — are predicted to have substantially greater effect than "
    "interventions that address the downstream symptoms of a problem whose upstream cause "
    "continues to operate. The computational digital twin described in the companion paper "
    "provides the quantitative framework for evaluating these predictions and prioritizing "
    "experimental investment."
)

doc.add_paragraph()
hr()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

heading('References', 1)

refs = [
    "López-Otín C, Blasco MA, Partridge L, Serrano M, Kroemer G. Hallmarks of aging: an expanding universe. Cell 2023; 186: 243–278.",
    "Tkemaladze JV, Chichinadze KN. Centriole, differentiation, and senescence. Rejuvenation Research 2010; 13: 339–342.",
    "Chichinadze K, Tkemaladze D. Centrosomal hypothesis of cellular aging and differentiation. Advances in Gerontology 2008; 21: 367–371.",
    "Tkemaladze J. Reduction, proliferation, and differentiation defects of stem cells over time: a consequence of selective accumulation of old centrioles in the stem cells? Molecular Biology Reports 2023; 50: 1951–1958. PMID: 36583780.",
    "Cairns J. Mutation selection and the natural history of cancer. Nature 1975; 255: 197–200.",
    "Ganem NJ, Godinho SA, Pellman D. A mechanism linking extra centrosomes to chromosomal instability. Nature 2009; 460: 278–282.",
    "Nigg EA, Holland AJ. Once and only once: mechanisms of centriole duplication and their deregulation in disease. Nature Reviews Molecular Cell Biology 2018; 19: 297–312.",
    "Bettencourt-Dias M, Glover DM. Centrosome biogenesis and function: centrosomics brings new understanding. Nature Reviews Molecular Cell Biology 2007; 8: 451–463.",
    "D'Angelo MA, Raices M, Bhargava P, Bhattacharya J, Bhargava S, Hetzer MW. Age-dependent deterioration of nuclear pore complexes causes a loss of nuclear integrity in postmitotic cells. Cell 2009; 136: 284–295.",
    "Prosser SL, Pelletier L. Mitotic spindle assembly in animal cells: a fine balancing act. Nature Reviews Molecular Cell Biology 2017; 18: 187–201.",
    "Yamashita YM, Mahowald AP, Perlin JR, Fuller MT. Asymmetric inheritance of mother versus daughter centrosome in stem cell division. Science 2007; 315: 518–521.",
    "Wang X, Tsai JW, Imai JH, et al. Asymmetric centrosome inheritance maintains neural progenitor pool during brain development. Nature 2009; 461: 947–955.",
    "Nachury MV, Mick DU. Establishing and regulating the composition of cilia for signal transduction. Nature Reviews Molecular Cell Biology 2019; 20: 389–405.",
    "Zhao M, Perry JM, Marshall H, et al. Megakaryocytes maintain homeostatic quiescence and promote post-injury regeneration of hematopoietic stem cells. Nature Medicine 2014; 20: 1321–1326.",
    "Doetsch F, García-Verdugo JM, Alvarez-Buylla A. Cellular composition and three-dimensional organization of the subventricular germinal zone in the adult mammalian brain. Journal of Neuroscience 1997; 17: 5046–5061.",
    "Jaiswal S, Fontanillas P, Flannick J, et al. Age-related clonal hematopoiesis associated with adverse outcomes. New England Journal of Medicine 2014; 371: 2488–2498.",
    "Genovese G, Kähler AK, Handsaker RE, et al. Clonal hematopoiesis and blood-cancer risk inferred from blood DNA sequence. New England Journal of Medicine 2014; 371: 2477–2487.",
    "Horvath S, Raj K. DNA methylation-based biomarkers and the epigenetic clock theory of ageing. Nature Reviews Genetics 2018; 19: 371–384.",
    "Sieben C, Banterle N, Douglass KM, Gönczy P, Manley S. Multicolor single-particle reconstruction of protein complexes. Nature Methods 2018; 15: 777–784.",
    "Steensma DP, Bejar R, Jaiswal S, et al. Clonal hematopoiesis of indeterminate potential and its distinction from myelodysplastic syndromes. Blood 2015; 126: 9–16.",
    "Kujoth GC, Hiona A, Pugh TD, et al. Mitochondrial DNA mutations, oxidative stress, and apoptosis in mammalian aging. Science 2005; 309: 481–484.",
    "Rettberg JR, Yao J, Brinton RD. Estrogen: a master regulator of bioenergetic systems in the brain and body. Frontiers in Neuroendocrinology 2014; 35: 8–30.",
    "Campisi J, d'Adda di Fagagna F. Cellular senescence: when bad things happen to good cells. Nature Reviews Molecular Cell Biology 2007; 8: 729–740.",
    "Lopes CS, Zheng H, Viales RR, et al. PLK4 trans-autoactivation controls centriole biogenesis in space and time. Developmental Cell 2015; 35: 222–235.",
    "Vertii A, Hehnly H, Doxsey S. The centrosome, a multitalented renaissance organelle. Cold Spring Harbor Perspectives in Biology 2016; 8: a025049.",
    "Habib SJ, Chen BC, Tsai FC, et al. A localized Wnt signal orients asymmetric stem cell division in vitro. Science 2013; 339: 1445–1448.",
    "Ritsma L, Ellenbroek SIJ, Zomer A, et al. Intestinal crypt homeostasis revealed at single-stem-cell level by in vivo live imaging. Nature 2014; 507: 362–365.",
    "Tkemaladze J. Old Centrioles Make Old Bodies. Annals of Rejuvenation Science 2025; 1(1). doi: 10.65649/yx9sn772.",
    "Tkemaladze J. Centrioles as intracellular timers of the cell cycle and cell fate. Longevity Horizon 2026; 2(3). doi: 10.65649/mxhcj531.",
    "Tkemaladze J. The Centriolar Theory of Differentiation Explains the Biological Meaning of the Centriolar Theory of Organismal Aging. Longevity Horizon 2025; 1(3). doi: 10.5281/zenodo.15057288.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{i}]  {ref}")
    r.font.size = Pt(9)

doc.add_paragraph()
hr()

# Сохранить
path = '/home/oem/Desktop/CDATA/CDATA_Theory_Full_Article.docx'
doc.save(path)
print(f'\n✅  Сохранено: {path}')
