#!/usr/bin/env python3
"""
Convert EIC_Pathfinder_CDATA_PartB.md to a professionally formatted .docx
following the FORGE-AI Horizon Europe template style.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

WORKDIR = "/home/oem/Desktop/Horizon"
MD_FILE  = os.path.join(WORKDIR, "EIC_Pathfinder_CDATA_PartB.md")
OUT_FILE = os.path.join(WORKDIR, "EIC_Pathfinder_CDATA_PartB.docx")

# ── Colour palette (Horizon Blue) ─────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x0F, 0x1F, 0x3D)   # #0f1f3d  — headings
BLUE_MED    = RGBColor(0x00, 0x50, 0x9E)   # Horizon blue — section heads
GOLD        = RGBColor(0xC9, 0xA8, 0x4C)   # #c9a84c  — accent
GREY_LIGHT  = RGBColor(0xF4, 0xF7, 0xFB)   # table row shade
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run_bold(para, text, size_pt=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run

def add_run_normal(para, text, size_pt=11, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.italic = italic
    return run

# ── Document setup ────────────────────────────────────────────────────────────
def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    return doc

# ── Title page ────────────────────────────────────────────────────────────────
def add_title_page(doc):
    # EIC header box
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, '0F1F3D')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run('Proposal Part B: Technical Description')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = WHITE

    doc.add_paragraph()

    # Project title box
    t2 = doc.add_table(rows=1, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell2 = t2.cell(0, 0)
    set_cell_bg(cell2, 'E8EFF7')
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after  = Pt(10)
    r2 = p2.add_run('CDATA')
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = BLUE_DARK
    p2.add_run('\n')
    r3 = p2.add_run('Centriolar Damage Accumulation Theory of Ageing:\nA Computational Cellular Digital Twin Validated Against Human Longevity Biomarkers')
    r3.font.size = Pt(12)
    r3.font.color.rgb = BLUE_DARK

    doc.add_paragraph()

    # Meta table
    meta = [
        ('Call', 'HORIZON-EIC-2026-PATHFINDEROPEN-01'),
        ('Instrument', 'EIC Pathfinder Open'),
        ('Requested funding', '€2,500,000 (36 months)'),
        ('Lead Beneficiary', 'Phasis Academy (Poti, Georgia)'),
        ('Beneficiary 2', 'Georgian Technical University — ABET BME (PIC: 983636358)'),
        ('Principal Investigator', 'Jaba Tkemaladze, MD'),
        ('Co-PI', 'Prof. Gela [Surname — pending GTU confirmation]'),
        ('Submission deadline', '12 May 2026'),
    ]
    t3 = doc.add_table(rows=len(meta), cols=2)
    t3.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        row = t3.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = BLUE_DARK
        if i % 2 == 0:
            set_cell_bg(row.cells[0], 'E8EFF7')
            set_cell_bg(row.cells[1], 'F4F7FB')
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_page_break()

# ── Heading styles ─────────────────────────────────────────────────────────────
def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BLUE_MED
    # underline via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A84C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = BLUE_DARK
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BLUE_DARK
    return p

def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    # Handle inline bold **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(11)
        else:
            r = p.add_run(part)
            r.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(10.5)
    return p

def add_table_from_md(doc, lines):
    """Parse a markdown table and render it as a formatted docx table."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or re.match(r'^\|[-| :]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # Pad rows
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = t.rows[i]
        is_header = (i == 0)
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            # Handle bold **text** in cells
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.size = Pt(9)
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.italic = True
                    r.font.size = Pt(9)
                else:
                    r = p.add_run(part)
                    r.font.size = Pt(9)
            if is_header:
                set_cell_bg(cell, '0F1F3D')
                for run in p.runs:
                    run.font.color.rgb = WHITE
                    run.bold = True
            elif i % 2 == 1:
                set_cell_bg(cell, 'F4F7FB')
    doc.add_paragraph()

# ── "In a nutshell" summary box ───────────────────────────────────────────────
def add_nutshell_box(doc, rows):
    """rows = list of (label, text)"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Table Grid'
    col_w = [Cm(3), Cm(13)]
    for i, (label, text) in enumerate(rows):
        row = t.rows[i]
        # Label cell
        lc = row.cells[0]
        set_cell_bg(lc, '0F1F3D')
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.color.rgb = WHITE
        lr.font.size = Pt(11)
        # Text cell
        tc = row.cells[1]
        set_cell_bg(tc, 'F4F7FB' if i % 2 == 0 else 'FFFFFF')
        tp = tc.paragraphs[0]
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = tp.add_run(part[2:-2])
                r.bold = True
                r.font.size = Pt(10)
            else:
                r = tp.add_run(part)
                r.font.size = Pt(10)
    doc.add_paragraph()

# ── Main parser ───────────────────────────────────────────────────────────────
def parse_and_render(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    nutshell_rows = []
    in_nutshell = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip YAML-like header lines
        if stripped.startswith('> ') and i < 10:
            i += 1
            continue

        # Skip horizontal rules
        if re.match(r'^[-*]{3,}$', stripped):
            i += 1
            continue

        # Page break on ---
        if stripped == '---' and i > 5:
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            add_h4(doc, stripped[5:])
            i += 1
            continue
        if stripped.startswith('### '):
            add_h3(doc, stripped[4:])
            i += 1
            continue
        if stripped.startswith('## '):
            add_h2(doc, stripped[3:])
            i += 1
            continue
        if stripped.startswith('# '):
            add_h1(doc, stripped[2:])
            i += 1
            continue

        # Markdown table
        if stripped.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            # Check indent level
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            add_bullet(doc, text, level=level)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            add_bullet(doc, text, level=0)
            i += 1
            continue

        # Code block (skip — Gantt is ASCII, render as monospace)
        if stripped.startswith('```'):
            i += 1  # skip opening ```
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                for cl in code_lines[:40]:  # limit Gantt
                    r = p.add_run(cl + '\n')
                    r.font.name = 'Courier New'
                    r.font.size = Pt(7)
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        if stripped:
            add_body(doc, stripped)

        i += 1

# ── Footer ────────────────────────────────────────────────────────────────────
def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run('CDATA EIC Pathfinder Open 2026 | Phasis Academy + GTU | PI: Jaba Tkemaladze MD')
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        fp.add_run('  |  Page ')
        # page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run2 = fp.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.size = Pt(8)

# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    print(f"Reading {MD_FILE}...")
    with open(MD_FILE, encoding='utf-8') as f:
        md_text = f.read()

    doc = setup_document()
    add_title_page(doc)

    # "In a nutshell" box (extract from markdown)
    nutshell_rows = [
        ('WHY',      'Ageing lacks a mechanistic root cause. CDATA identifies centrioles — permanently damage-accumulating, repair-deficient organelles — as the primary initiator.'),
        ('WHAT',     'First computational cellular digital twin (Cell-DT) integrating 7 biological ageing tracks and 20+ mechanisms, calibrated to 78.4-year human lifespan. Novel CAII biomarker. 4 CDATA-specific therapeutic approaches.'),
        ('HOW',      'WP1: In vitro PTM–CAII dose-response validation. WP2: Longitudinal human cohort (n=288, 66–80 yr). WP3: Cell-DT experimental parameterisation & personalised biological age. WP4: Therapeutic PoC + IP + translation.'),
        ('FOR WHOM', 'Biogerontologists, precision medicine clinicians, longevity biotech, pharma drug developers, EU healthy-ageing policymakers.'),
        ('BY WHOM',  'Phasis Academy (PI: Dr. Jaba Tkemaladze MD, CDATA originator, 2 publications, Cell-DT author) + GTU ABET BME (STED confocal, Horizon FORGE-AI track record, PIC 983636358).'),
    ]

    add_h2(doc, 'CDATA in a Nutshell')
    add_nutshell_box(doc, nutshell_rows)

    parse_and_render(doc, md_text)
    add_footer(doc)

    doc.save(OUT_FILE)
    print(f"✅ Saved: {OUT_FILE}")
    size_kb = os.path.getsize(OUT_FILE) // 1024
    print(f"   Size: {size_kb} KB")

if __name__ == '__main__':
    main()
