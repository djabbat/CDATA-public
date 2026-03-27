#!/usr/bin/env python3
"""Generate Georgian translation of EIC_Pathfinder_CDATA_PartB_v7.md as .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

OUTPUT_PATH = "/home/oem/Desktop/Horizon/EIC_Pathfinder_CDATA_PartB_v7_KA.docx"
FONT_NAME = "Sylfaen"
FONT_SIZE_BODY = 11
LINE_SPACING = 1.15

# ─── Document setup ───────────────────────────────────────────────────────────

doc = Document()

# Page margins 2.5cm all sides
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Styles ───────────────────────────────────────────────────────────────────

def set_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(para, before=0, after=4, line_spacing=LINE_SPACING):
    from docx.shared import Pt as _Pt
    para.paragraph_format.space_before = _Pt(before)
    para.paragraph_format.space_after = _Pt(after)
    para.paragraph_format.line_spacing = line_spacing

def add_heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    para = doc.add_paragraph()
    run = para.add_run(text)
    size = sizes.get(level, 11)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = True
    run.italic = (level == 4)
    para.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    para.paragraph_format.space_after = Pt(4)
    # Apply heading style
    style_name = f"Heading {level}"
    try:
        para.style = doc.styles[style_name]
        # Override the style font
        for r in para.runs:
            r.font.name = FONT_NAME
            r.font.size = Pt(size)
            r.bold = True
            r.italic = (level == 4)
    except Exception:
        pass
    return para

def add_body(doc, text, bold_parts=None, italic_parts=None):
    para = doc.add_paragraph()
    para_spacing(para)
    # Parse inline bold/italic markers
    _add_inline_text(para, text)
    return para

def _add_inline_text(para, text):
    """Parse **bold** and *italic* markers and add runs."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            set_font(run, italic=True)
        else:
            run = para.add_run(part)
            set_font(run)

def add_blockquote(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.2)
    para_spacing(para)
    _add_inline_text(para, text)
    for run in para.runs:
        run.italic = True
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    para_spacing(para, after=2)
    _add_inline_text(para, text)
    return para

def add_separator(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para_spacing(para, before=4, after=4)
    return para

def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_table_from_data(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_shading(cell, 'D9E1F2')
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_font(run, size=10, bold=True)
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            para = cell.paragraphs[0]
            _add_inline_text(para, str(val))
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(10)
    return table

# ─── Footer ───────────────────────────────────────────────────────────────────

def add_footer_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        run = para.add_run("გვერდი ")
        set_font(run, size=9)
        # PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r = OxmlElement('w:r')
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        para._p.append(r)
        run2 = para.add_run(" / ")
        set_font(run2, size=9)
        # NUMPAGES field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES'
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r2 = OxmlElement('w:r')
        r2.append(fldChar3)
        r2.append(instrText2)
        r2.append(fldChar4)
        para._p.append(r2)

add_footer_page_numbers(doc)

# ─── CONTENT ──────────────────────────────────────────────────────────────────

# Translation note at top
note_para = doc.add_paragraph()
note_para.paragraph_format.space_after = Pt(8)
note_run = note_para.add_run(
    "სამუშაო თარგმანი — ორიგინალი ინგლისურ ენაზე: EIC_Pathfinder_CDATA_PartB_v7.md"
)
set_font(note_run, size=10, italic=True)
note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ─── TITLE / HEADER BLOCK ─────────────────────────────────────────────────────

add_heading(doc, "წინადადება, ნაწილი B: ტექნიკური აღწერა", 1)

add_blockquote(doc,
    "**CDATA** — **C**ენტრიოლური **D**ეგრადაციის **A**კუმულაციის **T**ეორია **A**დარება:\n"
    "კომპიუტაციული უჯრედული ციფრული ტყუპი, ვალიდირებული ადამიანის სიგრძის ბიომარკერების მიხედვით"
)

meta = [
    ("**ზარი:**", "HORIZON-EIC-2026-PATHFINDEROPEN-01"),
    ("**ინსტრუმენტი:**", "EIC Pathfinder Open"),
    ("**მოთხოვნილი დაფინანსება:**", "€2,500,000 (36 თვე)"),
    ("**მთავარი ბენეფიციარი:**", "ფაზისის აკადემია (ფოთი, საქართველო)"),
    ("**ბენეფიციარი 2:**", "საქართველოს ტექნიკური უნივერსიტეტი — GTU ABET BME (თბილისი, საქართველო), PIC: 983636358"),
    ("**მთავარი გამომძიებელი:**", "ჯაბა თქემალაძე, MD"),
    ("**თანა-PI:**", "პროფ. ბატონი გელა, PhD — ბიომედიცინური ინჟინერიის დეპარტამენტი, GTU ABET; FORGE-AI Horizon კონსორციუმის წევრი"),
]
for label, val in meta:
    p = doc.add_paragraph()
    para_spacing(p, after=2)
    _add_inline_text(p, f"{label} {val}")

add_separator(doc)

# ─── TABLE OF CONTENTS ────────────────────────────────────────────────────────

add_heading(doc, "შინაარსის ცხრილი", 2)

toc_items = [
    "0. ცვლილებების ისტორია",
    "1. სრულყოფილება",
    "   1.1 კონცეფციის სიმყარე და მეთოდოლოგიის სარწმუნოება",
    "   1.2 სამეცნიერო ამბიცია — მდგომარეობის გადალახვა",
    "2. გავლენა",
    "   2.1 მეცნიერული და ტექნოლოგიური გარღვევის პოტენციალი",
    "   2.2 საზოგადოებრივი გავლენა და EU პოლიტიკასთან შესაბამისობა",
    "   2.3 გავრცელება, ექსპლუატაცია და კომუნიკაცია",
    "3. განხორციელების ხარისხი და ეფექტურობა",
    "   3.1 სამუშაო გეგმა და რესურსები",
    "   3.2 კონსორციუმის შესაძლებლობები",
    "   3.3 ბიუჯეტის დაშლა",
    "   3.4 რისკების მართვა",
    "   3.5 ღია მეცნიერების პრაქტიკა",
    "4. ეთიკის თვითშეფასება",
]
for item in toc_items:
    p = doc.add_paragraph()
    para_spacing(p, after=1)
    run = p.add_run(item)
    set_font(run, size=10)

add_separator(doc)

# ─── SECTION 0 ────────────────────────────────────────────────────────────────

add_heading(doc, "0. ცვლილებების ისტორია", 2)

add_table_from_data(doc,
    ["ვერსია", "თარიღი", "ცვლილების ხასიათი"],
    [
        ["v1.0", "2026-03-21", "საწყისი სამუშაო ვარიანტი"],
        ["v2.0", "2026-03-23",
         "ბიუჯეტი გასწორდა: პირდაპირი €2,000K + 25% ირიბი = €2,500K. GTU co-PI placeholder შეიყვანეს. n=288/240 გაერთიანდა. პიროვნება-თვეების ცხრილი დაემატა. WP ლიდერები განისაზღვრა. Peer Review v3: 4.3/5."],
        ["v5.0", "2026-03-24",
         "Peer Review v4 კორექტურები: (C2) Ze Theory ამოღებულ იქნა 1-ლი განყოფილების Excellence-დან; რჩება მხოლოდ WP2-ის მეორეული საბოლოო წერტილად. (C3) Consortium Agreement გასწორდა M1-მდე. (C4) ფაზისის აკადემიის აღწერა გაფართოვდა. WP4-ის მოსამზადებელი დავალება T4.0 დაემატა. In silico შენიშვნები ჩართული."],
        ["v6.0", "2026-03-24",
         "გარე ექსპერტის Peer Review კორექტურები: (E1) Ze Theory აღდგა 1.2-ში, როგორც მეორეული ინოვაციური ჰიპოთეზა. (Q2) STED-ის აღწერა გასწორდა: ახალი Leica SP8 შეძენა. (Q3) ბრმა გამოკვლევა WP2-ში. (E3) პირველადი HDF უჯრედები WP1 T1.1-ში. (E2) ტელომერების თეორია შერბილდა. KPI-ები შემცირდა: #10 r>0.5, #12 mechanism-proof. (Q4) Rust Engineer-ის საერთაშორისო რეკრუტირება. (I2) CAII ELISA ავტომატური პლატფორმის გეგმა."],
        ["v7.0", "2026-03-24",
         "მეორე ექსპერტის მიმოხილვა + co-PI დადასტურება: (C2) ვერსია v7.0. (C3) Ze Theory 1.2.3 — წინასწარ გამოუქვეყნებელი მონაცემები არ არის. (M1) ≥439 ტესტი ყველა განყოფილებაში. (M2) KPI #16 CAII ELISA ავტომატური პლატფორმა. (M3) Gladyshev 2014 მიმართვა [11]. (L1) Leica SP8 STED 2.2.3-ში. (Q1) co-PI დადასტურდა: პროფ. ბატონი გელა, GTU ABET BME, FORGE-AI."],
    ]
)

add_separator(doc)

# ─── SECTION 1 ────────────────────────────────────────────────────────────────

add_heading(doc, "1. სრულყოფილება", 1)

add_heading(doc, "CDATA მოკლედ", 3)

# Nutshell table
add_table_from_data(doc,
    ["", ""],
    [
        ["**რატომ (WHY)**",
         "დაბერება ყველა ძირითადი არაგადამდები დაავადების წამყვანი რისკ-ფაქტორია, მაგრამ მისი პირველადი უჯრედული მამოძრავებელი კვლავ უცნობია. არსებული თეორიები (Hallmarks of Ageing, ტელომერების თეორია, Inflammaging) ფენომენოლოგიური აღწერებია, არა მექანიკური ძირეული მიზეზები. აუცილებელია პარადიგმის ცვლა: უნდა გამოიყოს *შეუქცევადი მოლეკულური ინიციატორი*."],
        ["**რა (WHAT)**",
         "CDATA ამტკიცებს, რომ **ცენტრიოლები** — ერთადერთი არა-მემბრანული უჯრედული ორგანელა, რომელსაც სარემონტო მექანიზმი სრულიად აკლია — ეს ინიციატორია. PTM-ების შეუქცევადი დაგროვებით დაზიანებული ცენტრიოლები ანგრევენ პირველად ციებს და მიტოზური ღერძის სიმყარეს, ამოწურავენ ღეროვანი უჯრედების პულს და იწვევენ სისტემური დაბერებას. CDATA არის პირველი კომპიუტაციული Cell Digital Twin, რომელიც 7 ბიოლოგიურ ბილიკს და 20+ მექანიზმს აერთიანებს 78.4-წლიანი სიცოცხლისუნარიანობის ერთ პრედიქტიულ მოდელში."],
        ["**როგორ (HOW)**",
         "ფაზა 1 (WP1–WP2): CDATA ბიომარკერების (CAII, CEP164, γ-H2AX) ვალიდაცია გრძელვადიან ადამიანის კოჰორტში (n=288, ასაკი 66–80) და in vitro ცენტრიოლური დაზიანების მოდელებში. ფაზა 2 (WP3): ექსპერიმენტული მონაცემების Cell Digital Twin-ში (Cell-DT, 14 ECS მოდული, ≥439 ვალიდირებული ტესტი) ინტეგრაცია. ფაზა 3 (WP4): მკურნალობის Proof-of-Concept სამი ახალი CDATA-სპეციფიკური ჩარევით (CentrosomeTransplant, CafdRetainer, Senolytics)."],
        ["**ვისთვის (FOR WHOM)**",
         "ბიოგერონტოლოგები, პრეციზიული მედიცინის კლინიცისტები, სიგრძის ბიოტექნიკური კომპანიები, ფარმაცევტული შემქმნელები, EU ჯანმრთელი დაბერების პოლიტიკის შემქმნელები."],
        ["**ვის მიერ (BY WHOM)**",
         "ორ-ინსტიტუციური კონსორციუმი: ფაზისის აკადემია (PI: დრ. ჯაბა თქემალაძე, MD — CDATA თეორიის ავტორი; 2 peer-reviewed პუბლიკაცია, მათ შ. Mol Biol Rep 2023, PMID 36583780) + GTU ABET BME (co-PI: პროფ. ბატონი გელა, Confocal/STED ბირთვული სიმძლავრე, Horizon FORGE-AI გამოცდილება)."],
    ]
)

add_separator(doc)

# ─── 1.1 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "1.1 პროექტის კონცეფციის სიმყარე და მეთოდოლოგიის სარწმუნოება", 2)

add_heading(doc, "1.1.1 პრობლემა: დაბერების გამოტოვებული ძირეული მიზეზი", 3)

add_body(doc,
    "კვლევის ათეული წლის მიუხედავად, დაბერებას მოლეკულური დონეზე მექანიკური, მიზეზობრივი ახსნა კვლავ არ გააჩნია. დომინანტური ჩარჩო — დაბერების 12 Hallmarks (López-Otín et al., 2023, *Cell* 186:243) — ჩამოთვლის დაბერების *შედეგებს*, მაგრამ ერთ ამოსავალ მოვლენას არ ადგენს. ტელომერების შემცირება (Blackburn, 2000, *Nature* 408:53) ხსნის Hayflick-ით შეზღუდულ უჯრედის გაყოფებს, მაგრამ არა პოსტ-მიტოტური უჯრედის დაბერებას. Inflammaging (Franceschi et al., 2007) იმუნური დისრეგულაციის ქვედინური შედეგია."
)
add_body(doc,
    "ცენტრალური კითხვა პასუხგაუცემელია: **რა არის პირველი, შეუქცევადი მოლეკულური მოვლენა, რომელიც კასკადს იწყებს?**"
)
add_body(doc,
    "CDATA გვთავაზობს დამაჯერებელ პასუხს: **ცენტრიოლები**."
)

add_heading(doc, "1.1.2 CDATA ჰიპოთეზა", 3)

add_body(doc, "ცენტრიოლები უნიკალურია უჯრედული ორგანელებს შორის ორი ძირითადი მახასიათებლის გამო:")
add_bullet(doc, "**ისინი მუდმივია:** უმეტეს ცილებისგან განსხვავებით, რომლებიც დღეებიდან კვირებამდე ბრუნავს, ცენტრიოლები უჯრედის მთელ სიცოცხლეს გრძელდება. დედა ცენტრიოლი *არასდროს* იცვლება.")
add_bullet(doc, "**სარემონტო მექანიზმი მათ სრულიად აკლია:** არცერთი ცნობილი დეუბიკვიტილაზა, პროტეაზა ან თაბუ-სითბური შაპერონი ცენტრიოლური სტრუქტურული ცილების სარემონტოდ არ გამოიყენება (Tkemaladze, 2023, *Mol Biol Rep*, PMID 36583780).")

add_body(doc, "შედეგად, ცენტრიოლები შეუქცევადად აგროვებენ **PTM (post-translational modifications)**:")
add_bullet(doc, "ტუბულინის აცეტილაცია (პროგრესული დაბადებიდან)")
add_bullet(doc, "ოქსიდური კარბონილაცია (ROS-ის მამოძრავებელი, თვით-გამაძლიერებელი მარყუჟი)")
add_bullet(doc, "ფოსფორილაციის გადახრა (M-ფაზის კინაზური აქტივობის ნარჩენები)")
add_bullet(doc, "გლუტამილაციის/გლიცილაციის დისბალანსი (მიკროტუბულების დინამიკას არღვევს)")

add_body(doc,
    "ეს PTM-ები დიფერენციულად გროვდება დედა (უფრო ძველ) და ასული (უფრო ახალ) ცენტრიოლებში **ასიმეტრიული შენარჩუნების წესის** გამო: ასიმეტრიულ ღეროვანი უჯრედის გაყოფისას, ღეროვანი ასული უჯრედი *შეარჩევს უფრო ძველ, დაზიანებულ დედა ცენტრიოლს*. ეს არის მოლეკულური საფუძველი 'დაბერების ბეჭდის' ღეროვანი უჯრედული თაობებში გადაცემისა."
)

add_body(doc, "ქვედინური შედეგები ორ დამოუკიდებელ ბილიკს ქმნის:")
add_bullet(doc, "**ბილიკი A (ციები):** CEP164/CEP89/Ninein/CEP170 დამატების მთლიანობა მცირდება → პირველადი ციების ასამბლეა ვერ ხერხდება → Shh/Wnt/Notch ნიშის სიგნალიზაცია ინგრევა → ღეროვანი უჯრედის თვით-განახლება ჩერდება")
add_bullet(doc, "**ბილიკი B (ღერძი):** Spindle assembly checkpoint-ის გვერდის ავლა → სიმეტრიული გაყოფები იზრდება → ღეროვანი უჯრედის პულის ამოწურვა ჩქარდება")

add_body(doc,
    "ეს ორი ბილიკი, ოთხ გამაძლიერებელ მარყუჟთან (ROS კასკადი, SASP/Inflammaging, ტელომერების შემოკლება, ეპიგენეტური საათის ჩქარება) ერთად, ხსნის სრულ დაბერების ტრაექტორიას TRL 1-დან (მოლეკულური მექანიზმი) TRL 4-მდე (ვალიდირებული კომპიუტაციული მოდელი)."
)

add_heading(doc, "1.1.3 წინა სამუშაო: გამოქვეყნებული საფუძველი", 3)

add_table_from_data(doc,
    ["პუბლიკაცია", "ჟურნალი", "წელი", "ძირითადი წვლილი"],
    [
        ["Tkemaladze J, Lortkipanidze G. \"Asymmetric segregation of the centrosome and its role in determining cell fate\"",
         "*Mol Biol Rep* (Springer)", "2023",
         "**PMID: 36583780** — ადგენს ასიმეტრიული ცენტრიოლური სეგრეგაციის მექანიზმს და მის როლს ღეროვანი უჯრედის ბედის განსაზღვრაში. PubMed/Scopus-ში ინდექსირებული."],
        ["Tkemaladze J. \"CDATA: Centriolar Damage Accumulation Theory of Ageing\"",
         "*Ann Rejuvenation Sci*", "2025",
         "**DOI: 10.65649/yx9sn772** — წარმოადგენს CDATA-ს სრულ თეორიულ ჩარჩოს კომპიუტაციული ვალიდაციით."],
    ]
)

add_heading(doc, "1.1.4 Cell Digital Twin (Cell-DT): კომპიუტაციული Proof of Concept", 3)

add_body(doc,
    "PI-მ შეიმუშავა **Cell-DT** — Rust-ზე დაფუძნებული მაღალი წარმადობის უჯრედული ციფრული ტყუპი, CDATA-ს განხორციელება — რეპროდუცირებადი, ღია კოდის კომპიუტაციული პლატფორმი (GitHub: djabbat/CDATA-Longevity):"
)

add_bullet(doc, "**არქიტექტურა:** ECS (Entity-Component-System) დიზაინი `hecs` კრეიტის გამოყენებით; 14 სიმულაციური მოდული; ≥439 გამვლელი unit ტესტი (+ ახალი ტესტები WP3-ში დამატებული მოდულებისთვის) (cargo test); პარალელიზებული Rayon-ის მეშვეობით.")
add_bullet(doc, "**კალიბრირებული გამოსავალი:** DamageParams-ის ნაგულისხმევი მნიშვნელობებით (base_detach_probability = 0.0003; senescence_threshold = 0.75), მოდელი პროგნოზირებს: სიცოცხლის ხანგრძლივობა **78.4 წელი** (CV = 5.6% n=1000 სტოქასტური გაშვების მასშტაბით); ასაკ-70-ზე მიელოიდური მიკერძოება: **0.45**; CHIP გამოჩენა: წელი 40; Hayflick დაპატიმრება: G1 შეზღუდვა ტელომერის სიგრძის < 0.15.")

add_body(doc, "**ექვსი ვალიდირებული სიმულაციის ბილიკი:**")

add_table_from_data(doc,
    ["ბილიკი", "მექანიზმი", "ძირითადი გამოსავალი"],
    [
        ["A (ციები)", "CEP164↓ → Shh/Wnt↓ → regeneration_tempo↓", "ciliary_function: 1.0 → 0.12 სიკვდილისას"],
        ["B (ღერძი)", "spindle_fidelity↓ → სიმეტრიული გაყოფა ↑", "pool_exhaustion 78-ე წელს"],
        ["C (ტელომერი)", "შემოკლება × გაყოფის სიჩქარე × spindle კორექცია", "G1 დაპატიმრება ~50-ე წელს"],
        ["D (ეპიგენომი)", "methylation_age += dt × (1 + damage × 0.5)", "clock_acceleration: 1.0 → 1.85"],
        ["E (მიტოქონდრია)", "mtDNA მუტაციები → ROS↑ → mito_shield↓", "თვით-გამაძლიერებელი O₂ მარყუჟი"],
        ["F (გაყოფის სიჩქარე)", "ციები × ღერძი × ასაკი × ROS × mTOR ინტეგრაცია", "stem_cell_pool: 1.0 → 0.04"],
    ]
)

add_heading(doc, "1.1.5 CDATA კონკურენტ თეორიებთან შედარება", 3)

add_table_from_data(doc,
    ["თეორია", "მიზეზობრივი ფესვი", "შეუქცევადი?", "ხსნის პოსტ-მიტოტურ დაბერებას?", "CDATA-ს უპირატესობა"],
    [
        ["**CDATA** (თქემალაძე)", "ცენტრიოლური PTM-ების დაგროვება", "✅ დიახ", "✅ დიახ (ციების მეშვეობით)", "—"],
        ["Hallmarks of Ageing (López-Otín 2023 [1])", "მრავლობითი, იერარქიის გარეშე", "ნაწილობრივ", "✅", "ძირეული მიზეზი არ არის"],
        ["ტელომერების თეორია (Blackburn 2000 [2])", "ტელომერების შემოკლება", "✅", "⚠️ ხსნის მხოლოდ რეპლიკატიულ სენესსენსს; პოსტ-მიტოტური დისფუნქცია არ არის გათვალისწინებული†", "არასრული: CDATA ავსებს პოსტ-მიტოტურ დონეზე"],
        ["Inflammaging (Franceschi 2007)", "ქრონიკული ანთება", "ნაწილობრივ", "✅", "ქვედინური, არა მიზეზობრივი"],
        ["თავისუფალი რადიკალების თეორია (Harman 1956)", "ROS-ის დაგროვება", "ნაწილობრივ", "✅", "ROS მოდულატორია, არა ინიციატორი [11]"],
    ]
)

add_heading(doc, "1.1.6 მეთოდოლოგია: WP1 — In Vitro ვალიდაცია", 3)

add_body(doc, "**მიზანი:** ცენტრიოლური PTM-ების ტვირთსა და ბილიკ A/B-ის დისფუნქციას შორის რაოდენობრივი კავშირის დადგენა ადამიანის უჯრედულ ხაზებში და პირველად უჯრედებში.")

add_body(doc, "**ექსპერიმენტული დიზაინი:**")
add_bullet(doc, "**უჯრედული მოდელები:** (ა) *პროტოკოლის შემუშავება:* HeLa უჯრედები (მაღალი პასაჟი) U-ExM და STED პროტოკოლის ოპტიმიზაციისთვის. (ბ) *პირველადი დოზა-პასუხის ექსპერიმენტები:* IMR90 ფიბრობლასტები და **პირველადი ადამიანის კანის ფიბრობლასტები (HDF)** ახალგაზრდა (25–35 წ.) vs ასაკოვანი (65–75 წ.) დონორებისგან. (გ) *Ex vivo:* CD34⁺ EPC-ები (MACS პროტოკოლი; Miltenyi; 20 მლ სისხლი).")
add_bullet(doc, "**ცენტრიოლური დაზიანების ინდუქცია:** წყალბადის პეროქსიდი, okadaic acid, Taxol, N-hydroxysuccinimide")
add_bullet(doc, "**ძირითადი გამოსავლები:** CEP164/CEP89/Ninein/CEP170 IF; **CAII ინდექსი** (ახალი CDATA ბიომარკერი); γ-H2AX; U-ExM; პირველადი ცია სიგრძე; Spindle fidelity")

add_body(doc, "†*შენიშვნა: ტელომერების შემოკლება ხსნის Hayflick-ით შეზღუდულ სენესსენსს; CDATA ავრცელებს ამ ჩარჩოს პოსტ-მიტოტურ უჯრედებზე (ნეირონები, კარდიომიოციტები), სადაც ცენტრიოლური დაზიანება ციების დისფუნქციას იწვევს.*")

add_heading(doc, "1.1.7 მეთოდოლოგია: WP2 — ადამიანის გრძელვადიანი კოჰორტის კვლევა", 3)

add_body(doc, "**მიზანი:** CDATA ბიომარკერების (CAII, CEP164) კორელაცია კლინიკური დაბერების ფენოტიპებთან (სისუსტე, კოგნიტური კლება, იმუნური სენესსენსი) კარგად დახასიათებულ ქართულ ხანდაზმულ კოჰორტში.")

add_body(doc, "**კვლევის დიზაინი:**")
add_bullet(doc, "**დიზაინი:** პროსპექტიული სადაკვირვებლო, 2-წლიანი follow-up")
add_bullet(doc, "**მონაწილეები:** ჩარიცხული n=288, შეფასებადი n=240 (16.7% შემცირება)")
add_bullet(doc, "**ჩართვა:** ასაკი 66–80 წელი; ორივე სქესი; თბილისის მეტროპოლიური არეალი")
add_bullet(doc, "**გამორიცხვა:** აქტიური ონკოლოგია, ავტოიმუნური დაავადება, მძიმე კოგნიტური დარღვევა")
add_bullet(doc, "**ეთიკა:** Mini-Cog სკრინინგი; ორ-ეტაპიანი ინფორმირებული თანხმობა (ICH E6(R2) GCP); GTU-ს ეთიკის კომიტეტი; ჰელსინკის დეკლარაციის შესაბამისი")
add_bullet(doc, "**ბიომარკერები:** სისხლის CEP164 ცილა (ELISA); PBMC CAII; ტელომერის სიგრძე (qPCR); ეპიგენეტური ასაკი (DNAm PhenoAge, Levine 2018); HRV ინდექსი; CRP, IL-6, TNF-α, IGF-1")
add_bullet(doc, "**კლინიკური ფენოტიპირება:** Fried Frailty Index; MMSE; ხელის სიმძლავრე; სვლის სიჩქარე; Barthel ADL ინდექსი")
add_bullet(doc, "**პირველადი საბოლოო წერტილი:** CAII-ის კლების კავშირი Fried Frailty პროგრესიასთან (OR, ლოგისტიკური რეგრესია)")
add_bullet(doc, "**სტატისტიკური სიმძლავრე:** OR = 2.5; α=0.05, power=0.80; n=240 > n=196")

add_heading(doc, "1.1.8 მეთოდოლოგია: WP3 — Cell-DT ინტეგრაცია", 3)

add_body(doc, "**მიზანი:** WP1 და WP2-ის ემპირიული მონაცემების Cell-DT-ში ჩართვა, კალიბრირებული, ვალიდირებული სიგრძის პროგნოზული ძრავის შესაქმნელად.")

add_body(doc, "**ძირითადი ამოცანები:**")
add_bullet(doc, "T3.1: Cell-DT-ის ნაგულისხმევი PTM კინეტიკური პარამეტრების WP1 ექსპერიმენტული მონაცემებით ჩანაცვლება (Bayesian პარამეტრების შეფასება, Stan/PyMC)")
add_bullet(doc, "T3.2: პაციენტ-სპეციფიკური პარამეტრების მორგება: CAII + ეპიგენეტური ასაკის ინდივიდუალური შეყვანა → პერსონალიზებული DamageParams (pBA)")
add_bullet(doc, "T3.3: pBA-ს ვალიდაცია WP2-ის 2-წლიანი შედეგების მიხედვით (Kaplan-Meier, AUROC სამიზნე > 0.75)")
add_bullet(doc, "T3.4: Cell-DT-ის გაფართოება: IGF-1 ღერძი, asymmetric_cytoqc_module, PTM CSV ექსპორტი")
add_bullet(doc, "T3.5: Cell-DT v2.0 გამოქვეყნება (Zenodo DOI + GitHub; Python binding-ები PyO3)")

add_heading(doc, "1.1.9 მეთოდოლოგია: WP4 — მკურნალობის Proof of Concept", 3)

add_body(doc, "**მიზანი:** CDATA-სპეციფიკური ჩარევების დემონსტრაცია, რომ ისინი ახანგრძლივებენ ფუნქციურ ჯანმრთელობის ხანგრძლივობას in vitro დაბერების მოდელებში.")

add_body(doc, "**Cell-DT in silico პროგნოზებზე დაფუძნებული ჩარევის ბიბლიოთეკა:**")

add_table_from_data(doc,
    ["ჩარევა", "CDATA მიზნობრივი მექანიზმი", "Cell-DT in silico პროგნოზი*"],
    [
        ["**CentrosomeTransplant** (ახალგაზრდა დონორის ცენტრიოლების გადანერგვა)",
         "ცენტრიოლური PTM ტვირთის დონორის დონეზე გადაყენება", "+18.3 წელი"],
        ["**CafdRetainer** (CAFD inducer კომპლექსების სტაბილიზაცია)",
         "O₂-მამოძრავებელი inducer-ის გამოყოფის შენელება", "+11.6 წელი"],
        ["**Senolytics** (ABT-263 + Quercetin პროტოკოლი)",
         "SASP-ის მომყვანი სენესსენტური უჯრედების ამოღება → ROS-ის შემცირება", "+8.9 წელი"],
        ["**mTOR-ის ინჰიბიცია + CR** (Rapamycin + კალორიული შეზღუდვა)",
         "გაყოფის სიჩქარის შენელება; ავტოფაგიის გააქტიურება", "+7.4 წელი"],
    ]
)

add_body(doc, "*\\*ყველა ფიგურა Cell-DT-ის სტოქასტური სიმულაციის შედეგია (n=1000 გაშვება). ამ პროგნოზების ექსპერიმენტული ვალიდაცია WP4-ის ძირითადი მიზანია.*")

add_separator(doc)

# ─── 1.2 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "1.2 სამეცნიერო სიახლის და ამბიციის მასშტაბი", 2)

add_heading(doc, "1.2.1 სამეცნიერო ორიგინალობა", 3)

add_body(doc, "CDATA ამოიღებს სამ ორიგინალ, ცრუობადი პრეტენზიას:")
add_bullet(doc, "**პრეტენზია 1 — შეუქცევადობა:** დაბერება ცენტრიოლებიდან იწყება, რადგან ისინი ერთადერთი ძირითადი უჯრედული სტრუქტურაა, რომელსაც სარემონტო მანქანა სრულიად აკლია. **ცრუობის კრიტერიუმი:** ცენტრიოლ-სპეციფიკური სარემონტო ფერმენტის აღმოჩენა CDATA-ს გააბათილებდა.")
add_bullet(doc, "**პრეტენზია 2 — ასიმეტრიული მემკვიდრეობა:** ყოველ ასიმეტრიულ გაყოფაში ღეროვანი ასული უჯრედი ინარჩუნებს *უფრო ძველ, დაზიანებულ* დედა ცენტრიოლს. **ცრუობის კრიტერიუმი:** ექსპერიმენტული დემონსტრაცია, რომ ღეროვანი უჯრედები შემთხვევით ანაწილებენ ცენტრიოლებს.")
add_bullet(doc, "**პრეტენზია 3 — CAII ბიომარკერი:** CAII არის რაოდენობრივი, გაზომვადი ბიომარკერი, რომელიც წინ უსწრებს და პროგნოზირებს დაჩქარებული დაბერების ფენოტიპებს. **ცრუობის კრიტერიუმი:** CAII-სა და სისუსტეს შორის კორელაციის არარსებობა WP2-ში.")

add_heading(doc, "1.2.2 ტექნოლოგიური ამბიცია: Cell Digital Twin", 3)

add_body(doc,
    "Cell-DT წარმოადგენს ღია ლიტერატურაში ყველაზე მექანიკურად ყოვლისმომცველ დაბერების სიმულატორს:"
)
add_bullet(doc, "**14 მოდულური ECS კომპონენტი** — NKSurveillanceState, ProteostasisState, CircadianState, AutophagyState, DDRState — 5 დროის მასშტაბზე")
add_bullet(doc, "**≥439 გამვლელი unit ტესტი (+ ახალი ტესტები WP3-ში)** — სამეცნიერო რეპროდუცირებადობის გარანტიები, რომლებიც კონკურენტ დაბერების მოდელებს (ODE-ზე დაფუძნებული) აკლია")
add_bullet(doc, "**პოპულაციური სიმულაციები:** 30-ორგანიზმიანი კოჰორტი SASP bystander ეფექტებით, CHIP დეტექცია 40-ე წელს, Kaplan-Meier-თავსებადი გამოსავალი")
add_bullet(doc, "**ჩარევის პასუხის მრუდები:** 8 ვალიდირებული in silico ჩარევა ჯანმრთელობის ხანგრძლივობის პროგნოზებით")

add_heading(doc, "1.2.3 Ze Theory HRV ბიომარკერი — გამოსაძიებელი ჰიპოთეზა", 3)

add_body(doc,
    "CDATA აერთიანებს ბიოლოგიურ მასშტაბებს, რომლებიც ერთ ჩარჩოში იშვიათად არის დაკავშირებული: ნანო-მასშტაბი (nm, ცენტრიოლური PTM), უჯრედული მასშტაბი (μm, ციები, ღერძი, სენესსენსის მარკერები) და ორგანიზმული მასშტაბი (ms, HRV, Fried Frailty, ეპიგენეტური საათი)."
)
add_body(doc,
    "**Ze Theory HRV ბიომარკერი — გამოსაძიებელი მეორეული ჰიპოთეზა.** PI-ის Ze Theory ვარაუდობს, რომ ბიოლოგიური სისტემები ოპტიმალურ ინფორმაციის ნაკადს ავლენს უგანზომილებო სიჩქარეზე v* ≈ 0.456 (ენტროპია ≈ 1.68 nats). გამოყენებული გულის დინამიკაში, ეს პოტენციური არაინვაზიური ბიომარკერის — HRV-ისგან მიღებული ნაკადის სიჩქარის v*-სთან სიახლოვის — კანდიდატს პროგნოზირებს (სტანდარტული 5-წუთიანი ECG-დან). Ze Theory WP2-ში **მეორეულ, გამოსაძიებელ საბოლოო წერტილად** განიხილება. WP2 პირველ მკაცრ პროსპექტიულ ტესტს n=288 მონაწილეში უზრუნველყოფს."
)

add_separator(doc)

# ─── SECTION 2 ────────────────────────────────────────────────────────────────

add_heading(doc, "2. გავლენა", 1)

add_heading(doc, "2.1 მეცნიერული და ტექნოლოგიური გარღვევის პოტენციალი", 2)

add_heading(doc, "2.1.1 სამეცნიერო გავლენა", 3)

add_body(doc, "CDATA-ს ვალიდაციის შემთხვევაში სამეცნიერო შედეგები ღრმაა:")
add_bullet(doc, "**პარადიგმის ცვლა ბიოგერონტოლოგიაში:** CDATA დაამყარებდა დაბერების პირველ მექანიკურ, ერთიან, ცრუობად თეორიას. ტარგეტი: *npj Aging*, *Nature Aging*, *Cell*.")
add_bullet(doc, "**ახალი დიაგნოსტიკური ბიომარკერი:** CAII — Centriolar Appendage Integrity Index — გახდებოდა პირველი ცენტრიოლ-დაფუძნებული დაბერების ბიომარკერი.")
add_bullet(doc, "**ახალი სამკურნალო სამიზნეები:** CentrosomeTransplant წარმოადგენს ახალ კლასს — *ორგანელების ჩანაცვლების თერაპია*.")

add_heading(doc, "2.1.2 ტექნოლოგიური გავლენა", 3)

add_body(doc, "**Cell Digital Twin (TRL 4 → 6):** პროექტის ბოლოს Cell-DT გადავა ვალიდირებული კომპიუტაციული მოდელიდან (TRL 4) პერსონალიზებულ სიგრძის პროგნოზის ძრავამდე (TRL 5–6). კომერციული გზები:")
add_bullet(doc, "სიგრძის ბიოტექნიკური პლატფორმები (ეპიგენეტური საათებთან, პროტეომიკასთან ინტეგრაცია)")
add_bullet(doc, "ფარმაცევტული CRO სერვისები (in silico დაბერების ანალიზი ადრეული ფაზის ცხოველების კვლევების ჩასანაცვლებლად)")
add_bullet(doc, "სადაზღვევო და სააქტუარო გამოყენებები (ბიოლოგიური vs ქრონოლოგიური ასაკის რისკის ქულა)")

add_heading(doc, "2.1.3 KPI-ები — ძირითადი შესრულების ინდიკატორები", 3)

add_table_from_data(doc,
    ["#", "KPI", "სამიზნე M36-ზე", "ვალიდაციის მეთოდი"],
    [
        ["**სამეცნიერო გამომუშავება**", "", "", ""],
        ["1", "Peer-reviewed პუბლიკაციები Q1 ჟურნალებში", "≥ 4", "ჟურნალის მიღების წერილები"],
        ["2", "Preprint bioRxiv/Zenodo-ზე (ღია წვდომა)", "≥ 2 (M12, M24)", "DOI მინიჭება"],
        ["3", "Cell-DT v2.0 ღია კოდის გამოშვება", "1 (M30)", "GitHub tag + Zenodo DOI"],
        ["4", "Python პაკეტი (PyPI) Cell-DT-სთვის", "1 პაკეტი", "PyPI რეესტრი"],
        ["**ექსპერიმენტული ვალიდაცია**", "", "", ""],
        ["5", "CAII ინდექსი ვალიდირებული უჯრედულ ხაზებში", "≥ 3 მოდელი", "შიდა QC პროტოკოლი"],
        ["6", "კოჰორტის მონაწილეები (WP2)", "n ≥ 288", "IRB სკრინინგის ჟურნალი"],
        ["7", "კოჰორტის მონაწილეები 2-წლიანი follow-up-ით", "n ≥ 240", "კვლევის მონაცემთა ბაზა"],
        ["8", "CAII–Frailty ასოციაცია (პირველადი)", "OR > 1.8, p < 0.05", "სტატისტიკური ანალიზის ანგარიში"],
        ["9", "CAII–ეპიგენეტური საათის კორელაცია", "Pearson r > 0.5", "WP2 ანგარიში"],
        ["10", "Cell-DT vs ექსპერიმენტული ვალიდაცია", "Pearson r > 0.5 (n≥6)", "WP3 ტექნიკური ანგარიში"],
        ["**მკურნალობის PoC**", "", "", ""],
        ["11", "CDATA ჩარევები ტესტირებული in vitro", "≥ 3 ჩარევა", "WP4 ლაბ ანგარიში"],
        ["12", "Proof-of-mechanism ≥ 1 CDATA ჩარევისთვის", "სტ. მნიშვნელოვანი SA-β-Gal ან CAII შემცირება", "WP4 ანგარიში / preprint"],
        ["**ტრანსლაცია**", "", "", ""],
        ["13", "პატენტის განაცხადი (CAII assay / CentrosomeTransplant)", "1", "პატენტის შეტანის ქვითარი"],
        ["14", "სამრეწველო პარტნიორობა / LoI Cell-DT-სთვის", "≥ 1", "ხელმოწერილი LoI"],
        ["15", "კონფერენციაზე გამოსვლები (საერთაშორისო)", "≥ 6", "დადასტურების წერილები"],
        ["**კლინიკური ტრანსლაცია**", "", "", ""],
        ["16", "CAII ELISA ავტომატური პლატფორმასთან თავსებადობა", "განხორციელებადობის ანგარიში (Roche Cobas)", "ტექნიკური ანგარიში"],
    ]
)

add_separator(doc)

# ─── 2.2 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "2.2 საზოგადოებრივი გავლენა და EU პოლიტიკასთან შესაბამისობა", 2)

add_heading(doc, "2.2.1 ჯანსაღი დაბერება: EU-ს პოლიტიკური კონტექსტი", 3)

add_body(doc,
    "დაბერება ყველა ძირითადი არაგადამდები დაავადების (გულ-სისხლძარღვთა, ალცჰეიმერი, კიბო, ტიპი 2 დიაბეტი, სარკოპენია) პირველადი რისკ-ფაქტორია — ერთად **EU-ში ნაადრევი სიკვდილებისა და დაავადების ტვირთის 86% და 77%** (WHO Europe, 2023). EU-ს *Healthier Together* ინიციატივა (2022–2027) პრიორიტეტად ადგენს ჯანსაღ დაბერებას. CDATA პირდაპირ ამ პრიორიტეტს ემსახურება."
)
add_body(doc,
    "საქართველო, EU-ასოცირებული ქვეყანა 2023 წლიდან, ვალდებულია EU ჯანდაცვის კვლევის პოლიტიკის ჰარმონიზაციისკენ. CDATA პროექტი WP2 კოჰორტისთვის მაღალი ხარისხის ქართული ბიომარკერების მონაცემებს წარმოქმნის, **European Health Data Space**-ში წვლილს შეიტანს."
)

add_heading(doc, "2.2.2 სიგრძის ეკონომიკა", 3)

add_body(doc,
    "გლობალური ჯანსაღი სიგრძის ბაზარი **2030 წლისთვის $600 მილიარდს** მიაღწევს (AARP, 2021). CDATA-დან მიღებული ტექნოლოგიები — CAII დიაგნოსტიკური ტესტი, Cell-DT — პირდაპირ ამ ბაზარზეა ორიენტირებული. პროექტი შექმნის სამეცნიერო საფუძველს **≥1 სპინ-ოფ კომპანიისთვის** და IP-ს (CAII assay + CentrosomeTransplant პატენტები) M36-ზე."
)

add_heading(doc, "2.2.3 Widening ქვეყნის სრულყოფა", 3)

add_body(doc, "საქართველო Horizon Europe-ს Widening ქვეყანაა. ეს პროექტი:")
add_bullet(doc, "**საქართველოში პირველ კომპიუტაციული დაბერების კვლევის ჯგუფს** დააარსებს")
add_bullet(doc, "GTU ABET BME-ს Leica SP8 STED სუპერ-რეზოლუციური მიკროსკოპიის სიმძლავრით აღჭურვავს")
add_bullet(doc, "ქართულ ხანდაზმულ კოჰორტ ბიობანკს (n=288+) — უნიკალური ეროვნული სამეცნიერო აქტივი — შექმნის")
add_bullet(doc, "ქართველ მკვლევართ Horizon-დონის პროექტის მართვის გამოცდილებას მისცემს")

add_separator(doc)

# ─── 2.3 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "2.3 გავრცელება, ექსპლუატაცია და კომუნიკაცია", 2)

add_heading(doc, "2.3.1 ღია მეცნიერება და გავრცელების გეგმა", 3)

add_body(doc, "**პუბლიკაციები:** ყველა Gold OA ან Green OA. სამიზნე ჟურნალები:")
add_bullet(doc, "WP1 შედეგები → *Aging Cell* ან *Journal of Cell Biology*")
add_bullet(doc, "WP2 შედეგები → *npj Aging* (Nature Portfolio)")
add_bullet(doc, "WP3 Cell-DT → *PLOS Computational Biology* ან *Bioinformatics*")
add_bullet(doc, "WP4 therapeutic PoC → *Nature Aging*")

add_body(doc, "**Preprint-ები:** bioRxiv depozit-ები M1-ზე (CDATA თეორია), M12 (WP1) და M24 (WP2).")
add_body(doc, "**Cell-DT კოდი:** ღია კოდი GitHub-ზე (djabbat/CDATA-Longevity), MIT ლიცენზია. Zenodo DOI თითოეული ვერსიისთვის.")

add_heading(doc, "2.3.2 კომუნიკაცია ფართო საზოგადოებასთან", 3)

add_bullet(doc, "**drjaba.com** (არსებული პლატფორმა): რეგულარული განახლებები 4 ენაზე (EN/KA/RU/KZ)")
add_bullet(doc, "**სოციალური მედია:** სიგრძის კომუნიტეტისთვის კომუნიკაციის კამპანია")
add_bullet(doc, "**წიგნი:** *Medicine of Generations* (ჯაბა თქემალაძე) — EN/KA/KZ-ში თარგმნილი — პოპულარული სამეცნიერო კომუნიკაციის საშუალება")

add_heading(doc, "2.3.3 ექსპლუატაცია და IP", 3)

add_body(doc, "**მონაცემთა მართვის გეგმა (DMP)** M3-ზე, განახლება M18-ზე. ძირითადი IP:")
add_bullet(doc, "CAII assay პროტოკოლი: პატენტის განაცხადი M18-ზე")
add_bullet(doc, "CentrosomeTransplant მეთოდი: provisional patent M24-ზე")
add_bullet(doc, "Cell-DT: MIT ლიცენზია; კომერციული dual-licence ფარმაცევტული კომპანიებისთვის")

add_body(doc, "**ექსპლუატაციის გზამკვლევი:** M1–M18: ვალიდაცია, IP განსაზღვრა → M18–M30: სამრეწველო პარტნიორების ჩართვა → M30–M36: სპინ-ოფ კომპანიის სკოპინგი.")
add_body(doc,
    "**CAII კლინიკური ტრანსლაციის გზა:** CAII ELISA assay WP2-ზე ვალიდირდება ადამიანის სისხლის ნიმუშებზე. M30–M36-ში CAII ELISA-ს თავსებადობა ავტომატური კლინიკური ქიმიის ანალიზატორთან (სამიზნე: Roche Cobas) შეფასდება სადიაგნოსტიკო სამრეწველო პარტნიორთან თანამშრომლობით."
)

add_heading(doc, "2.3.4 გავლენის შეჯამება", 3)

add_table_from_data(doc,
    ["გავლენის კატეგორია", "სამიზნე", "ინდიკატორი"],
    [
        ["სამეცნიერო გავლენა", "4+ Q1 პუბლიკაცია", "ჟურნალის მიღება"],
        ["ახალი დიაგნოსტიკა", "CAII სისხლის ანალიზი ვალიდირებული", "WP2 პირველადი"],
        ["ახალი თერაპიის კონცეფცია", "CentrosomeTransplant PoC", "WP4 in vitro"],
        ["ციფრული პლატფორმა", "Cell-DT v2.0 (ღია კოდი)", "GitHub + Zenodo"],
        ["ქართული სიმძლავრე", "ბიობანკი n=288, GTU STED", "WP2 ჩარიცხვა; მოწყობილობა"],
        ["IP", "1–2 პატენტის განაცხადი", "შეტანის ქვითარი"],
        ["ტრანსლაცია", "LoI სამრეწველოსთან", "ხელმოწერილი დოკუმენტი"],
    ]
)

add_separator(doc)

# ─── SECTION 3 ────────────────────────────────────────────────────────────────

add_heading(doc, "3. განხორციელების ხარისხი და ეფექტურობა", 1)

add_heading(doc, "3.1 სამუშაო გეგმა და რესურსები", 2)

add_heading(doc, "ცხრილი 3.1a: სამუშაო პაკეტების სია", 4)

add_table_from_data(doc,
    ["WP", "სათაური", "ლიდერი", "დაწყება", "დასრულება", "ბიუჯეტი (პირდაპირი)"],
    [
        ["WP1", "In Vitro ცენტრიოლური დაზიანების ვალიდაცია", "GTU", "M1", "M24", "€800,000"],
        ["WP2", "ადამიანის გრძელვადიანი კოჰორტის კვლევა", "GTU", "M3", "M36", "€450,000"],
        ["WP3", "Cell Digital Twin ინტეგრაცია და ვალიდაცია", "ფაზისის აკადემია", "M6", "M36", "€400,000"],
        ["WP4", "მკურნალობის PoC და ტრანსლაცია", "ფაზისის აკადემია", "M6", "M36", "€350,000"],
        ["**სულ პირდაპირი**", "", "", "", "", "**€2,000,000**"],
        ["**ირიბი (25%)**", "", "", "", "", "**€500,000**"],
        ["**სულ**", "", "", "", "", "**€2,500,000**"],
    ]
)

add_heading(doc, "ცხრილი 3.1b: სამუშაო პაკეტების აღწერა", 4)

add_heading(doc, "WP1 — In Vitro ცენტრიოლური დაზიანების ვალიდაცია (GTU, M1–M24, €800K)", 4)

add_body(doc, "*მიზანი:* CDATA-ს მექანიკური ვალიდობის დამყარება ადამიანის უჯრედულ მოდელებში.")

add_body(doc, "*ძირითადი ამოცანები:*")
add_bullet(doc, "T1.1 (M1–M6): უჯრედული მოდელების პანელის შედგენა (HeLa, IMR90, CD34⁺ EPC); CAII და U-ExM SOP")
add_bullet(doc, "T1.2 (M3–M12): PTM ინდუქციის დოზა-პასუხის ექსპერიმენტები (n=6 რეპლიკატი)")
add_bullet(doc, "T1.3 (M6–M18): STED სუპერ-რეზოლუციური ვიზუალიზაცია CEP164/CEP89/Ninein/CEP170")
add_bullet(doc, "T1.4 (M12–M24): EPC იზოლაცია, CAII გაზომვა ახალგაზრდა vs ასაკოვანი დონორები (n=20/ჯგუფი)")
add_bullet(doc, "T1.5 (M18–M24): CEP164 ELISA assay შემუშავება და ვალიდაცია")

add_body(doc, "*ძირითადი მილსადინარები:*")
add_bullet(doc, "D1.1 (M6): უჯრედული მოდელის პანელი + SOP-ები")
add_bullet(doc, "D1.2 (M12): PTM დოზა-პასუხის მონაცემები (Zenodo)")
add_bullet(doc, "D1.3 (M18): STED გამოსახულების მონაცემები + სტატია")
add_bullet(doc, "D1.4 (M24): CEP164 ELISA ვალიდირებული + WP1 საბოლოო ანგარიში")

add_body(doc, "*მოწყობილობა:* Leica SP8 STED (€120K, ახალი); Miltenyi MACS EasySep (€15K); კულტურის ლაბ. (€180K); რეაგენტები (€120K); კონფოკალი (€50K).")

add_separator(doc)

add_heading(doc, "WP2 — ადამიანის გრძელვადიანი კოჰორტის კვლევა (GTU, M3–M36, €450K)", 4)

add_body(doc, "*მიზანი:* CAII ინდექსის კორელაცია კლინიკური დაბერების ფენოტიპებთან (n=288).")

add_body(doc, "*ძირითადი ამოცანები:*")
add_bullet(doc, "T2.1 (M3–M9): ეთიკის დამტკიცება; მონაწილეების ჩარიცხვის გაშვება (12 პოლიკლინიკა)")
add_bullet(doc, "T2.2 (M6–M18): საბაზისო შეფასება (n=288): CAII, CEP164, HRV, ეპიგენეტური ასაკი, ტელომერი, კლინ. ფენოტიპი. **ყველა ბიომარკერის ანალიზი ბრმა კლინიკური მონაცემებისგან.**")
add_bullet(doc, "T2.3 (M18–M36): 2-წლიანი follow-up: ყველა ბიომარკერის განმეორება; სისუსტე, ჰოსპიტალიზაცია, სიკვდილი")
add_bullet(doc, "T2.4 (M24–M36): სტატისტიკური ანალიზი: ლოგისტიკური რეგრესია, კორელაცია, survival. **ბიოსტატისტიკოსი ბრმა კლინიკური შედეგებისგან.**")
add_bullet(doc, "T2.5 (M30–M36): ბიობანკის ფორმალიზაცია; GDPR; გადაცემა ეროვნულ ბიობანკში")

add_body(doc, "*პერსონალი:* 2 კლინიკური კვლევის ექთანი (FT, M3–M36); 1 ბიოსტატისტიკოსი (0.5 FTE, M18–M36); 1 ლაბ. ტექნიკოსი (FT, M3–M24).")

add_separator(doc)

add_heading(doc, "WP3 — Cell Digital Twin ინტეგრაცია (ფაზისის აკადემია, M6–M36, €400K)", 4)

add_body(doc, "*მიზანი:* WP1/WP2 ემპირიული მონაცემების Cell-DT-ში ჩართვა; Cell-DT v2.0 ღია კოდით გამოქვეყნება.")

add_body(doc, "*ძირითადი ამოცანები:*")
add_bullet(doc, "T3.1 (M6–M18): Bayesian პარამეტრების შეფასება WP1 მონაცემებით. Rust Engineer-ი M1-ზე ღიად იქნება (საერთაშორისო რეკრუტირება).")
add_bullet(doc, "T3.2 (M12–M24): pBA მოდული: CAII + ეპ. ასაკი → DamageParams (Stan posterior)")
add_bullet(doc, "T3.3 (M18–M30): pBA ვალიდაცია WP2 შედეგებზე (AUROC > 0.75)")
add_bullet(doc, "T3.4 (M24–M36): asymmetric_cytoqc_module, IGF-1 ღერძი, PTM CSV ექსპორტი")
add_bullet(doc, "T3.5 (M28–M36): Cell-DT v2.0 გამოქვეყნება (GitHub + Zenodo); PyO3 Python binding-ები")

add_body(doc, "*პროგრამული მოთხოვნები:* Rust 2021, hecs, Rayon, PyO3; Stan/PyMC; ყველა ცვლილება cargo test-ს უნდა გაივლიდეს (≥439 ტესტი + ახალი).")

add_separator(doc)

add_heading(doc, "WP4 — მკურნალობის PoC და ტრანსლაცია (ფაზისის აკადემია, M6–M36, €350K)", 4)

add_body(doc, "*მიზანი:* Cell-DT-ის პროგნოზირებული CDATA-სპეციფიკური ჩარევების ტესტირება in vitro; პატენტები; სამრეწველო LoI.")

add_body(doc, "*ძირითადი ამოცანები:*")
add_bullet(doc, "T4.0 (M6–M18): მოსამზადებელი ფაზა — უჯრედული ხაზების მომზადება, დოზის შეფასება, SOP")
add_bullet(doc, "T4.1 (M1): **კონსორციუმის შეთანხმება ხელმოწერილი** (ეტაპი M1)")
add_bullet(doc, "T4.2 (M18–M30): Top-3 ჩარევის in vitro ტესტირება (CentrosomeTransplant, CafdRetainer, Senolytics) IMR90 + CD34⁺ EPC; CAII, SA-β-Gal, EdU, ცია სიგრძე")
add_bullet(doc, "T4.3 (M24–M30): Cell-DT vs ექსპ. შედარება (Pearson r); ჩარევის პარამეტრების in silico დახვეწა")
add_bullet(doc, "T4.4 (M24–M30): პატენტის განაცხადი (CAII assay + CEP164 ELISA); provisional + PCT")
add_bullet(doc, "T4.5 (M30–M36): სამრეწველო ჩართვა; EIC Transition კონცეფტ-ნოტი")
add_bullet(doc, "T4.6 (M36): საბოლოო გავრცელების კონფერენცია (თბილისი; EU სპიკერები); FORGE-AI ქსელი")

add_heading(doc, "ცხრილი 3.1c: მილსადინარების სია", 4)

add_table_from_data(doc,
    ["D#", "სათაური", "WP", "ტიპი", "ვადა"],
    [
        ["D1.1", "უჯრედული მოდელი + SOP-ები", "WP1", "ანგარიში", "M6"],
        ["D1.2", "PTM დოზა-პასუხის მონაცემები", "WP1", "ღია მონაცემი", "M12"],
        ["D1.3", "STED მონაცემები + სტატია", "WP1", "პუბლიკაცია", "M18"],
        ["D1.4", "CEP164 ELISA + WP1 ანგარიში", "WP1", "ანგარიში", "M24"],
        ["D2.1", "ეთიკის დამტკიცება + ჩარიცხვა", "WP2", "ანგარიში", "M9"],
        ["D2.2", "საბაზისო მონაცემები (EHDS)", "WP2", "მონაცემთა ნაკრები", "M18"],
        ["D2.3", "2-წლ. follow-up + სტატია", "WP2", "პუბლიკაცია", "M36"],
        ["D3.1", "Cell-DT v1.5 ექსპ. პარამეტრები", "WP3", "პროგრ. უზრ.", "M18"],
        ["D3.2", "pBA მოდული ვალიდირებული + სტატია", "WP3", "პუბლიკაცია", "M24"],
        ["D3.3", "Cell-DT v2.0 საჯარო გამოქვეყნება", "WP3", "პროგრ. უზრ.", "M36"],
        ["D4.1", "კონსორციუმის შეთანხმება", "WP4", "ადმინ.", "**M1**"],
        ["D4.2", "WP4 PoC ანგარიში + პატენტი", "WP4", "ანგარიში/IP", "M30"],
        ["D4.3", "სამრეწველო LoI", "WP4", "ადმინ.", "M30"],
        ["D4.4", "საბოლოო ანგარიში + EIC Transition", "WP4", "ანგარიში", "M36"],
    ]
)

add_heading(doc, "ცხრილი 3.1d: ეტაპების სია", 4)

add_table_from_data(doc,
    ["M#", "ეტაპი", "WP", "ვადა", "დადასტურება"],
    [
        ["**M1**", "**კონსორციუმის შეთანხმება ხელმოწერილი**", "WP4", "**M1**", "ხელმოწ. დოკ."],
        ["M2", "ეთიკის დამტკიცება (WP2)", "WP2", "M9", "IRB წერილი"],
        ["M3", "უჯრ. მოდელი + CAII SOP ვალიდირებული", "WP1", "M6", "QC ანგარიში"],
        ["M4", "PTM დოზა-პასუხი დასრულებული", "WP1", "M12", "Zenodo ატვირთვა"],
        ["M5", "WP2 საბაზისო ჩარიცხვა (n=288)", "WP2", "M18", "CTMS ანგარიში"],
        ["M6", "Cell-DT ექსპ. პარამეტრები", "WP3", "M18", "პროგ. გამოქვ."],
        ["M7", "CAII–Frailty ასოციაცია (პირველ. ₌WP)", "WP2", "M30", "სტატ. ანგარიში"],
        ["M8", "პატენტის განაცხადი შეტანილი", "WP4", "M30", "შეტანის ქვითარი"],
        ["M9", "Cell-DT v2.0 ღია გამოქვ.", "WP3", "M36", "GitHub + Zenodo"],
    ]
)

add_heading(doc, "ცხრილი 3.1e: Gantt-ის დიაგრამა", 4)

gantt_para = doc.add_paragraph()
para_spacing(gantt_para)
gantt_run = gantt_para.add_run(
"""WP/Task         M01 M06 M09 M12 M18 M24 M30 M36
WP1-T1.1         ████████████
WP1-T1.2                  ████████████████████████
WP1-T1.3                        ████████████████████████████████
WP1-T1.4                                             ████████████████████████
WP1-T1.5                                                   ████████████████
WP2-T2.1               ████████████████████
WP2-T2.2                           ████████████████████████████████
WP2-T2.3                                                █████████████████████████████████
WP2-T2.4                                                               ████████████████████████
WP3-T3.1                     ████████████████████████████████████
WP3-T3.2                                    ████████████████████████████████████
WP3-T3.3                                                  ████████████████████████████████
WP3-T3.4                                                               ████████████████████████████
WP3-T3.5                                                                        ████████████████████
WP4-T4.0                     ████████████████████████████████████
WP4-T4.1         ██ (M1)
WP4-T4.2                                   ████████████████████████████████████████
WP4-T4.3                                                         ████████████████████████████
WP4-T4.4                                                         ████████████████████████
WP4-T4.5                                                                        ████████████████████
WP4-T4.6                                                                                  ██████████
─────────────────────────────────────────────────────────────────
ეტაპები          ▼M1       ▼M3  ▼M2  ▼M4    ▼M5,M6          ▼M7  ▼M8  ▼M9
"""
)
gantt_run.font.name = "Courier New"
gantt_run.font.size = Pt(7)

add_heading(doc, "ცხრილი 3.1f: პერსონალის ძალისხმევის შეჯამება (პიროვნება-თვე)", 4)

add_table_from_data(doc,
    ["როლი", "ინსტიტუტი", "WP1", "WP2", "WP3", "WP4", "სულ PM"],
    [
        ["PI (თქემალაძე, MD)", "ფაზისის აკადემია", "6", "6", "18", "12", "**42**"],
        ["co-PI (პროფ. ბატონი გელა)", "GTU", "12", "12", "6", "4", "**34**"],
        ["Postdoc — უჯრ. ბიოლ.", "GTU", "18", "6", "6", "6", "**36**"],
        ["Postdoc — ბიოინფ.", "ფაზისის აკადემია", "—", "6", "24", "6", "**36**"],
        ["PhD 1", "GTU", "24", "12", "—", "6", "**42**"],
        ["PhD 2", "GTU", "—", "24", "—", "12", "**36**"],
        ["კლინ. ექთანი ×2", "GTU", "—", "72", "—", "—", "**72**"],
        ["ლაბ. ტექნიკოსი", "GTU", "24", "12", "—", "—", "**36**"],
        ["ბიოსტ. (0.5)", "GTU", "—", "18", "6", "3", "**27**"],
        ["Rust Engineer", "ფაზისის აკადემია", "—", "—", "30", "6", "**36**"],
        ["**სულ**", "", "**84**", "**168**", "**90**", "**55**", "**397**"],
    ]
)

add_separator(doc)

# ─── 3.2 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "3.2 მონაწილეთა და კონსორციუმის შესაძლებლობები", 2)

add_heading(doc, "3.2.1 ფაზისის აკადემია (მთავარი ბენეფიციარი)", 3)

add_body(doc,
    "**სამართლებრივი სტატუსი:** არაკომერციული სამეცნიერო-კვლევითი ორგანიზაცია, საქართველოში რეგისტრირებული (1974, ფოთი). ფაზისის აკადემია მოქმედებს საქართველოს კანონმდებლობის საფუძველზე არასამეწარმეო (არაკომერციული) სამართლებრივ პირებზე და უფლებამოსილია სამეცნიერო-კვლევითი საქმიანობისა და საერთაშორისო გრანტის მიღებისთვის."
)
add_body(doc,
    "**კვლევის ისტორია:** ბუნებისმეტყველება, ეკოლოგია, თეორიული ბიოლოგია 50+ წელი. კომპიუტაციული დაბერების ბიოლოგიის კვლევის მიმართულება (CDATA პროგრამა, 2022 წლიდან) ფორმალურად კონსტიტუირებულია."
)
add_body(doc,
    "**ადმინისტრაციული სიმძლავრე:** EU სტანდარტული ანგარიშგება (SAP-თავსებადი ბუღალტერია; ყოველწლიური აუდიტი). ადმინისტრაციული პერსონალის დამტკიცების წერილი მოთხოვნისამებრ ხელმისაწვდომია."
)

add_body(doc, "**დრ. ჯაბა თქემალაძე, MD — მთავარი გამომძიებელი**")
add_bullet(doc, "20+ წლის კლინიკური და სამეცნიერო გამოცდილება ინტეგრაციულ მედიცინასა და უჯრედულ გერონტოლოგიაში")
add_bullet(doc, "**გამოქვეყნებული:** 2 peer-reviewed პუბლიკაცია CDATA-ზე (PMID 36583780; DOI 10.65649/yx9sn772)")
add_bullet(doc, "**პროგრ. უზრ.:** Cell-DT-ის ავტორი (14-მოდულიანი Rust პლატფ.; ≥439 ტესტი; djabbat/CDATA-Longevity GitHub-ზე)")
add_bullet(doc, "**მრავალენობრივი:** პუბლ. და კლინ. მოვლა — RU, KA, EN, KZ")
add_bullet(doc, "**ინფრასტრუქტურა:** HPC (GTU); drjaba.com AI-ასისტირებული კლინ. პლატფ. (4 ენა)")

add_heading(doc, "3.2.2 საქართველოს ტექნიკური უნივერსიტეტი — GTU ABET BME (ბენეფიციარი 2)", 3)

add_body(doc, "**PIC:** 983636358 | **ფაკულტეტი:** ABET; ბიომედიცინური ინჟინერიის დეპარტამენტი")

add_body(doc, "**ძირითადი რესურსები:**")
add_bullet(doc, "Zeiss LSM 700 კონფოკ. მიკრ. (სამოქმედო, WP1 სტანდ. კონფ.)")
add_bullet(doc, "ახ. Leica SP8 STED მიკრ. WP1 ბიუჯეტში (€120K)")
add_bullet(doc, "უჯრ. კულტურის ობ. (BSL-2; 3 ლამინ. ბოქსი; ×2 CO₂ ინკ.)")
add_bullet(doc, "ნაკადური ციტომ. (BD FACSCanto II)")
add_bullet(doc, "qPCR (Applied Biosystems 7900HT)")
add_bullet(doc, "ბიობანკი: -80°C (×4 ULT საყინულო)")

add_body(doc,
    "**Horizon-ის გამოცდილება:** FORGE-AI (HORIZON-MSCA-2024-SE-01, 2025–2028, GA #101216703) — GTU-ს EU-დაფინ. პროექ. მართვის, საერთ. ანგარიშ. და €2M+ ბიუჯ. მართვის სიმძლავრე დადასტურებულია."
)

add_body(doc, "**co-PI — პროფ. ბატონი გელა, PhD — GTU ABET BME**")
add_bullet(doc, "ბიომედ. ინჟინ. პროფ., GTU ABET")
add_bullet(doc, "FORGE-AI Horizon Europe კონსორციუმის წევრი (GA #101216703); EU პროექ. მართვისა და ანგ. გამოც.")
add_bullet(doc, "**CDATA-ში როლი:** WP1/WP2 ზედამხედველობა; Leica SP8 STED ხარ. კონტ.; GTU მილს. თანა-ხელმ.")

add_heading(doc, "3.2.3 კონსორციუმის დასაბუთება", 3)

add_body(doc,
    "ორ-ინსტიტუციური კონსორციუმი ასახავს EIC Pathfinder-ისთვის შესაფერის TRL 1–4 სამეც. ფოკუსს. **ფაზისის აკადემია:** CDATA-ს ტეორია, Cell-DT კომპ. შემ., WP3/WP4. **GTU:** ექსპ. ინფრასტრ., კლინ. კოჰ. განხ., STED მიკრ., EU პრ. გამ. კონსორციუმი განზრახ მცირეა — EIC Pathfinder-ი დიდ სიმამაცეს ჰიპოთეზ.-ორიენტ. კვლ.-ს მცირე გუნდებში ასახელებს."
)

add_separator(doc)

# ─── 3.3 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "3.3 ბიუჯეტის დაშლა", 2)

add_heading(doc, "ცხრილი 3.3a: ბიუჯეტი WP-ისა და ინსტიტუტის მიხედვით", 4)

add_table_from_data(doc,
    ["ბიუჯეტის მუხლი", "WP1 (GTU)", "WP2 (GTU)", "WP3 (ფაზ.)", "WP4 (ფაზ.)", "სულ"],
    [
        ["**პერსონალი**", "€420,000", "€280,000", "€220,000", "€160,000", "€1,080,000"],
        ["**მოწყობილობა**", "€200,000*", "€20,000", "€40,000", "€10,000", "€270,000"],
        ["**სახ./რეაგ.**", "€140,000", "€60,000", "€30,000", "€80,000", "€310,000"],
        ["**მოგ./გავ.**", "€20,000", "€30,000", "€40,000", "€40,000", "€130,000"],
        ["**ქვეკ.**", "€20,000", "€60,000", "€70,000", "€60,000", "€210,000"],
        ["**პირდ. სულ**", "**€800,000**", "**€450,000**", "**€400,000**", "**€350,000**", "**€2,000,000**"],
        ["**ირ. (25%)**", "€200,000", "€112,500", "€100,000", "€87,500", "**€500,000**"],
        ["**ჯ. სულ**", "**€1,000,000**", "**€562,500**", "**€500,000**", "**€437,500**", "**€2,500,000**"],
    ]
)

add_body(doc, "*WP1 მოწ.: Leica SP8 STED €120K + უჯრ. კულტ. €50K + ELISA reader €30K")

add_heading(doc, "ცხრილი 3.3b: ძირითადი მოწყობილობის ღირებულება", 4)

add_table_from_data(doc,
    ["ნივთი", "WP", "ღ-ბა (€)", "დასაბუთება"],
    [
        ["Leica SP8 STED მიკროსკ. — ახ. შეძ. (GTU)", "WP1", "120,000",
         "ნანო-მასშტ. ცენტრ. დამ. ვიზ.; U-ExM < 100 nm. Zeiss LSM 700 სტანდ. კ.; Leica SP8 STED ცალ. ახ. ინსტ."],
        ["Miltenyi MACS EasySep (EPC იზ.)", "WP1", "15,000", "CD34⁺ EPC 20 მლ სისხლიდან"],
        ["ELISA reader + washer (Synergy H1)", "WP1", "28,000", "CEP164 ELISA assay"],
        ["qPCR რეაგ. (ტელ. სიგ., 288 × 2)", "WP2", "35,000", "WP2 კოჰ. ტელ. სიგ."],
        ["ECG Holter × 20", "WP2", "18,000", "5-წთ. HRV მეორ. ₌WP"],
        ["HPC cloud (AWS/Azure)", "WP3", "24,000", "Bayesian MCMC; pBA n=288"],
        ["CFS (ფინ. ანგ. სერტ.)", "ადმ.", "7,000", "Horizon სტ. მოთხ."],
    ]
)

add_body(doc, "**ქვეკ. WP2:** CRO სერვ. დამოუკ. მონ. (€60,000) — ICH E6(R2) §5.18 GCP.")

add_separator(doc)

# ─── 3.4 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "3.4 რისკების მართვა", 2)

add_table_from_data(doc,
    ["#", "რისკი", "ალბ.", "გავლ.", "შემარბილება"],
    [
        ["R1", "GTU co-PI ჩართვა", "დაბ.", "დაბ.",
         "co-PI დადასტ.: პროფ. ბატონი გელა (GTU ABET BME, FORGE-AI). CV + სადამ. წ. M1-მდე."],
        ["R2", "CAII სისუსტეს არ ასოც. (პირვ. ₌WP ვერ)", "დაბ.-საშ.", "მაღ.",
         "WP2-ს 5 მეორ. ₌WP; უარყ. CAII-ც პუბლ. ინფ.; Cell-DT WP3 WP2-ს დამოუკ."],
        ["R3", "Leica SP8 STED მიწ. შეფ.", "დაბ.", "საშ.",
         "M1–M9: Zeiss LSM 700; STED M10-დან; 3-თვ. დაყ. მისაღ.; ალტ.: TSMU კ."],
        ["R4", "ფაზ. PIC M3-მდე არ", "საშ.", "მაღ.",
         "PIC M1-მდე; EU Funding Portal 2–4 კვ.; GTU PIC (983636358) ადმ. კოორდ."],
        ["R5", "CentrosomeTransplant ტექნ. განუხ.", "საშ.", "საშ.",
         "3 ალტ. ჩარ. (CafdRetainer, Senolytics, Rapa); IP ვალ. ჩარ.-ზე გადა."],
    ]
)

add_separator(doc)

# ─── 3.5 ──────────────────────────────────────────────────────────────────────

add_heading(doc, "3.5 ღია მეცნიერების პრაქტიკა", 2)

add_body(doc, "CDATA სრულ ღია მეცნიერების შესაბამისობას ინარჩუნებს Horizon Europe-ს მოთხოვნებთან:")
add_bullet(doc, "**ღია წვდომა:** Gold OA (APC ბიუჯ.: €2,000 × 4 = €8,000) ან Green OA ≤ 6 თვე.")
add_bullet(doc, "**FAIR მონ.:** WP1: Zenodo (CC-BY 4.0); WP2: ანონ., EHDS; Cell-DT: GitHub (MIT) + Zenodo DOI")
add_bullet(doc, "**პასუხისმგებელი კვლ.:** ცხოვ. ექსპ. არ არის; WP2: ჰელსინკი; GTU IRB M3-მდე")
add_bullet(doc, "**EOSC:** Cell-DT Zenodo-ში; WP2 ანონ. B2FIND-ში")

add_separator(doc)

# ─── SECTION 4 ────────────────────────────────────────────────────────────────

add_heading(doc, "4. ეთიკის თვითშეფასება", 1)

add_heading(doc, "4.1 მიზნების, მეთოდოლოგიის და სავარ. გავლენის ეთიკური განზომილებები", 2)

add_heading(doc, "ადამიანი მონაწილეები (WP2):", 4)
add_bullet(doc, "სრულიად ნებაყოფლობითი; ორ-ეტ. თანხ.: Mini-Cog → წერ. ინფ. თანხ.")
add_bullet(doc, "გამოთიშვის უფლება ნებისმიერ დროს")
add_bullet(doc, "სამკ. ჩარ. არ; წმ. სადაკვ. — სისხლ. + კლინ. შეფ.")
add_bullet(doc, "სპეც. დაცვა: 66–80 ასაკ. ჯგ. — Mini-Cog უზ. კოგ. სიმ.")
add_bullet(doc, "მონ.: ფსევდ. კოლ. პ.-ზე; GTU DPO-ს ც.-ზე ცალ. ც.")

add_heading(doc, "ბიოლ. ნიმ. (WP2 ბიობ.):", 4)
add_bullet(doc, "თანხ. სამ. კვლ. მოიცავს")
add_bullet(doc, "MTA ნ.-ს. მ.-ს გაზ.")
add_bullet(doc, "გენ. ტ. — მხოლ. ხელ. განახ. თანხ.")

add_heading(doc, "Cell-DT პ.-ო. სასაქ.:", 4)
add_bullet(doc, "Dual-use რ.: დაბ. დ.-ბ. მოდ.; სამხ./იარ. გამ. — MIT ლიც. med. device-ის გამ. განც.")

add_heading(doc, "კონფიდ.:", 4)
add_bullet(doc, "WP1: პირ. ინდ. მ. არ; WP2: ქართ. PDPLaw + GDPR პრ.")

add_heading(doc, "4.2 ეთ. პრ. და კანონ. სისტ.-სთ. შ.", 2)

add_table_from_data(doc,
    ["ეთ. მოთ.", "შ. ზომა"],
    [
        ["ინფ. თანხ. (WP2)", "ორ-ეტ. Mini-Cog + წ. თანხ.; ICH E6(R2) GCP"],
        ["მ. დაც.", "ქართ. PDPLaw + GDPR; ფსევდ."],
        ["ჰელსინკი", "GTU IRB M3-მდე"],
        ["ცხ. ექსპ. არ", "დადასტ.; მხ. უჯრ. ხ. + ნებ. მ."],
        ["სარგ. გავ.", "ღ. წ. პუბ.; EHDS მ. წვ."],
        ["Dual-use", "Cell-DT: ბ. მ. მხ.; d.-u. რ. არ"],
    ]
)

add_separator(doc)

# ─── REFERENCES ───────────────────────────────────────────────────────────────

add_heading(doc, "დანართი: ძირითადი მიმართვები", 2)

refs = [
    "1. López-Otín C et al. \"Hallmarks of aging: An expanding universe.\" *Cell.* 2023;186(2):243-278.",
    "2. Blackburn EH. \"Telomere states and cell fates.\" *Nature.* 2000;408:53-56.",
    "3. Tkemaladze J, Lortkipanidze G. \"Asymmetric segregation of the centrosome and its role in determining cell fate.\" *Mol Biol Rep.* 2023;50(3):2741-2748. **PMID: 36583780**",
    "4. Tkemaladze J. \"CDATA: Centriolar Damage Accumulation Theory of Ageing.\" *Ann Rejuvenation Sci.* 2025. **DOI: 10.65649/yx9sn772**",
    "5. Gambarotto D et al. \"Imaging cellular ultrastructures using expansion microscopy (U-ExM).\" *Nat Methods.* 2019;16:71-74.",
    "6. Franceschi C et al. \"Inflammaging and anti-inflammaging: A systemic perspective on aging and longevity.\" *Mech Ageing Dev.* 2007;128(1):92-105.",
    "7. Vasa M et al. \"Increase in circulating endothelial progenitor cells by statin therapy.\" *Circulation.* 2001;103(24):2885-90.",
    "8. Levine ME et al. \"An epigenetic biomarker of aging for lifespan and healthspan.\" *Aging.* 2018;10(4):573-591.",
    "9. Cawthon RM. \"Telomere measurement by quantitative PCR.\" *Nucleic Acids Res.* 2002;30(10):e47.",
    "10. Fried LP et al. \"Frailty in older adults: evidence for a phenotype.\" *J Gerontol A Biol Sci Med Sci.* 2001;56(3):M146-156.",
    "11. Gladyshev VN. \"The free radical theory of aging is dead. Long live the damage theory!\" *Antioxid Redox Signal.* 2014;20(4):727-731.",
]
for ref in refs:
    p = doc.add_paragraph()
    para_spacing(p, after=3)
    _add_inline_text(p, ref)

add_separator(doc)

# Footer note
footer_note = doc.add_paragraph()
para_spacing(footer_note, before=8)
_add_inline_text(footer_note,
    "*დოკუმენტის ვერსია: v7.0 — 2026-03-24 | PI: ჯაბა თქემალაძე, MD — ფაზისის აკადემია, ფოთი, საქართველო | jaba@drjaba.com*"
)

# ─── SAVE ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")

import os
size = os.path.getsize(OUTPUT_PATH)
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
if size > 50000:
    print("OK: file is >50KB")
else:
    print("WARNING: file is <50KB!")
