#!/usr/bin/env python3
"""
Translate EIC_Pathfinder_CDATA_PartB_v7.md → Russian → .docx
"""

import re
import sys
import time
import requests
from pathlib import Path

# ── 1. Read source ──────────────────────────────────────────────────────────
SRC = Path("/home/oem/Desktop/Horizon/EIC_Pathfinder_CDATA_PartB_v7.md")
OUT = Path("/home/oem/Desktop/Horizon/EIC_Pathfinder_CDATA_PartB_v7_RU.docx")

text = SRC.read_text(encoding="utf-8")
lines = text.splitlines()
print(f"Source: {len(lines)} lines")

# ── 2. Terms that must NOT be translated ────────────────────────────────────
PROTECTED = [
    "CDATA", "CAII", "Cell-DT", "Ze Theory", "WP1", "WP2", "WP3", "WP4",
    "PTM", "STED", "HDF", "HeLa", "KPI", "TRL", "SASP", "IRB", "GDPR",
    "GCP", "ECS", "PyO3", "Rayon", "hecs", "ELISA", "FORGE-AI", "GTU",
    "BME", "ABET", "EIC", "Horizon Europe", "IMR90", "CEP164", "CEP89",
    "CEP170", "Ninein", "U-ExM", "MACS", "PBMC", "HRV", "MMSE",
    "AUROC", "AUROC", "CHIP", "HSC", "EHDS", "FAIR", "EOSC", "MIT",
    "Phasis Academy", "Jaba Tkemaladze", "Batoni Gela",
    "EPC", "CentrosomeTransplant", "CafdRetainer", "Senolytics",
    "DDRState", "NKSurveillanceState", "ProteostasisState",
    "CircadianState", "AutophagyState", "pBA", "DamageParams",
    "CAFD", "SA-β-Gal", "EdU", "CTMS", "DMP", "CRO", "IVD",
    "MCMC", "Stan", "PyMC", "PyPI", "TSMU", "FORGE", "PCT",
]

# ── 3. DeepSeek API helper ───────────────────────────────────────────────────
DS_KEY = "sk-a759a615a20147b19849be295f096675"
DS_URL = "https://api.deepseek.com/v1/chat/completions"
DS_MODEL = "deepseek-chat"

SYSTEM_PROMPT = (
    "You are a scientific translator. Refine this Russian translation of a Horizon Europe "
    "grant application. Keep technical terms in English (CDATA, CAII, Cell-DT, WP1-4, PTM, "
    "STED, KPI, TRL, ECS, SASP, GTU, ABET, FORGE-AI, ELISA, GCP, GDPR, IRB, "
    "CentrosomeTransplant, CafdRetainer, Ze Theory, EPC, HDF, IMR90, HeLa, U-ExM, "
    "CEP164, CEP89, CEP170, Ninein, MACS, PBMC, HRV, MMSE, AUROC, CHIP, HSC, EHDS, "
    "FAIR, EOSC, MIT, PyO3, Rayon, hecs, pBA, DamageParams, SA-β-Gal, EdU). "
    "Keep all names unchanged: Jaba Tkemaladze, Batoni Gela, Phasis Academy. "
    "Keep all €, %, numbers, M1–M36, n=288 unchanged. "
    "Keep all citation markers like [1]–[11] unchanged. "
    "Keep all journal names in English. "
    "Keep Markdown formatting (##, **, *, >, -, |) intact. "
    "Return ONLY the refined Russian text, no commentary, no preamble."
)

def deepseek_refine(draft: str, retries: int = 3) -> str:
    """Call DeepSeek API to refine a draft Russian translation."""
    user_msg = f"Draft translation:\n{draft}"
    payload = {
        "model": DS_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg},
        ],
        "temperature": 0.3,
        "max_tokens": 4096,
    }
    headers = {"Authorization": f"Bearer {DS_KEY}", "Content-Type": "application/json"}
    for attempt in range(retries):
        try:
            resp = requests.post(DS_URL, json=payload, headers=headers, timeout=120)
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"].strip()
        except Exception as e:
            print(f"  DeepSeek attempt {attempt+1} failed: {e}")
            if attempt < retries - 1:
                time.sleep(5)
    print("  Falling back to Google draft for this chunk.")
    return draft

# ── 4. Google Translate helper ───────────────────────────────────────────────
from deep_translator import GoogleTranslator

def google_translate(chunk: str) -> str:
    """Translate a chunk to Russian using GoogleTranslator."""
    # Split into ≤4500-char pieces (Google limit ~5000)
    MAX = 4500
    if len(chunk) <= MAX:
        try:
            return GoogleTranslator(source="en", target="ru").translate(chunk)
        except Exception as e:
            print(f"  Google translate error: {e}")
            return chunk
    # Chunk is too long — split by lines
    parts = []
    buf = ""
    for ln in chunk.splitlines(keepends=True):
        if len(buf) + len(ln) > MAX:
            if buf:
                try:
                    parts.append(GoogleTranslator(source="en", target="ru").translate(buf))
                except Exception as e:
                    print(f"  Google translate partial error: {e}")
                    parts.append(buf)
                buf = ""
        buf += ln
    if buf:
        try:
            parts.append(GoogleTranslator(source="en", target="ru").translate(buf))
        except Exception as e:
            print(f"  Google translate final error: {e}")
            parts.append(buf)
    return "".join(parts)

# ── 5. Chunk source lines (≈50 lines each) ──────────────────────────────────
CHUNK_SIZE = 50
chunks = []
for i in range(0, len(lines), CHUNK_SIZE):
    chunks.append("\n".join(lines[i : i + CHUNK_SIZE]))

print(f"Chunks to translate: {len(chunks)}")

# ── 6. Translate each chunk ──────────────────────────────────────────────────
translated_chunks = []
for idx, chunk in enumerate(chunks):
    print(f"Chunk {idx+1}/{len(chunks)} ({len(chunk)} chars)...", end=" ", flush=True)
    # Google draft
    draft = google_translate(chunk)
    print("Google ✓", end=" ", flush=True)
    # DeepSeek refinement
    refined = deepseek_refine(draft)
    print("DeepSeek ✓")
    translated_chunks.append(refined)
    time.sleep(0.5)  # be gentle with APIs

full_ru = "\n".join(translated_chunks)
print(f"\nFull Russian text: {len(full_ru)} chars")

# ── 7. Build .docx ───────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default paragraph style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# ── Footer (Page X of Y) ────────────────────────────────────────────────────
def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        # "Страница " field
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)
        run2 = para.add_run(" из ")
        run2.font.size = Pt(9)
        run2.font.name = "Calibri"
        run3 = para.add_run()
        run3.font.size = Pt(9)
        run3.font.name = "Calibri"
        fld_begin2 = OxmlElement("w:fldChar")
        fld_begin2.set(qn("w:fldCharType"), "begin")
        run3._r.append(fld_begin2)
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = "NUMPAGES"
        run3._r.append(instr2)
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_end2)
        # Prefix
        run_prefix = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
        run_prefix.insert(0, rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = "Страница "
        run_prefix.append(t)
        para._p.insert(list(para._p).index(run._r), run_prefix)

add_footer(doc)

# ── First paragraph: translator note ────────────────────────────────────────
note = doc.add_paragraph()
note_run = note.add_run(
    "Рабочий перевод. Оригинал: EIC_Pathfinder_CDATA_PartB_v7.md"
)
note_run.italic = True
note_run.font.name = "Calibri"
note_run.font.size = Pt(10)

# ── Heading styles ──────────────────────────────────────────────────────────
def set_heading_style(para, level):
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    run = para.runs[0] if para.runs else para.add_run()
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 11))
    if level == 4:
        run.font.italic = True
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)

# ── Inline bold/italic parser ────────────────────────────────────────────────
def apply_inline(para, text: str):
    """Parse **bold** and *italic* and add runs to paragraph."""
    # Pattern: **text** | *text* | plain
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for m in pattern.finditer(text):
        # plain before
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.font.name = "Calibri"
            r.font.size = Pt(11)
        if m.group(0).startswith("**"):
            r = para.add_run(m.group(2))
            r.bold = True
        else:
            r = para.add_run(m.group(3))
            r.italic = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.name = "Calibri"
        r.font.size = Pt(11)

# ── Table builder ────────────────────────────────────────────────────────────
def add_table_from_md(doc, table_lines):
    """Parse markdown table lines and add a python-docx Table."""
    rows = []
    for ln in table_lines:
        if re.match(r"^\s*\|[-| :]+\|\s*$", ln):
            continue  # separator row
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # Normalise
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            # Strip markdown bold from cell
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            p = cell.paragraphs[0]
            run = p.add_run(clean)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if ri == 0:
                run.bold = True
                # Header row shading D9E1F2
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E1F2")
                tcPr.append(shd)

# ── Horizontal divider ───────────────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ── Main parser: iterate lines ────────────────────────────────────────────────
ru_lines = full_ru.splitlines()

i = 0
while i < len(ru_lines):
    ln = ru_lines[i]

    # --- Horizontal divider ---
    if re.match(r"^---+\s*$", ln):
        add_hr(doc)
        i += 1
        continue

    # --- Headings ---
    m = re.match(r"^(#{1,4})\s+(.*)", ln)
    if m:
        level = len(m.group(1))
        heading_text = m.group(2).strip()
        # Clean bold from heading text
        heading_text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
        heading_text = re.sub(r"\*(.+?)\*", r"\1", heading_text)
        para = doc.add_paragraph()
        run = para.add_run(heading_text)
        run.font.name = "Calibri"
        set_heading_style(para, level)
        i += 1
        continue

    # --- Table detection ---
    if ln.strip().startswith("|"):
        table_lines = []
        while i < len(ru_lines) and ru_lines[i].strip().startswith("|"):
            table_lines.append(ru_lines[i])
            i += 1
        add_table_from_md(doc, table_lines)
        continue

    # --- Code block ---
    if ln.strip().startswith("```"):
        i += 1
        code_lines = []
        while i < len(ru_lines) and not ru_lines[i].strip().startswith("```"):
            code_lines.append(ru_lines[i])
            i += 1
        i += 1  # closing ```
        para = doc.add_paragraph()
        run = para.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        continue

    # --- Blockquote ---
    if ln.strip().startswith(">"):
        qt = ln.strip().lstrip(">").strip()
        # Remove bold markers for quote
        qt = re.sub(r"\*\*(.+?)\*\*", r"\1", qt)
        qt = re.sub(r"\*(.+?)\*", r"\1", qt)
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.5)
        run = para.add_run(qt)
        run.italic = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        i += 1
        continue

    # --- Bullet list ---
    if re.match(r"^\s*[-*]\s+", ln):
        item_text = re.sub(r"^\s*[-*]\s+", "", ln)
        para = doc.add_paragraph(style="List Bullet")
        apply_inline(para, item_text)
        i += 1
        continue

    # --- Numbered list ---
    if re.match(r"^\s*\d+\.\s+", ln):
        item_text = re.sub(r"^\s*\d+\.\s+", "", ln)
        para = doc.add_paragraph(style="List Number")
        apply_inline(para, item_text)
        i += 1
        continue

    # --- Empty line ---
    if not ln.strip():
        i += 1
        continue

    # --- Normal paragraph ---
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    # spacing
    para.paragraph_format.line_spacing = Pt(11 * 1.15)
    apply_inline(para, ln)
    i += 1

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
size = OUT.stat().st_size
print(f"\nSaved: {OUT}")
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
if size >= 60 * 1024:
    print("✓ Exceeds 60 KB threshold")
else:
    print("⚠ WARNING: File is smaller than 60 KB — check content!")
