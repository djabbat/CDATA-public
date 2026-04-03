#!/usr/bin/env python3
"""
Translate EIC_Pathfinder_CDATA_PartB_v7.md to Georgian and save as .docx
"""

import re
import time
import requests
from pathlib import Path

from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE = "/home/oem/Desktop/Horizon/EIC_Pathfinder_CDATA_PartB_v7.md"
OUTPUT = "/home/oem/Desktop/Horizon/EIC_Pathfinder_CDATA_PartB_v7_KA.docx"

DEEPSEEK_KEY = "sk-a759a615a20147b19849be295f096675"
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"

# Terms to keep in English (protected via placeholder system)
KEEP_ENGLISH = [
    "CDATA", "CAII", "Cell-DT", "Ze Theory", "WP1", "WP2", "WP3", "WP4",
    "PTM", "STED", "HDF", "HeLa", "KPI", "TRL", "SASP", "IRB", "GDPR",
    "GCP", "ECS", "PyO3", "Rayon", "hecs", "ELISA", "FORGE-AI", "GTU",
    "BME", "ABET", "EIC", "Horizon Europe", "IMR90", "MACS", "U-ExM",
    "SA-β-Gal", "EdU", "ECG", "HRV", "AUROC", "CHIP", "CTMS", "EHDS",
    "EOSC", "FAIR", "DMP", "APC", "IVD", "CRO", "HPC", "MIT", "PyPI",
    "GitHub", "Zenodo", "bioRxiv", "PMID", "DOI", "PIC", "MSCA",
    "NKSurveillanceState", "ProteostasisState", "CircadianState",
    "AutophagyState", "DDRState", "DamageParams", "base_detach_probability",
    "senescence_threshold", "regeneration_tempo", "ciliary_function",
    "pool_exhaustion", "methylation_age", "stem_cell_pool", "mito_shield",
    "cargo test", "hecs", "async_cytoqc_module", "asymmetric_cytoqc_module",
    "IGF-1", "mTOR", "NAD", "CEP164", "CEP89", "CEP170", "Ninein",
    "γ-H2AX", "CD34", "CentrosomeTransplant", "CafdRetainer", "Senolytics",
    "CAFD", "CHIP", "pBA", "Stan", "PyMC", "Bayesian",
    "Aging Cell", "npj Aging", "Nature Aging", "PLOS Computational Biology",
    "Bioinformatics", "Mol Biol Rep", "Ann Rejuvenation Sci",
    "Cell", "Nature", "Nat Methods", "Nucleic Acids Res",
    "drjaba.com", "djabbat/CDATA-Longevity",
    "Phasis Academy", "Georgian Technical University",
    "FORGE-AI", "BD FACSCanto", "Applied Biosystems", "Roche Cobas",
    "Miltenyi", "Synergy H1", "AWS", "Azure",
]

# Specific Georgian name forms
GEORGIAN_NAMES = {
    "Jaba Tkemaladze": "ჯაბა თქემალაძე",
    "Dr. Jaba Tkemaladze": "დოქტ. ჯაბა თქემალაძე",
    "Batoni Gela": "ბატონი გელა",
    "Prof. Batoni Gela": "პროფ. ბატონი გელა",
    "Phasis Academy": "ფაზისის აკადემია",
}


def protect_terms(text):
    """Replace protected terms with numbered placeholders."""
    placeholders = {}
    protected = text

    # Protect citations like [1]-[11]
    citation_pattern = r'\[\d+\](?:-\[\d+\])?'
    for m in re.finditer(citation_pattern, protected):
        token = f"__CITE_{len(placeholders)}__"
        placeholders[token] = m.group(0)

    for k, v in placeholders.items():
        protected = protected.replace(v, k)

    # Protect €, %, numbers with units
    money_pattern = r'€[\d,]+(?:,\d{3})*(?:K|M)?'
    new_placeholders = {}
    for m in re.finditer(money_pattern, protected):
        token = f"__MONEY_{len(placeholders) + len(new_placeholders)}__"
        new_placeholders[token] = m.group(0)
    for k, v in new_placeholders.items():
        protected = protected.replace(v, k)
    placeholders.update(new_placeholders)

    # Protect n=NNN patterns
    n_pattern = r'n=[\d,]+'
    new_placeholders = {}
    for m in re.finditer(n_pattern, protected):
        token = f"__NPAT_{len(placeholders) + len(new_placeholders)}__"
        new_placeholders[token] = m.group(0)
    for k, v in new_placeholders.items():
        protected = protected.replace(v, k)
    placeholders.update(new_placeholders)

    # Protect M1-M36 month markers
    month_pattern = r'\bM\d{1,2}\b'
    new_placeholders = {}
    for m in re.finditer(month_pattern, protected):
        token = f"__MONTH_{len(placeholders) + len(new_placeholders)}__"
        new_placeholders[token] = m.group(0)
    for k, v in new_placeholders.items():
        protected = protected.replace(v, k)
    placeholders.update(new_placeholders)

    # Protect kept English terms (longer first to avoid partial matches)
    sorted_terms = sorted(KEEP_ENGLISH, key=len, reverse=True)
    for term in sorted_terms:
        if term in protected:
            token = f"__TERM_{len(placeholders)}__"
            # Only replace whole-word occurrences
            escaped = re.escape(term)
            new_text = re.sub(r'(?<!\w)' + escaped + r'(?!\w)', token, protected)
            if new_text != protected:
                placeholders[token] = term
                protected = new_text

    # Protect Georgian name forms
    for eng, geo in GEORGIAN_NAMES.items():
        if eng in protected:
            token = f"__NAME_{len(placeholders)}__"
            placeholders[token] = geo  # Store the Georgian form directly
            protected = protected.replace(eng, token)

    return protected, placeholders


def restore_terms(text, placeholders):
    """Restore placeholders back to original terms."""
    restored = text
    for token, original in placeholders.items():
        restored = restored.replace(token, original)
    return restored


def google_translate_chunk(text, retries=3):
    """Translate a chunk using GoogleTranslator."""
    for attempt in range(retries):
        try:
            translator = GoogleTranslator(source='en', target='ka')
            # Google has a limit of ~5000 chars per call
            if len(text) > 4800:
                # Split into smaller pieces
                parts = []
                lines = text.split('\n')
                batch = ""
                for line in lines:
                    if len(batch) + len(line) + 1 > 4800:
                        if batch:
                            parts.append(translator.translate(batch))
                        batch = line
                    else:
                        batch = batch + '\n' + line if batch else line
                if batch:
                    parts.append(translator.translate(batch))
                return '\n'.join(parts)
            else:
                return translator.translate(text)
        except Exception as e:
            print(f"  Google Translate attempt {attempt+1} failed: {e}")
            time.sleep(2)
    return text  # Return original if all attempts fail


def deepseek_refine(draft_ka, chunk_index):
    """Refine Georgian translation using DeepSeek API."""
    prompt = (
        "You are a scientific translator specializing in Georgian. "
        "Refine this Georgian translation of a Horizon Europe grant application. "
        "Keep technical terms in English (CDATA, CAII, Cell-DT, WP1-4, PTM, STED, KPI, TRL, ECS, SASP, GTU, ABET, FORGE-AI, ELISA, GCP, GDPR, IRB). "
        "Keep all numbers, citations [1]-[11], budget figures unchanged. "
        "Use proper Georgian scientific vocabulary. "
        "Return only the refined Georgian text, no commentary.\n\n"
        f"Draft translation:\n{draft_ka}"
    )

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 4096
    }

    for attempt in range(3):
        try:
            resp = requests.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=60)
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        except Exception as e:
            print(f"  DeepSeek attempt {attempt+1} failed: {e}")
            time.sleep(3)

    print(f"  DeepSeek failed for chunk {chunk_index}, using Google draft")
    return draft_ka


def translate_chunk(chunk_text, chunk_index):
    """Full translation pipeline for one chunk."""
    print(f"\n--- Chunk {chunk_index} ({len(chunk_text)} chars) ---")

    # Protect terms
    protected, placeholders = protect_terms(chunk_text)

    # Google Translate draft
    print("  Google Translate...")
    google_draft = google_translate_chunk(protected)

    # DeepSeek refinement
    print("  DeepSeek refinement...")
    refined = deepseek_refine(google_draft, chunk_index)

    # Restore protected terms
    result = restore_terms(refined, placeholders)
    return result


def split_into_chunks(lines, chunk_size=50):
    """Split lines into chunks, respecting block boundaries."""
    chunks = []
    current_chunk = []

    for line in lines:
        current_chunk.append(line)
        if len(current_chunk) >= chunk_size:
            # Try to end at a blank line or heading for clean breaks
            chunks.append('\n'.join(current_chunk))
            current_chunk = []

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks


# ============================================================
# DOCX GENERATION
# ============================================================

def add_page_numbers(doc):
    """Add 'გვერდი X / Y' footer."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Clear existing runs
        for run in para.runs:
            run.text = ""

        # Add field codes for page X of Y
        run = para.add_run()
        run.font.size = Pt(9)
        run.text = "გვერდი "

        # Add PAGE field
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run._r.append(instr)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_end)

        run2 = para.add_run(" / ")
        run2.font.size = Pt(9)

        run3 = para.add_run()
        run3.font.size = Pt(9)

        fld_begin2 = OxmlElement('w:fldChar')
        fld_begin2.set(qn('w:fldCharType'), 'begin')
        run3._r.append(fld_begin2)

        instr2 = OxmlElement('w:instrText')
        instr2.set(qn('xml:space'), 'preserve')
        instr2.text = ' NUMPAGES '
        run3._r.append(instr2)

        fld_end2 = OxmlElement('w:fldChar')
        fld_end2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld_end2)


def set_cell_shading(cell, fill_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def apply_inline_formatting(para, text, font_name='Sylfaen', font_size=11):
    """Parse **bold**, *italic* inline markdown and add runs."""
    # Pattern to match bold (**text**), italic (*text*), or plain text
    pattern = r'(\*\*.*?\*\*|\*[^*]+\*|[^*]+)'
    for match in re.finditer(pattern, text):
        segment = match.group(0)
        if segment.startswith('**') and segment.endswith('**'):
            run = para.add_run(segment[2:-2])
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        elif segment.startswith('*') and segment.endswith('*') and len(segment) > 2:
            run = para.add_run(segment[1:-1])
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(font_size)
        else:
            run = para.add_run(segment)
            run.font.name = font_name
            run.font.size = Pt(font_size)


def add_paragraph_with_format(doc, text, style_name, font_name='Sylfaen', font_size=11,
                               bold=False, italic=False, alignment=None, space_before=0,
                               space_after=6, left_indent=None):
    """Add a paragraph with specified formatting."""
    para = doc.add_paragraph(style=style_name)
    if alignment is not None:
        para.alignment = alignment
    pPr = para.paragraph_format
    pPr.space_before = Pt(space_before)
    pPr.space_after = Pt(space_after)
    if left_indent is not None:
        pPr.left_indent = Cm(left_indent)
    pPr.line_spacing = Pt(font_size * 1.15)

    apply_inline_formatting(para, text, font_name, font_size)

    for run in para.runs:
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    return para


def parse_and_add_table(doc, table_lines):
    """Parse a markdown table and add it to the document."""
    # Filter out separator rows (---|--- lines)
    data_rows = []
    header_done = False
    is_header_row = True

    for line in table_lines:
        stripped = line.strip()
        if not stripped or stripped.startswith('|'):
            # Check if it's a separator line
            if re.match(r'^\|[-| :]+\|$', stripped):
                is_header_row = False
                continue
            # Parse cells
            cells = [c.strip() for c in stripped.split('|')]
            cells = [c for c in cells if c != '' or len(cells) > 2]
            # Remove leading/trailing empty from split
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            if cells:
                data_rows.append((cells, is_header_row and not header_done))
                if is_header_row and not header_done:
                    header_done = True

    if not data_rows:
        return

    # Determine max columns
    max_cols = max(len(row[0]) for row in data_rows)

    table = doc.add_table(rows=len(data_rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, (cells, is_header) in enumerate(data_rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= max_cols:
                break
            cell = row.cells[col_idx]

            if is_header:
                set_cell_shading(cell, 'D9E1F2')

            # Clean markdown bold from cell text
            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)

            para = cell.paragraphs[0]
            run = para.add_run(clean_text)
            run.font.name = 'Sylfaen'
            run.font.size = Pt(9)
            if is_header:
                run.bold = True


def md_to_docx(translated_text, output_path):
    """Convert translated markdown to a formatted DOCX."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Add footer with page numbers
    add_page_numbers(doc)

    # Add first paragraph note
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note_para.add_run(
        "სამუშაო თარგმანი. ორიგინალი: EIC_Pathfinder_CDATA_PartB_v7.md"
    )
    note_run.font.name = 'Sylfaen'
    note_run.font.size = Pt(10)
    note_run.italic = True
    note_para.paragraph_format.space_after = Pt(12)

    lines = translated_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Code block (Gantt chart etc.) - render as monospace
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            # Add as a single paragraph with Courier
            for cl in code_lines:
                cp = doc.add_paragraph()
                cr = cp.add_run(cl)
                cr.font.name = 'Courier New'
                cr.font.size = Pt(7)
                cp.paragraph_format.space_after = Pt(0)
            continue

        # Detect table block
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            parse_and_add_table(doc, table_lines)
            continue

        # Headings
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(level=4)
            p.clear()
            run = p.add_run(text)
            run.font.name = 'Sylfaen'
            run.font.size = Pt(11)
            run.bold = True
            run.italic = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(level=3)
            p.clear()
            run = p.add_run(text)
            run.font.name = 'Sylfaen'
            run.font.size = Pt(12)
            run.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(text)
            run.font.name = 'Sylfaen'
            run.font.size = Pt(14)
            run.bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(level=1)
            p.clear()
            run = p.add_run(text)
            run.font.name = 'Sylfaen'
            run.font.size = Pt(16)
            run.bold = True
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        # Block quote
        if line.startswith('> '):
            text = line[2:].strip()
            # Remove ** for block quotes
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.name = 'Sylfaen'
            run.font.size = Pt(11)
            run.italic = True
            i += 1
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('  - ') or line.startswith('    - '):
            indent_level = 0
            stripped = line.lstrip()
            if line.startswith('    - '):
                indent_level = 2
            elif line.startswith('  - '):
                indent_level = 1
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.75 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(3)
            apply_inline_formatting(p, text, 'Sylfaen', 11)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(11 * 1.15)
            apply_inline_formatting(p, line.strip(), 'Sylfaen', 11)

        i += 1

    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    return output_path


def main():
    print("=== Georgian Translation Pipeline ===")
    print(f"Source: {SOURCE}")
    print(f"Output: {OUTPUT}")

    # Read source
    with open(SOURCE, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    print(f"Total lines: {len(lines)}")

    # Split into chunks of ~50 lines
    chunks = split_into_chunks(lines, chunk_size=50)
    print(f"Total chunks: {len(chunks)}")

    # Translate each chunk
    translated_chunks = []
    for idx, chunk in enumerate(chunks):
        translated = translate_chunk(chunk, idx + 1)
        translated_chunks.append(translated)

    # Assemble full translation
    full_translation = '\n'.join(translated_chunks)
    print(f"\nFull translation length: {len(full_translation)} chars")

    # Save draft translation for reference
    draft_path = "/home/oem/Desktop/Horizon/EIC_PartB_v7_KA_draft.txt"
    with open(draft_path, 'w', encoding='utf-8') as f:
        f.write(full_translation)
    print(f"Draft saved: {draft_path}")

    # Generate DOCX
    print("\nGenerating DOCX...")
    md_to_docx(full_translation, OUTPUT)

    # Report file size
    size = Path(OUTPUT).stat().st_size
    print(f"\nFile size: {size:,} bytes ({size/1024:.1f} KB)")
    if size > 60 * 1024:
        print("SUCCESS: File exceeds 60KB minimum")
    else:
        print(f"WARNING: File is {size/1024:.1f} KB (below 60KB minimum)")

    return size


if __name__ == '__main__':
    main()
